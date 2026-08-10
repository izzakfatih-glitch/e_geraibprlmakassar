"""
Aplikasi Web Penggabung Proposal PKKPRL
=========================================
Alur: Upload 2 PDF -> halaman Review (preview dokumen penuh + form koreksi
data) -> klik "Generate Dokumen Final" -> file Word diunduh.

HTML ditanam langsung di dalam file ini (tidak pakai folder templates/)
supaya tidak ada masalah TemplateNotFound di berbagai platform hosting.

CARA MENJALANKAN (LOKAL / TES):
    pip install -r requirements.txt --break-system-packages
    export ANTHROPIC_API_KEY="sk-ant-..."   # opsional, untuk fallback ekstraksi
    python3 app.py
    -> buka http://localhost:5000 di browser
"""
import os
import re
import uuid
import shutil
import traceback
from flask import Flask, request, render_template_string, send_file, after_this_request, session, redirect, url_for
import mammoth

# Login Google (authlib) bersifat OPSIONAL: kalau library ini gagal di-import
# atau gagal diinisialisasi (misal karena versi Python di server belum kompatibel),
# aplikasi tetap harus jalan normal -- cuma fitur Login yang otomatis nonaktif.
try:
    from authlib.integrations.flask_client import OAuth
    AUTHLIB_AVAILABLE = True
except Exception:
    traceback.print_exc()
    print("[startup] PERINGATAN: authlib gagal di-import, fitur Login Google dinonaktifkan.")
    OAuth = None
    AUTHLIB_AVAILABLE = False

from extract import extract_proposal_with_fallback, extract_laporan_with_fallback
from generate_docx import build_document
from review_fields import FIELD_GROUPS, form_field_name, get_value, apply_form_values
import job_store

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
# DATA_DIR: folder untuk data yang perlu BERTAHAN lintas-redeploy (riwayat &
# draft per-pengguna). Kalau env var DATA_DIR diisi (arahkan ke mount path
# Railway Volume, mis. "/data"), riwayat tidak akan hilang saat redeploy.
# Kalau tidak diisi, fallback ke folder aplikasi biasa (ephemeral -- hilang
# tiap redeploy, sama seperti sebelumnya).
DATA_DIR = os.environ.get("DATA_DIR") or BASE_DIR
UPLOAD_DIR = os.path.join(BASE_DIR, "uploads")
OUTPUT_DIR = os.path.join(BASE_DIR, "outputs")
JOBS_DIR = os.path.join(BASE_DIR, "jobs")
HISTORY_FILE = os.path.join(DATA_DIR, "history.jsonl")
DRAFTS_DIR = os.path.join(DATA_DIR, "drafts")
DRAFT_MAX_AGE_DAYS = 4  # riwayat isian per-pengguna otomatis terhapus setelah sekian hari
STAFF_SHEET_CSV_URL = os.environ.get("STAFF_SHEET_CSV_URL") or (
    "https://docs.google.com/spreadsheets/d/1X7YS72vG6wF0XqxClfeZkc_zpeJxaGKfuB3Nw7M1QXM"
    "/export?format=csv&gid=0"
)
os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)
os.makedirs(JOBS_DIR, exist_ok=True)
os.makedirs(DRAFTS_DIR, exist_ok=True)
print(f"[startup] DATA_DIR (riwayat/draft tersimpan di sini) = {DATA_DIR}")


MAX_CONTENT_LENGTH = 30 * 1024 * 1024  # 30 MB batas unggah per file gabungan

app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = MAX_CONTENT_LENGTH
# PENTING: SECRET_KEY wajib SAMA di semua worker proses (Procfile pakai -w 4,
# artinya ada 4 proses gunicorn berbeda) supaya cookie sesi yang ditandatangani
# oleh 1 worker bisa tetap dikenali/valid saat request berikutnya ditangani
# worker LAIN -- kalau tidak, user akan dianggap "belum login" lagi setiap kali
# request-nya kebetulan jatuh ke worker yang berbeda (bug: login muncul lagi
# saat pindah halaman). Paling aman: set FLASK_SECRET_KEY di environment
# variable Railway dengan nilai TETAP. Kalau belum diset, fallback di bawah ini
# tetap KONSISTEN antar-worker (bukan acak per-proses) supaya tidak bug,
# walaupun secara keamanan sebaiknya tetap diganti dengan nilai rahasia sendiri.
app.secret_key = os.environ.get("FLASK_SECRET_KEY") or "e-gerai-kkprl-bprl-makassar-2026-ganti-dgn-secret-sendiri"
# Sesi login dibuat bertahan SANGAT LAMA (10 tahun / praktis tanpa batas waktu)
# supaya staf cukup login SEKALI di awal, lalu semua menu langsung bisa dipakai
# terus-menerus tanpa diminta login ulang, sampai mereka klik "Keluar" sendiri.
import datetime as _dt
app.config["PERMANENT_SESSION_LIFETIME"] = _dt.timedelta(days=3650)

# ---- Konfigurasi OAuth Google (Sign in with Google) ----
# GOOGLE_CLIENT_ID & GOOGLE_CLIENT_SECRET wajib diisi lewat environment variable
# di Railway (dibuat lewat Google Cloud Console -> OAuth 2.0 Client IDs).
google_oauth = None
if AUTHLIB_AVAILABLE:
    try:
        oauth = OAuth(app)
        google_oauth = oauth.register(
            name="google",
            client_id=os.environ.get("GOOGLE_CLIENT_ID"),
            client_secret=os.environ.get("GOOGLE_CLIENT_SECRET"),
            server_metadata_url="https://accounts.google.com/.well-known/openid-configuration",
            client_kwargs={"scope": "openid email profile"},
        )
    except Exception:
        traceback.print_exc()
        print("[startup] PERINGATAN: inisialisasi OAuth Google gagal, fitur Login dinonaktifkan.")
        google_oauth = None

print(f"[startup] BASE_DIR = {BASE_DIR}")
print(f"[startup] Isi BASE_DIR = {os.listdir(BASE_DIR)}")


# ---------------------------------------------------------------------------
# Riwayat penggunaan (history) -- catatan RINGAN saja (metadata: nama pemohon,
# nama perusahaan, waktu, dan siapa yang memproses). TIDAK menyimpan file
# dokumennya sama sekali, sesuai kebijakan privasi aplikasi ini.
#
# CATATAN PENTING: file history.jsonl ini disimpan di disk lokal container.
# Di Railway, disk lokal bersifat SEMENTARA -- akan terhapus setiap kali
# aplikasi di-redeploy/restart, KECUALI Anda memasang "Volume" (persistent
# storage) di pengaturan Railway dan mengarahkannya ke folder BASE_DIR ini.
# ---------------------------------------------------------------------------
import json
import datetime

WITA = datetime.timezone(datetime.timedelta(hours=8))  # Waktu Indonesia Tengah (Makassar)


def log_history_entry(nama_pemohon, nama_perusahaan, sumber):
    entry = {
        "waktu": datetime.datetime.now(WITA).strftime("%Y-%m-%d %H:%M:%S"),
        "nama_pemohon": nama_pemohon or "-",
        "nama_perusahaan": nama_perusahaan or "-",
        "sumber": sumber,
        "diproses_oleh": (session.get("user") or {}).get("name", "Belum Login"),
        "email": (session.get("user") or {}).get("email", ""),
    }
    try:
        with open(HISTORY_FILE, "a", encoding="utf-8") as f:
            f.write(json.dumps(entry, ensure_ascii=False) + "\n")
    except Exception:
        traceback.print_exc()


def read_history(limit=200):
    if not os.path.exists(HISTORY_FILE):
        return []
    entries = []
    try:
        with open(HISTORY_FILE, "r", encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                if not line:
                    continue
                try:
                    entries.append(json.loads(line))
                except json.JSONDecodeError:
                    continue
    except Exception:
        traceback.print_exc()
        return []
    return list(reversed(entries))[:limit]


# ---------------------------------------------------------------------------
# Riwayat pengisian PER-PENGGUNA ("Riwayat Saya") -- beda dari riwayat global
# (history.jsonl) yang isinya cuma metadata ringkas dan permanen. Ini
# menyimpan SELURUH data isian form (prop_data, tanpa gambar) supaya staf
# bisa melanjutkan isian yang belum selesai tanpa mengulang dari awal.
# Otomatis terhapus sendiri setelah DRAFT_MAX_AGE_DAYS hari supaya tidak
# membebani penyimpanan server. Riwayat global (semua pengguna) TIDAK
# terpengaruh oleh penghapusan ini.
# ---------------------------------------------------------------------------
def _safe_kode(kode_nama):
    return re.sub(r"[^A-Za-z0-9_-]", "_", str(kode_nama or "anon"))


def _draft_images_dir(safe_kode, job_id):
    return os.path.join(DRAFTS_DIR, f"{safe_kode}__{job_id}_images")


def save_draft(kode_nama, job_id, prop_data, prop_images=None):
    """Simpan seluruh isian form (prop_data) SEKALIGUS gambar yang sudah
    diunggah (prop_images) supaya tidak hilang kalau pengguna menutup
    halaman / melanjutkan isian di lain waktu. Gambar disimpan sebagai file
    terpisah di folder <kode>__<job_id>_images/ (bukan di-embed ke JSON)
    supaya file draft tetap ringan dan cepat dibaca."""
    try:
        safe_kode = _safe_kode(kode_nama)

        images_meta = []
        if prop_images:
            img_dir = _draft_images_dir(safe_kode, job_id)
            os.makedirs(img_dir, exist_ok=True)
            for i, im in enumerate(prop_images):
                ext = im.get("ext") or "png"
                fname = f"{i}_{im.get('tag', 'lainnya')}.{ext}"
                with open(os.path.join(img_dir, fname), "wb") as imgf:
                    imgf.write(im["bytes"])
                images_meta.append({"tag": im.get("tag", "lainnya"), "filename": fname})

        entry = {
            "waktu": datetime.datetime.now(WITA).strftime("%Y-%m-%d %H:%M:%S"),
            "job_id": job_id,
            "kode_nama": kode_nama,
            "nama_pemohon": prop_data.get("Nama Pemohon", ""),
            "nama_perusahaan": prop_data.get("Nama Perusahaan/Instansi", ""),
            "prop_data": prop_data,
            "prop_images_meta": images_meta,
        }
        path = os.path.join(DRAFTS_DIR, f"{safe_kode}__{job_id}.json")
        with open(path, "w", encoding="utf-8") as f:
            json.dump(entry, f, ensure_ascii=False)
    except Exception:
        traceback.print_exc()


def load_draft_images(kode_nama, job_id):
    """Baca kembali gambar yang tersimpan untuk satu draft (hasil save_draft
    sebelumnya) dalam format list [{'tag','bytes','ext'}] -- format yang
    sama seperti prop_images biasa, siap dipakai langsung oleh
    build_document()."""
    safe_kode = _safe_kode(kode_nama)
    draft = load_draft(kode_nama, job_id)
    if not draft:
        return []
    images = []
    img_dir = _draft_images_dir(safe_kode, job_id)
    for meta in draft.get("prop_images_meta", []) or []:
        fpath = os.path.join(img_dir, meta.get("filename", ""))
        if not os.path.exists(fpath):
            continue
        try:
            with open(fpath, "rb") as f:
                data_bytes = f.read()
            ext = os.path.splitext(meta["filename"])[1].lstrip(".").lower() or "png"
            images.append({"tag": meta.get("tag", "lainnya"), "bytes": data_bytes, "ext": ext})
        except Exception:
            continue
    return images


def merge_draft_images(kode_nama, job_id, prop_images):
    """Gabungkan gambar yang BARU diunggah pada submit ini (prop_images)
    dengan gambar yang sudah tersimpan sebelumnya lewat tombol \"Simpan\"
    (job_id yang sama) -- supaya kalau pengguna sudah pernah mengunggah
    gambar tapi lupa mengunggah ulang di sesi berikutnya, gambar lama tetap
    dipakai (tidak hilang). Tag yang di-upload ulang pada submit ini akan
    MENGGANTIKAN (bukan menumpuk) gambar lama dengan tag yang sama."""
    if not kode_nama or not job_id:
        return prop_images
    saved = load_draft_images(kode_nama, job_id)
    if not saved:
        return prop_images
    new_tags = {im.get("tag") for im in prop_images}
    merged = list(prop_images)
    for im in saved:
        if im.get("tag") not in new_tags:
            merged.append(im)
    return merged


def cleanup_expired_drafts():
    try:
        cutoff = time.time() - (DRAFT_MAX_AGE_DAYS * 86400)
        for fname in os.listdir(DRAFTS_DIR):
            path = os.path.join(DRAFTS_DIR, fname)
            try:
                if os.path.getmtime(path) < cutoff:
                    os.remove(path)
                    # folder gambar draft ini (kalau ada) ikut dihapus
                    if fname.endswith(".json"):
                        img_dir = path[:-len(".json")] + "_images"
                        shutil.rmtree(img_dir, ignore_errors=True)
            except OSError:
                continue
    except Exception:
        traceback.print_exc()


def read_user_drafts(kode_nama, limit=50):
    cleanup_expired_drafts()
    safe_kode = _safe_kode(kode_nama)
    prefix = f"{safe_kode}__"
    entries = []
    try:
        for fname in os.listdir(DRAFTS_DIR):
            if not fname.startswith(prefix):
                continue
            path = os.path.join(DRAFTS_DIR, fname)
            try:
                with open(path, "r", encoding="utf-8") as f:
                    entries.append(json.load(f))
            except Exception:
                continue
    except Exception:
        traceback.print_exc()
        return []
    entries.sort(key=lambda e: e.get("waktu", ""), reverse=True)
    return entries[:limit]


def load_draft(kode_nama, job_id):
    safe_kode = _safe_kode(kode_nama)
    path = os.path.join(DRAFTS_DIR, f"{safe_kode}__{job_id}.json")
    if not os.path.exists(path):
        return None
    try:
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        traceback.print_exc()
        return None


# Login sederhana pegawai -- kredensial diambil LANGSUNG dari Google Sheet
# ("Kode Nama" = username, "Kode petugas" = password), di-cache sebentar
# supaya tidak nge-fetch Google di setiap request. Kalau ada pegawai baru
# ditambahkan di sheet, otomatis kepakai begitu cache kedaluwarsa (tanpa
# perlu ubah kode / redeploy).
# ---------------------------------------------------------------------------
import csv
import io as _io
import time

_staff_cache = {"data": {}, "fetched_at": 0}
STAFF_CACHE_TTL = 300  # detik (5 menit)


def parse_koordinat_file(file_storage):
    """Baca titik koordinat otomatis dari file Excel (.xlsx), CSV, Word
    (.docx), atau gambar (PNG/JPG) yang diupload -- 2 kolom pertama tiap
    baris dianggap Longitude dan Latitude. Baris/sel yang bukan angka
    otomatis dilewati (termasuk baris judul/header). Mengembalikan tuple
    (list [nomor, longitude, latitude], pesan_status_atau_None)."""
    filename = (file_storage.filename or "").lower()
    data = file_storage.read()
    hasil = []
    pesan = None

    def add_pair(a, b):
        try:
            float(str(a).replace(",", "."))
            float(str(b).replace(",", "."))
        except (TypeError, ValueError):
            return
        hasil.append([str(len(hasil) + 1), str(a).strip(), str(b).strip()])

    try:
        if filename.endswith((".xlsx", ".xlsm")):
            import openpyxl
            wb = openpyxl.load_workbook(_io.BytesIO(data), data_only=True)
            for ws in wb.worksheets:
                for row in ws.iter_rows(values_only=True):
                    vals = [v for v in row if v is not None and str(v).strip() != ""]
                    if len(vals) >= 2:
                        add_pair(vals[0], vals[1])
        elif filename.endswith(".csv"):
            text = data.decode("utf-8-sig", errors="ignore")
            for line in text.splitlines():
                parts = [p for p in re.split(r"[,;\t]", line.strip()) if p.strip()]
                if len(parts) >= 2:
                    add_pair(parts[0], parts[1])
        elif filename.endswith(".docx"):
            from docx import Document as _DocxDocument
            doc = _DocxDocument(_io.BytesIO(data))
            for table in doc.tables:
                for row in table.rows:
                    vals = [c.text.strip() for c in row.cells if c.text.strip()]
                    if len(vals) >= 2:
                        add_pair(vals[0], vals[1])
            if not hasil:
                for para in doc.paragraphs:
                    nums = re.findall(r"-?\d+[.,]\d+", para.text)
                    if len(nums) >= 2:
                        add_pair(nums[0], nums[1])
        elif filename.endswith((".png", ".jpg", ".jpeg")):
            hasil, pesan = parse_koordinat_dari_gambar(data, filename)
    except Exception:
        traceback.print_exc()
        return [], "Terjadi kesalahan saat membaca file. Silakan isi koordinat manual."

    return hasil, pesan


def parse_koordinat_dari_gambar(image_bytes, filename):
    """Baca titik koordinat dari gambar (mis. screenshot tabel koordinat atau
    peta) memakai Claude Vision API. Kalau ANTHROPIC_API_KEY tidak diset,
    otomatis dilewati (dikembalikan list kosong, tidak error) -- konsisten
    dengan pola fallback LLM lain di aplikasi ini (lihat llm_fallback.py).
    Mengembalikan tuple (list_titik, pesan_status) supaya pemanggil bisa
    kasih tahu user kenapa gagal/berhasil, bukan cuma diam-diam kosong."""
    from llm_fallback import api_key_available
    if not api_key_available():
        return [], "ANTHROPIC_API_KEY belum diset di server -- pembacaan otomatis dari gambar dilewati."

    MAX_BYTES = 5 * 1024 * 1024  # 5 MB, batas aman untuk request Claude API
    if len(image_bytes) > MAX_BYTES:
        return [], "Ukuran gambar terlalu besar (maks 5 MB) untuk dibaca otomatis."

    try:
        import base64
        from anthropic import Anthropic
        client = Anthropic()
        ext = "jpeg" if filename.lower().endswith((".jpg", ".jpeg")) else "png"
        b64 = base64.b64encode(image_bytes).decode("utf-8")
        resp = client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=1500,
            messages=[{
                "role": "user",
                "content": [
                    {"type": "image", "source": {"type": "base64", "media_type": f"image/{ext}", "data": b64}},
                    {"type": "text", "text": (
                        "Gambar ini berisi daftar titik koordinat (longitude, latitude) batas area, "
                        "bisa berupa tabel, daftar teks, atau anotasi pada peta/citra satelit. Baca "
                        "semua titik koordinat yang terlihat dalam format desimal derajat (bukan "
                        "derajat-menit-detik), urut sesuai urutan tampil di gambar (atas ke bawah, "
                        "atau sesuai penomoran titik kalau ada). Kalau koordinat tertulis dalam format "
                        "derajat-menit-detik (mis. 119°27'30\"E), konversi ke desimal derajat. "
                        "Balas HANYA dengan JSON array polos (tanpa markdown/backtick/penjelasan "
                        "apa pun), format persis: "
                        '[["longitude1","latitude1"],["longitude2","latitude2"],...]. '
                        "Kalau tidak ada titik koordinat yang terbaca sama sekali, balas: []"
                    )},
                ],
            }],
        )
        text = "".join(b.text for b in resp.content if b.type == "text").strip()
        text = text.replace("```json", "").replace("```", "").strip()
        # Kalau model tetap menyelipkan teks lain di luar instruksi, ambil
        # bagian yang berbentuk array JSON saja (dari '[' pertama sampai ']' terakhir).
        m = re.search(r"\[.*\]", text, re.DOTALL)
        if m:
            text = m.group(0)
        pairs = json.loads(text)

        hasil = []
        for p in pairs:
            if not (isinstance(p, (list, tuple)) and len(p) >= 2):
                continue
            try:
                lon = float(str(p[0]).replace(",", "."))
                lat = float(str(p[1]).replace(",", "."))
            except (TypeError, ValueError):
                continue
            # Validasi kasar: koordinat wilayah Indonesia (longitude 90-142, latitude -12..8)
            if not (90 <= lon <= 142 and -12 <= lat <= 8):
                continue
            hasil.append([str(len(hasil) + 1), str(p[0]).strip(), str(p[1]).strip()])

        if not hasil:
            return [], "Gambar berhasil dibaca AI, tapi tidak ditemukan titik koordinat yang valid di dalamnya."
        return hasil, f"Berhasil membaca {len(hasil)} titik koordinat dari gambar."
    except Exception:
        traceback.print_exc()
        return [], "Terjadi kesalahan saat membaca gambar lewat AI. Silakan isi koordinat manual."


def fetch_staff_list(force=False):
    now = time.time()
    if not force and _staff_cache["data"] and (now - _staff_cache["fetched_at"]) < STAFF_CACHE_TTL:
        return _staff_cache["data"]

    import urllib.request

    try:
        req = urllib.request.Request(STAFF_SHEET_CSV_URL, headers={"User-Agent": "Mozilla/5.0"})
        with urllib.request.urlopen(req, timeout=8) as resp:
            raw = resp.read().decode("utf-8-sig")
        reader = csv.DictReader(_io.StringIO(raw))
        staff = {}
        for row in reader:
            nama = (row.get("Nama") or "").strip()
            kode_nama = (row.get("Kode Nama") or "").strip()
            kode_petugas = (row.get("Kode petugas") or "").strip()
            if nama and kode_nama and kode_petugas:
                staff[kode_nama] = {"nama": nama, "password": kode_petugas}
        if staff:
            _staff_cache["data"] = staff
            _staff_cache["fetched_at"] = now
        return _staff_cache["data"]
    except Exception:
        traceback.print_exc()
        # Kalau fetch gagal (mis. sedang tidak ada internet), tetap pakai
        # cache lama yang masih ada di memori (kalau ada), daripada error total.
        return _staff_cache["data"]

BASE_CSS = """
:root { --navy:#1F4E79; --bg:#f4f6f8; }
* { box-sizing: border-box; }
body {
  font-family: -apple-system, Segoe UI, Roboto, Arial, sans-serif;
  background: var(--bg); margin: 0; padding: 0; color: #222;
}
.wrap { max-width: 640px; margin: 0 auto; padding: 32px 20px 60px; }
.wrap.wide { max-width: 900px; }
h1 { color: var(--navy); font-size: 22px; margin-bottom: 4px; }
p.sub { color: #555; margin-top: 0; font-size: 14px; }
.card { background: #fff; border-radius: 12px; padding: 24px;
        box-shadow: 0 2px 10px rgba(0,0,0,0.06); margin-top: 20px; }
label { display:block; font-weight:600; margin-bottom:6px; margin-top:14px; font-size: 13px; color:#333; }
input[type=file] { display:block; width:100%; padding: 10px; border: 1px dashed #aaa;
                    border-radius: 8px; background:#fafafa; font-size: 13px; }
input[type=text] { display:block; width:100%; padding: 9px 10px; border: 1px solid #ccc;
                    border-radius: 6px; font-size: 13px; }
button, .btn { margin-top: 20px; width: 100%; background: var(--navy); color:#fff;
       border:none; padding: 14px; border-radius: 8px; font-size: 15px;
       font-weight: 600; cursor:pointer; display:block; text-align:center; text-decoration:none; }
button:hover, .btn:hover { background:#163a5c; }
.note { font-size: 12px; color:#777; margin-top: 16px; line-height:1.5; }
.flash { background:#fff3cd; border:1px solid #ffe08a; padding:12px 14px;
         border-radius:8px; margin-top:16px; font-size:13px; color:#7a5b00; }
.spinner { display:none; text-align:center; margin-top:18px; font-size:13px; color:var(--navy); }
details { border:1px solid #e2e2e2; border-radius:8px; margin-top:12px; padding: 4px 14px 12px; background:#fbfbfd; }
summary { cursor:pointer; font-weight:700; color:var(--navy); padding:10px 0; font-size:14px; }
.preview-box { border:1px solid #ddd; border-radius:10px; background:#fff; padding: 30px 40px;
               max-height: 720px; overflow-y: auto; font-size: 14px; line-height:1.5; }
.preview-box table { border-collapse: collapse; width:100%; margin: 10px 0; }
.preview-box td { border:1px solid #999; padding:5px 8px; vertical-align: top; }
.preview-box img { max-width: 100%; height:auto; margin: 8px 0; }
.preview-box p { margin: 6px 0; }
.top-actions { position: sticky; top:0; background: var(--bg); padding: 10px 0; z-index:5; }
"""

LANDING_CSS = """
:root {
  --navy:#123A63; --blue:#1E63C7; --blue2:#2F7FE0; --bg:#eef3f8;
  --line:#e3e9f0; --ink:#1c2b3a; --muted:#5b6b7c; --green:#1f9d55; --green-bg:#eafaf0;
}
* { box-sizing: border-box; }
html, body { margin:0; padding:0; }
body { font-family: -apple-system, "Segoe UI", Roboto, Arial, sans-serif; background: var(--bg); color: var(--ink); }
a { text-decoration:none; }

/* ---- Header ---- */
.site-header { background:#fff; border-bottom:1px solid var(--line); padding:18px 40px;
  display:flex; align-items:center; justify-content:space-between; flex-wrap:wrap; gap:20px; }
.brand { display:flex; align-items:center; gap:14px; }
.brand-badge { width:44px; height:44px; border-radius:50%; flex:none;
  background: conic-gradient(from 220deg, var(--blue2), var(--navy), var(--blue2));
  display:flex; align-items:center; justify-content:center; }
.brand-badge svg { width:22px; height:22px; }
.brand-text .t1 { font-size:19px; font-weight:800; color:var(--navy); line-height:1.1; }
.brand-text .t2 { font-size:11.5px; font-weight:700; color:var(--blue); letter-spacing:.03em; }
.header-right { display:flex; align-items:center; gap:18px; }
.gov-block { display:flex; align-items:center; gap:12px; }
.gov-badge { width:52px; height:52px; border-radius:50%; flex:none; background:#f0f4f9;
  border:1px solid var(--line); display:flex; align-items:center; justify-content:center; padding:6px; }
.brand-logo-img { height:76px; width:auto; display:block; }
.gov-badge img { max-width:100%; max-height:100%; object-fit:contain; }
.gov-badge.rect { width:auto; height:48px; min-width:48px; border-radius:8px; padding:5px 10px;
  background:#fff; border:1px solid var(--line); }
.gov-badge.rect img { height:100%; width:auto; }
.gov-text { font-size:12px; line-height:1.35; color:var(--navy); font-weight:700; max-width:170px; }
.header-divider { width:1px; align-self:stretch; background:var(--line); }
.brl-text .t1 { font-size:15px; font-weight:800; color:var(--navy); line-height:1.1; }
.brl-text .t2 { font-size:10.5px; font-weight:700; color:var(--muted); letter-spacing:.05em; }

.navbar { display:flex; align-items:center; gap:26px; flex-wrap:wrap; }
.nav-link { display:flex; align-items:center; gap:6px; font-size:14px; font-weight:700;
  color:#334862; padding:8px 2px; border-bottom:2px solid transparent; white-space:nowrap; }
.nav-link svg { width:16px; height:16px; }
.nav-link:hover { color:var(--blue); }
.nav-link.active { color:var(--blue); border-bottom-color:var(--blue); }
.login-btn { display:flex; align-items:center; gap:8px; font-size:13.5px; font-weight:700;
  color:var(--blue); border:1.5px solid #cfe0f5; background:#f3f8ff; padding:9px 18px;
  border-radius:999px; white-space:nowrap; }
.login-btn svg { width:16px; height:16px; }
.login-btn:hover { background:#e8f1fd; }
.user-badge { display:flex; align-items:center; gap:10px; font-size:13.5px; }
.user-avatar { width:32px; height:32px; border-radius:50%; border:2px solid #e8f1fd; flex:none; }
.user-greet { font-weight:700; color:var(--navy); white-space:nowrap; }
.logout-link { font-size:12px; font-weight:700; color:var(--blue); border:1.5px solid #cfe0f5;
  background:#f3f8ff; padding:7px 14px; border-radius:999px; white-space:nowrap; }
.logout-link:hover { background:#e8f1fd; }
@media (max-width: 1180px) { .navbar { order:3; width:100%; justify-content:center; padding-top:10px;
  border-top:1px solid var(--line); } }

/* ---- Hero ---- */
.hero { position:relative; overflow:hidden; background:linear-gradient(135deg,#eaf2fb 0%,#cfe1f6 55%,#a9cdec 100%);
  padding:36px 32px 50px; }
.hero-image { padding:0; }
.hero-image .hero-inner { display:block; max-width:1600px; margin:0 auto; }
.hero-banner-img { width:100%; height:auto; display:block; }
.hero-inner { max-width:1600px; margin:0 auto; display:flex; align-items:center; gap:36px; flex-wrap:wrap; }

.hero-logo { flex:0 0 auto; }
.hero-logo img { height:150px; width:auto; display:block; }
@media (max-width: 1300px) { .hero-logo img { height:110px; } }
@media (max-width: 700px) { .hero-logo { width:100%; text-align:center; } .hero-logo img { height:90px; } }

.hero-copy { flex:1 1 380px; min-width:280px; border-left:3px solid #ffb020; padding-left:24px; }
.hero-copy h1 { font-size:30px; font-weight:800; color:var(--navy); margin:0 0 10px; letter-spacing:-.01em; }
.hero-copy p { font-size:14.5px; color:#33495e; max-width:560px; line-height:1.6; margin:0 0 22px; }
.feature-row { display:flex; gap:22px; flex-wrap:wrap; }
.feature { display:flex; align-items:center; gap:10px; }
.feature .ic { width:38px; height:38px; border-radius:10px; background:#fff; flex:none;
  display:flex; align-items:center; justify-content:center; box-shadow:0 2px 6px rgba(18,58,99,.12); }
.feature .ic svg { width:19px; height:19px; color:var(--blue); }
.feature span { font-size:12.5px; font-weight:700; color:var(--navy); line-height:1.25; display:block; max-width:110px; }

.hero-art { flex:0 0 auto; display:flex; justify-content:center; }
.illust-wrap { position:relative; width:300px; }
.illust-screen { background:linear-gradient(160deg,#1c4faa,#0d2a5c); border-radius:16px;
  aspect-ratio:4/3; display:flex; align-items:center; justify-content:center;
  box-shadow:0 20px 40px rgba(10,30,60,.28); position:relative; overflow:hidden; }
.illust-stand { width:110px; height:10px; background:#c3cad2; border-radius:0 0 8px 8px; margin:0 auto; }
.illust-ai-badge { position:absolute; top:-14px; right:-10px; z-index:2;
  background:linear-gradient(135deg,#ffc857,#ff8a00); color:#fff; font-weight:900; font-size:17px;
  width:46px; height:46px; border-radius:12px; display:flex; align-items:center; justify-content:center;
  box-shadow:0 8px 16px rgba(255,140,0,.35); transform:rotate(-6deg); }
.illust-mag { position:absolute; bottom:-10px; right:18px; width:44px; height:44px; border-radius:50%;
  background:#fff; display:flex; align-items:center; justify-content:center; box-shadow:0 6px 14px rgba(18,58,99,.25); }
.illust-mag svg { width:22px; height:22px; color:var(--blue); }
@media (max-width: 700px) { .illust-wrap { width:220px; } }

/* ---- Main grid ---- */
.main-wrap { max-width:1600px; margin:-24px auto 0; padding:0 40px 40px; position:relative; z-index:2; }
.grid { display:grid; grid-template-columns:1fr 1fr 340px; gap:26px; align-items:start; }
@media (max-width: 980px) { .grid { grid-template-columns:1fr; } }

.upload-card { background:#fff; border-radius:16px; padding:24px; box-shadow:0 6px 24px rgba(18,58,99,.08); }
.step-head { display:flex; align-items:center; gap:12px; margin-bottom:6px; }
.step-num { width:26px; height:26px; border-radius:50%; background:var(--blue); color:#fff;
  font-size:13px; font-weight:800; display:flex; align-items:center; justify-content:center; flex:none; }
.step-icon { width:44px; height:44px; border-radius:12px; background:#eaf1fc; flex:none;
  display:flex; align-items:center; justify-content:center; }
.step-icon svg { width:22px; height:22px; color:var(--blue); }
.step-title { font-size:15.5px; font-weight:800; color:var(--navy); }
.step-desc { font-size:12.5px; color:var(--muted); margin:2px 0 16px; }

.dropzone { border:2px dashed #b9cbe0; border-radius:12px; background:#f7fafd; padding:26px 14px;
  text-align:center; cursor:pointer; transition:.15s; }
.dropzone:hover, .dropzone.dragover { border-color:var(--blue); background:#eef5fd; }
.dropzone svg { width:30px; height:30px; color:var(--blue); margin-bottom:8px; }
.dropzone .dz-title { font-size:13.5px; font-weight:700; color:var(--navy); }
.dropzone .dz-sub { font-size:12px; color:var(--muted); margin-top:2px; }
.dropzone .dz-max { font-size:11px; color:var(--muted); margin-top:8px; }
.dropzone input[type=file] { display:none; }

.file-chip { display:none; margin-top:12px; align-items:center; justify-content:space-between;
  background:var(--green-bg); border:1px solid #cdeedb; border-radius:10px; padding:9px 12px; }
.file-chip.show { display:flex; }
.file-chip .fc-left { display:flex; align-items:center; gap:8px; min-width:0; }
.file-chip .fc-left svg { width:16px; height:16px; color:var(--green); flex:none; }
.file-chip .fc-name { font-size:12.5px; font-weight:600; color:#1c2b3a; overflow:hidden;
  text-overflow:ellipsis; white-space:nowrap; max-width:160px; }
.file-chip .fc-size { font-size:11.5px; color:var(--muted); flex:none; margin-left:8px; }
.file-chip .fc-remove { background:none; border:none; cursor:pointer; color:var(--muted);
  font-size:15px; line-height:1; padding:2px 4px; }

.gen-btn { grid-column:1 / span 2; }
@media (max-width: 980px) { .gen-btn { grid-column:1; } }
.gen-btn button { width:100%; background:linear-gradient(90deg,var(--blue2),var(--navy)); color:#fff;
  border:none; padding:16px; border-radius:12px; font-size:15.5px; font-weight:800; cursor:pointer;
  display:flex; align-items:center; justify-content:center; gap:8px; box-shadow:0 6px 18px rgba(30,99,199,.3); }
.gen-btn button svg { width:20px; height:20px; flex:none; }
.gen-btn button:hover { filter:brightness(1.05); }
.gen-btn button:disabled { opacity:.55; cursor:not-allowed; }
.gen-btn .gen-note { text-align:center; font-size:11.5px; color:var(--muted); margin-top:10px; }
.gen-btn .spinner { display:none; text-align:center; font-size:12.5px; color:var(--blue); margin-top:10px; font-weight:700; }

.flow-card { background:#fff; border-radius:16px; padding:22px 20px; box-shadow:0 6px 24px rgba(18,58,99,.08); }
.flow-card h3 { font-size:14px; font-weight:800; color:var(--blue); margin:0 0 16px; }
.flow-step { display:flex; gap:12px; position:relative; padding-bottom:22px; }
.flow-step:last-child { padding-bottom:0; }
.flow-step::before { content:""; position:absolute; left:15px; top:34px; bottom:0; width:2px;
  background:repeating-linear-gradient(to bottom, #c9d6e6 0 4px, transparent 4px 8px); }
.flow-step:last-child::before { display:none; }
.flow-dot { width:32px; height:32px; border-radius:50%; flex:none; display:flex; align-items:center;
  justify-content:center; font-weight:800; font-size:13px; color:#fff; z-index:1; }
.flow-dot svg { width:16px; height:16px; }
.flow-step:nth-of-type(1) .flow-dot { background:var(--blue); }
.flow-step:nth-of-type(2) .flow-dot { background:#3fa7d6; }
.flow-step:nth-of-type(3) .flow-dot { background:#7e5bd6; }
.flow-step:nth-of-type(4) .flow-dot { background:linear-gradient(135deg,#ffc857,#ff8a00); }
.flow-body .ft { font-size:13px; font-weight:800; color:var(--navy); }
.flow-body .fd { font-size:11.5px; color:var(--muted); margin-top:2px; line-height:1.4; }

.error-banner { grid-column:1/-1; background:#fff3cd; border:1px solid #ffe08a; padding:12px 16px;
  border-radius:10px; font-size:13px; color:#7a5b00; margin-bottom:2px; }

/* ---- Trust strip ---- */
.trust-strip { max-width:1600px; margin:0 auto 40px; padding:0 40px; }
.trust-inner { background:#fff; border-radius:16px; padding:22px 26px; box-shadow:0 6px 24px rgba(18,58,99,.06);
  display:flex; align-items:center; flex-wrap:wrap; gap:26px; justify-content:space-between; }
.trust-item { display:flex; align-items:center; gap:10px; min-width:170px; }
.trust-item .ic { width:38px; height:38px; border-radius:10px; background:#eaf1fc; flex:none;
  display:flex; align-items:center; justify-content:center; }
.trust-item .ic svg { width:18px; height:18px; color:var(--blue); }
.trust-item .tt { font-size:12.5px; font-weight:800; color:var(--navy); }
.trust-item .td { font-size:11px; color:var(--muted); }
.trust-brand { display:flex; align-items:center; gap:10px; }
.trust-brand .tt { font-size:14px; font-weight:800; color:var(--navy); }
.trust-brand .td { font-size:11px; color:var(--muted); max-width:220px; }
"""

ICONS = {
  "doc": '<svg width="20" height="20" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><path d="M14 2H6a2 2 0 0 0-2 2v16a2 2 0 0 0 2 2h12a2 2 0 0 0 2-2V8z"/><path d="M14 2v6h6"/></svg>',
  "wave": '<svg width="20" height="20" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><path d="M2 12c2-3 4-3 6 0s4 3 6 0 4-3 6 0"/><path d="M2 18c2-3 4-3 6 0s4 3 6 0 4-3 6 0"/></svg>',
  "upload": '<svg width="20" height="20" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><path d="M12 16V4M12 4l-4 4M12 4l4 4"/><path d="M4 16v3a2 2 0 0 0 2 2h12a2 2 0 0 0 2-2v-3"/></svg>',
  "bolt": '<svg width="20" height="20" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><path d="M13 2 3 14h7l-1 8 10-12h-7z"/></svg>',
  "shield": '<svg width="20" height="20" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><path d="M12 22s8-4 8-10V5l-8-3-8 3v7c0 6 8 10 8 10z"/></svg>',
  "shield-check": '<svg width="20" height="20" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><path d="M12 22s8-4 8-10V5l-8-3-8 3v7c0 6 8 10 8 10z"/><path d="m9 12 2 2 4-4"/></svg>',
  "lock": '<svg width="20" height="20" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><rect x="3" y="11" width="18" height="10" rx="2"/><path d="M7 11V7a5 5 0 0 1 10 0v4"/></svg>',
  "check-circle": '<svg width="20" height="20" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><path d="M22 11.08V12a10 10 0 1 1-5.93-9.14"/><path d="m22 4-10 10-3-3"/></svg>',
  "x": '<svg width="20" height="20" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><path d="M18 6 6 18M6 6l12 12"/></svg>',
  "cloud": '<svg width="20" height="20" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><path d="M18 10h-1.26A8 8 0 1 0 9 20h9a5 5 0 0 0 0-10z"/></svg>',
  "gear": '<svg width="20" height="20" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><circle cx="12" cy="12" r="3"/><path d="M19.4 15a1.65 1.65 0 0 0 .33 1.82l.06.06a2 2 0 1 1-2.83 2.83l-.06-.06a1.65 1.65 0 0 0-1.82-.33 1.65 1.65 0 0 0-1 1.51V21a2 2 0 0 1-4 0v-.09A1.65 1.65 0 0 0 9 19.4a1.65 1.65 0 0 0-1.82.33l-.06.06a2 2 0 1 1-2.83-2.83l.06-.06A1.65 1.65 0 0 0 4.6 15a1.65 1.65 0 0 0-1.51-1H3a2 2 0 0 1 0-4h.09A1.65 1.65 0 0 0 4.6 9a1.65 1.65 0 0 0-.33-1.82l-.06-.06a2 2 0 1 1 2.83-2.83l.06.06A1.65 1.65 0 0 0 9 4.6a1.65 1.65 0 0 0 1-1.51V3a2 2 0 0 1 4 0v.09a1.65 1.65 0 0 0 1 1.51 1.65 1.65 0 0 0 1.82-.33l.06-.06a2 2 0 1 1 2.83 2.83l-.06.06A1.65 1.65 0 0 0 19.4 9c.14.35.4.65.73.83.3.17.65.26 1 .26H21a2 2 0 0 1 0 4h-.09a1.65 1.65 0 0 0-1.51 1z"/></svg>',
  "download": '<svg width="20" height="20" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><path d="M12 4v12m0 0-4-4m4 4 4-4"/><path d="M4 20h16"/></svg>',
  "boat": '<svg width="20" height="20" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><path d="M2 18h20l-2 4H4z"/><path d="M4 18V9l8-6 8 6v9"/><path d="M12 3v15"/></svg>',
  "home": '<svg width="20" height="20" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><path d="m3 9 9-7 9 7v11a2 2 0 0 1-2 2H5a2 2 0 0 1-2-2z"/><path d="M9 22V12h6v10"/></svg>',
  "chevron-down": '<svg width="20" height="20" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2.5" stroke-linecap="round" stroke-linejoin="round"><path d="m6 9 6 6 6-6"/></svg>',
  "chart-bar": '<svg width="20" height="20" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><path d="M3 3v18h18"/><rect x="7" y="12" width="3" height="6"/><rect x="12" y="8" width="3" height="10"/><rect x="17" y="5" width="3" height="13"/></svg>',
  "user": '<svg width="20" height="20" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><circle cx="12" cy="8" r="4"/><path d="M4 21c0-4.4 3.6-8 8-8s8 3.6 8 8"/></svg>',
  "search": '<svg width="20" height="20" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><circle cx="11" cy="11" r="7"/><path d="m21 21-4.3-4.3"/></svg>',
  "book": '<svg width="20" height="20" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><path d="M4 19.5A2.5 2.5 0 0 1 6.5 17H20"/><path d="M6.5 2H20v20H6.5A2.5 2.5 0 0 1 4 19.5v-15A2.5 2.5 0 0 1 6.5 2z"/></svg>',
  "life-buoy": '<svg width="20" height="20" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><circle cx="12" cy="12" r="10"/><circle cx="12" cy="12" r="4"/><path d="m4.9 4.9 4.24 4.24m5.72 5.72 4.24 4.24m0-14.2-4.24 4.24m-5.72 5.72L4.9 19.1"/></svg>',
}

HEADER_HTML = """
<header class="site-header">
  <div class="brand">
    <div class="gov-badge"><img src="/static/logo-kkp.png" alt="Kementerian Kelautan dan Perikanan"></div>
    <div class="header-divider"></div>
    <div class="gov-badge rect"><img src="/static/logo-djprl.png" alt="DJPRL"></div>
    <div class="header-divider"></div>
    <div class="gov-text">BALAI PENATAAN<br>RUANG LAUT (BPRL)<br>MAKASSAR</div>
  </div>

  <nav class="navbar">
    <a href="/" class="nav-link active">""" + ICONS["home"] + """ Beranda</a>
    <a href="#" class="nav-link" onclick="return false;">Layanan """ + ICONS["chevron-down"] + """</a>
    <a href="#" class="nav-link" onclick="return false;">Informasi """ + ICONS["chevron-down"] + """</a>
    <a href="#" class="nav-link" onclick="return false;">""" + ICONS["book"] + """ Panduan</a>
    <a href="#" class="nav-link" onclick="return false;">""" + ICONS["life-buoy"] + """ Bantuan</a>
    <a href="/history" class="nav-link">""" + ICONS["chart-bar"] + """ Laporan</a>
  </nav>

  {% if session.user %}
  <div class="user-badge">
    <a href="/riwayat-saya" class="logout-link" style="color:var(--navy);background:#fff;">Riwayat Saya</a>
    {% if session.user.picture %}<img src="{{ session.user.picture }}" class="user-avatar" alt="">{% endif %}
    <span class="user-greet">Halo, {{ session.user.name }}</span>
    <a href="/logout" class="logout-link">Keluar</a>
  </div>
  {% else %}
  <a href="/login-pegawai" class="login-btn">""" + ICONS["user"] + """ Login</a>
  {% endif %}
</header>
"""

UPLOAD_HTML = """<!DOCTYPE html>
<html lang="id"><head><meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>e-GeRAI KKPRL &mdash; Generate &amp; Asistensi Dokumen KKPRL</title>
<meta name="description" content="Platform layanan digital terintegrasi untuk Konsultasi, Asistensi, Pendampingan, Informasi &amp; Generate Dokumen KKPRL secara cepat, tepat, efisien dan efektif.">
<meta property="og:title" content="e-GeRAI KKPRL &mdash; Generate &amp; Asistensi Dokumen KKPRL">
<meta property="og:description" content="Platform layanan digital terintegrasi untuk Konsultasi, Asistensi, Pendampingan, Informasi &amp; Generate Dokumen KKPRL secara cepat, tepat, efisien dan efektif.">
<meta property="og:type" content="website">
<style>""" + LANDING_CSS + """</style></head>
<body>
""" + HEADER_HTML + """

<section class="hero hero-image">
  <div class="hero-inner">
    <img src="/static/hero-banner.png" alt="Generate &amp; Asistensi Dokumen KKPRL" class="hero-banner-img">
  </div>
</section>

<div class="main-wrap">
  <form method="POST" action="/review" enctype="multipart/form-data" id="genForm">
    <div class="grid">
      {% if error %}<div class="error-banner">\u26A0 {{ error }}</div>{% endif %}

      <div class="upload-card">
        <div class="step-head">
          <div class="step-num">1</div>
          <div class="step-icon">""" + ICONS["doc"] + """</div>
        </div>
        <div class="step-title">Draft Proposal PKKPRL (PDF/Word)</div>
        <div class="step-desc">Unggah file PDF atau Word proposal yang akan digabungkan. Belum punya file-nya?
        <a href="/proposal-manual" style="color:var(--blue); font-weight:700;">Isi Formulir di sini</a>.</div>
        <div class="dropzone" id="dz1">
          """ + ICONS["cloud"] + """
          <div class="dz-title">Drag &amp; Drop PDF/Word di sini</div>
          <div class="dz-sub">atau klik untuk memilih file</div>
          <div class="dz-max">Maksimum 10 MB &middot; PDF atau .docx</div>
          <input type="file" name="proposal" id="proposal" accept="application/pdf,.docx" required>
        </div>
        <div class="file-chip" id="chip1">
          <div class="fc-left">""" + ICONS["check-circle"] + """<span class="fc-name" id="name1"></span></div>
          <span class="fc-size" id="size1"></span>
          <button type="button" class="fc-remove" data-target="1">""" + ICONS["x"] + """</button>
        </div>
      </div>

      <div class="upload-card">
        <div class="step-head">
          <div class="step-num">2</div>
          <div class="step-icon">""" + ICONS["wave"] + """</div>
        </div>
        <div class="step-title">Laporan Kondisi Eksisting / Hidro-Oseanografi (PDF/Word)</div>
        <div class="step-desc">Unggah file PDF atau Word laporan hidro-oseanografi.</div>
        <div class="dropzone" id="dz2">
          """ + ICONS["cloud"] + """
          <div class="dz-title">Drag &amp; Drop PDF/Word di sini</div>
          <div class="dz-sub">atau klik untuk memilih file</div>
          <div class="dz-max">Maksimum 10 MB &middot; PDF atau .docx</div>
          <input type="file" name="laporan" id="laporan" accept="application/pdf,.docx" required>
        </div>
        <div class="file-chip" id="chip2">
          <div class="fc-left">""" + ICONS["check-circle"] + """<span class="fc-name" id="name2"></span></div>
          <span class="fc-size" id="size2"></span>
          <button type="button" class="fc-remove" data-target="2">""" + ICONS["x"] + """</button>
        </div>
      </div>

      <div class="flow-card">
        <h3>Alur Proses</h3>
        <div class="flow-step">
          <div class="flow-dot">""" + ICONS["cloud"] + """</div>
          <div class="flow-body"><div class="ft">Upload Proposal</div><div class="fd">Unggah file PDF Proposal PKKPRL</div></div>
        </div>
        <div class="flow-step">
          <div class="flow-dot">""" + ICONS["wave"] + """</div>
          <div class="flow-body"><div class="ft">Upload Laporan</div><div class="fd">Unggah file PDF Laporan Hidro-Oseanografi</div></div>
        </div>
        <div class="flow-step">
          <div class="flow-dot">""" + ICONS["gear"] + """</div>
          <div class="flow-body"><div class="ft">Generate Dokumen</div><div class="fd">Sistem menggabungkan dokumen secara otomatis</div></div>
        </div>
        <div class="flow-step">
          <div class="flow-dot">""" + ICONS["download"] + """</div>
          <div class="flow-body"><div class="ft">Download Dokumen</div><div class="fd">Dokumen Word siap diunduh dan diedit</div></div>
        </div>
      </div>

      <div class="gen-btn">
        <button type="submit">""" + ICONS["bolt"] + """ Generate &amp; Preview Dokumen Word</button>
        <div class="gen-note">Sistem akan memproses dan membuat dokumen Word final secara otomatis</div>
        <div class="spinner" id="spinner">\u23F3 Memproses dokumen, mohon tunggu...</div>
      </div>
    </div>
  </form>
</div>

<script>
function fmtSize(bytes) {
  if (bytes >= 1024*1024) return (bytes/(1024*1024)).toFixed(2) + " MB";
  return Math.ceil(bytes/1024) + " KB";
}
function setupDropzone(n) {
  var dz = document.getElementById('dz' + n);
  var input = dz.querySelector('input[type=file]');
  var chip = document.getElementById('chip' + n);
  var nameEl = document.getElementById('name' + n);
  var sizeEl = document.getElementById('size' + n);

  function showFile(file) {
    if (!file) { chip.classList.remove('show'); dz.style.display = 'block'; return; }
    nameEl.textContent = file.name;
    nameEl.title = file.name;
    sizeEl.textContent = fmtSize(file.size);
    chip.classList.add('show');
    dz.style.display = 'none';
  }

  dz.addEventListener('click', function(e) {
    if (e.target.closest('.file-chip')) return;
    input.click();
  });
  input.addEventListener('change', function() {
    if (input.files && input.files[0]) showFile(input.files[0]);
  });
  dz.addEventListener('dragover', function(e) { e.preventDefault(); dz.classList.add('dragover'); });
  dz.addEventListener('dragleave', function() { dz.classList.remove('dragover'); });
  dz.addEventListener('drop', function(e) {
    e.preventDefault();
    dz.classList.remove('dragover');
    if (e.dataTransfer.files && e.dataTransfer.files[0]) {
      input.files = e.dataTransfer.files;
      showFile(input.files[0]);
    }
  });
  document.querySelector('.fc-remove[data-target="' + n + '"]').addEventListener('click', function() {
    input.value = '';
    dz.style.display = 'block';
    chip.classList.remove('show');
  });
}
setupDropzone(1);
setupDropzone(2);

document.getElementById('genForm').addEventListener('submit', function() {
  document.getElementById('spinner').style.display = 'block';
});
</script>
</body></html>"""


REVIEW_CSS = """
.review-hero { background:linear-gradient(135deg,#eaf2fb 0%,#cfe1f6 55%,#a9cdec 100%);
  padding:26px 32px 34px; }
.review-hero h1 { font-size:24px; font-weight:800; color:var(--navy); margin:0 0 6px; }
.review-hero p { font-size:13.5px; color:#33495e; margin:0; max-width:640px; line-height:1.5; }
.review-wrap { max-width:1600px; margin:-16px auto 40px; padding:0 40px; position:relative; z-index:2; }
.review-grid { display:grid; grid-template-columns:1.15fr .85fr; gap:22px; align-items:start; }
@media (max-width: 980px) { .review-grid { grid-template-columns:1fr; } }

.sticky-bar { position:sticky; top:0; z-index:6; background:var(--bg); padding:14px 0 10px; }
.sticky-bar-row { display:flex; gap:10px; }
.sticky-bar button { width:100%; background:linear-gradient(90deg,var(--blue2),var(--navy)); color:#fff;
  border:none; padding:15px; border-radius:12px; font-size:14.5px; font-weight:800; cursor:pointer;
  display:flex; align-items:center; justify-content:center; gap:8px; box-shadow:0 6px 18px rgba(30,99,199,.28); }
.sticky-bar button svg { width:19px; height:19px; flex:none; }
.sticky-bar button:hover { filter:brightness(1.05); }
.sticky-bar button.btn-draft { background:#fff; color:var(--navy); border:1.5px solid #cfe0f5;
  box-shadow:none; flex:none; width:auto; white-space:nowrap; padding:15px 20px; }
.sticky-bar button.btn-draft:hover { background:#f3f8ff; filter:none; }
.back-link { display:inline-flex; align-items:center; gap:6px; font-size:12.5px; font-weight:700;
  color:var(--blue); margin-top:14px; }

.review-card { background:#fff; border-radius:16px; padding:22px 24px; box-shadow:0 6px 24px rgba(18,58,99,.08);
  margin-bottom:18px; }
.review-card h3 { font-size:14.5px; font-weight:800; color:var(--navy); margin:0 0 14px;
  display:flex; align-items:center; gap:8px; }
.review-card h3 svg { width:18px; height:18px; color:var(--blue); }

.acc-item { border:1px solid var(--line); border-radius:12px; margin-bottom:10px; overflow:hidden; }
.acc-item:last-child { margin-bottom:0; }
.acc-item summary { cursor:pointer; list-style:none; padding:13px 16px; font-size:13.5px; font-weight:700;
  color:var(--navy); background:#f7fafd; display:flex; align-items:center; justify-content:space-between; }
.acc-item summary::-webkit-details-marker { display:none; }
.acc-item summary::after { content:"\\25BE"; color:var(--blue); font-size:12px; transition:.15s; }
.acc-item[open] summary::after { transform:rotate(180deg); }
.acc-body { padding:14px 16px 4px; }
.field-row { margin-bottom:12px; }
.field-row label { display:block; font-size:12px; font-weight:700; color:var(--muted); margin-bottom:5px; }
.field-row input[type=text] { width:100%; padding:11px 13px; border:1px solid #d3dde7; border-radius:8px;
  font-size:18px; color:var(--ink); background:#fff; }
.field-row input[type=text]:focus { outline:none; border-color:var(--blue); box-shadow:0 0 0 3px rgba(47,127,224,.15); }

.preview-panel { position:sticky; top:80px; }
.preview-box { border:1px solid var(--line); border-radius:12px; background:#fff; padding:24px 28px;
  max-height:70vh; overflow-y:auto; font-size:13.5px; line-height:1.55; }
.preview-box table { border-collapse:collapse; width:100%; margin:10px 0; }
.preview-box td { border:1px solid #c7d1db; padding:5px 8px; vertical-align:top; }
.preview-box img { max-width:100%; height:auto; margin:8px 0; }
.preview-box p { margin:6px 0; }

.checkbox-row { display:flex; align-items:center; gap:10px; margin-bottom:12px; }
.checkbox-row input[type=checkbox] { width:18px; height:18px; accent-color:var(--blue); flex:none; }
.checkbox-row label { font-size:13px; font-weight:600; color:var(--ink); }
.manual-hero { background:linear-gradient(135deg,#eaf2fb 0%,#cfe1f6 55%,#a9cdec 100%); padding:26px 32px 34px; }
.manual-hero h1 { font-size:24px; font-weight:800; color:var(--navy); margin:0 0 6px; }
.manual-hero p { font-size:13.5px; color:#33495e; margin:0; max-width:640px; line-height:1.5; }
.manual-upload-card { background:#fff; border-radius:16px; padding:22px 24px; box-shadow:0 6px 24px rgba(18,58,99,.08); margin-bottom:18px; }

.file-field-row { margin-bottom:16px; }
.file-field-row label { display:block; font-size:12.5px; font-weight:700; color:var(--navy); margin-bottom:3px; }
.file-field-row .ff-hint { font-size:11px; color:var(--muted); margin-bottom:6px; }
.file-field-row input[type=file] { width:100%; font-size:12.5px; padding:8px; border:1px solid #d3dde7;
  border-radius:8px; background:#f7fafd; }
.field-row textarea { width:100%; padding:11px 13px; border:1px solid #d3dde7; border-radius:8px;
  font-size:18px; font-family:monospace; color:var(--ink); background:#fff; resize:vertical; min-height:96px; }
.field-row textarea:focus { outline:none; border-color:var(--blue); box-shadow:0 0 0 3px rgba(47,127,224,.15); }
.field-example { font-size:17px; color:var(--muted); margin-top:6px; display:flex; align-items:center; gap:10px; flex-wrap:wrap; line-height:1.4; }
.field-example .ex-text { font-style:italic; }
.field-example .ex-fill { font-size:14px; font-weight:700; color:var(--blue); background:#eaf1fc;
  border:1px solid #cfe0f5; border-radius:8px; padding:4px 12px; cursor:pointer; white-space:nowrap; }
.field-example .ex-fill:hover { background:#dcebfa; }

.img-paste-zone { border:2px dashed #b9cbe0; border-radius:10px; background:#f7fafd; padding:14px;
  cursor:pointer; transition:.15s; outline:none; }
.img-paste-zone:hover, .img-paste-zone:focus, .img-paste-zone.dragover { border-color:var(--blue); background:#eef5fd; }
.img-paste-placeholder { font-size:11.5px; color:var(--muted); text-align:center; }
.img-paste-placeholder svg { width:18px; height:18px; display:block; margin:0 auto 4px; color:var(--blue); }
.img-preview-list { display:flex; flex-wrap:wrap; gap:8px; margin-top:8px; }
.img-preview-list:empty { margin-top:0; }
.img-thumb { position:relative; width:64px; height:64px; border-radius:8px; overflow:hidden;
  border:1px solid var(--line); flex:none; }
.img-thumb img { width:100%; height:100%; object-fit:cover; display:block; }
.img-thumb .img-remove { position:absolute; top:2px; right:2px; width:18px; height:18px; border-radius:50%;
  background:rgba(0,0,0,.55); color:#fff; border:none; font-size:11px; line-height:1; cursor:pointer;
  display:flex; align-items:center; justify-content:center; }
.img-thumb-saved { border-color:#8fd19e; }
.img-thumb-badge { position:absolute; bottom:0; left:0; right:0; background:rgba(30,120,60,.85); color:#fff;
  font-size:8.5px; font-weight:700; text-align:center; padding:2px 0; }
.img-preview-saved:empty { display:none; margin:0; }

.img-input-row { display:flex; gap:8px; align-items:stretch; justify-content:center; max-width:30%; margin:0 auto; }
.img-upload-btn { flex:none; width:110px; display:flex; align-items:center; justify-content:center; gap:6px; background:#fff;
  border:1.5px solid var(--line); border-radius:10px; padding:0 10px; height:40px;
  font-size:12.5px; font-weight:700; color:var(--navy); cursor:pointer; white-space:nowrap; }
.img-upload-btn:hover { background:#f3f8ff; border-color:#cfe0f5; }
.img-upload-btn svg { width:16px; height:16px; color:var(--blue); flex:none; }
.img-paste-target { flex:none; width:90px; height:40px; border:2px dashed #b9cbe0; border-radius:10px;
  background:#f7fafd; padding:0 10px; cursor:text; outline:none; transition:.15s;
  display:flex; align-items:center; justify-content:center; gap:5px; }
.img-paste-target:hover, .img-paste-target:focus, .img-paste-target.dragover { border-color:var(--blue); background:#eef5fd; }
.img-paste-target svg { width:15px; height:15px; color:var(--blue); flex:none; }
.img-paste-target .ipt-text { font-size:11px; color:var(--muted); line-height:1.25; white-space:nowrap; }
.img-paste-target .ipt-text b { color:var(--navy); }

@media (max-width: 700px) { .img-input-row { max-width:80%; } }

.select-field { width:100%; padding:11px 13px; border:1px solid #d3dde7; border-radius:8px;
  font-size:18px; color:var(--ink); background:#fff; }
.select-field:focus { outline:none; border-color:var(--blue); box-shadow:0 0 0 3px rgba(47,127,224,.15); }
.select-other-input { margin-top:8px; display:none; }
.select-other-input.show { display:block; }
.eco-children { padding-left:16px; border-left:3px solid var(--line); margin:-4px 0 14px; transition:.2s; }
.eco-children.locked { opacity:.5; }
.eco-children.locked input, .eco-children.locked select { background:#f3f5f7; cursor:not-allowed; }
.eco-lock-note { font-size:11px; color:var(--muted); font-style:italic; margin:2px 0 10px; }

.dukung-item { margin-bottom:6px; }
.dukung-detail { padding:10px 0 6px 26px; border-left:2px solid #e0e8f0; margin:2px 0 8px 8px; }
.dukung-detail .field-row { margin-bottom:8px; }
.dukung-custom-row { display:flex; gap:8px; align-items:flex-start; margin-bottom:8px; }
.dukung-custom-row input[type=text] { flex:1; }
.dukung-custom-row .dukung-remove { flex:none; background:#fff0f0; border:1px solid #f3c6c6; color:#c0392b;
  border-radius:8px; padding:9px 12px; cursor:pointer; font-size:12px; font-weight:700; }

.money-field { display:flex; align-items:stretch; border:1px solid #d3dde7; border-radius:8px; overflow:hidden; }
.money-field .money-prefix { background:#f0f4f9; color:var(--navy); font-weight:700; font-size:18px;
  padding:11px 12px; border-right:1px solid #d3dde7; display:flex; align-items:center; }
.money-field .money-input { border:none; border-radius:0; flex:1; padding:11px 13px; font-size:18px; }
.money-field .money-input:focus { outline:none; }
.money-field:focus-within { border-color:var(--blue); box-shadow:0 0 0 3px rgba(47,127,224,.15); }
.date-picker-input { width:100%; padding:11px 13px; border:1px solid #d3dde7; border-radius:8px; font-size:18px; color:var(--ink); background:#fff; }
.date-picker-input:focus { outline:none; border-color:var(--blue); box-shadow:0 0 0 3px rgba(47,127,224,.15); }
.field-note { font-size:11px; color:var(--muted); margin-top:4px; }

.history-table { width:100%; border-collapse:collapse; font-size:13px; }
.history-table th, .history-table td { padding:10px 12px; border-bottom:1px solid var(--line); text-align:left; }
.history-table th { color:var(--muted); font-weight:700; font-size:11px; text-transform:uppercase; letter-spacing:.03em; }
.history-table tr:hover td { background:#f7fafd; }
.history-empty { text-align:center; padding:40px 20px; color:var(--muted); font-size:13.5px; }
.history-login-gate { text-align:center; padding:60px 20px; }
.history-login-gate p { color:var(--muted); font-size:14px; margin:10px 0 20px; }
"""


FIELD_HELP = {
    ("prop", "Nama Pemohon"): "Mohon isi nama perwakilan perusahaan/instansi (Kepala) sebagai PIC yang bertanggung jawab dalam permohonan KKPRL.",
    ("prop", "Jabatan Pemohon"): "Jabatan dari perwakilan perusahaan/instansi (Kepala) yang bertanggung jawab dalam permohonan KKPRL.",
    ("prop", "Nama Perusahaan/Instansi"): "Ditulis dengan benar, tanpa disingkat.",
    ("prop", "NIB"): "Jika tidak ada, isi dengan tanda -",
    ("prop", "NPWP"): "NPWP milik perusahaan/Instansi, atau milik pemohon jika perseorangan.",
    ("prop", "Nomor Telepon Selular"): "Nomor yang dapat dihubungi untuk informasi lebih lanjut.",
    ("prop", "Surat Elektronik"): "Email aktif — akan dipakai untuk mengirim Draft Proposal yang telah selesai.",
    ("prop", "Luas Kebutuhan Ruang"): "Isi berupa angka (dalam hektar).",
    ("prop_loc", "0"): "Nama Desa/Kelurahan dari lokasi yang dimohonkan.",
    ("prop_loc", "1"): "Nama Kecamatan dari lokasi yang dimohonkan.",
    ("prop_loc", "2"): "Nama Kabupaten/Kota dari lokasi yang dimohonkan.",
    ("prop", "mangrove_persen"): "Dilampirkan dalam bentuk persentase (%).",
    ("prop", "desa_luas_ha"): "Luas penduduk desa atau desa terdekat dari lokasi yang dimohonkan.",
    ("prop", "desa_penduduk"): "Jumlah penduduk desa atau desa terdekat dari lokasi yang dimohonkan.",
    ("prop", "tenaga_kerja_asing"): "Jika tidak ada, isi dengan 0.",
}

EXAMPLE_HINTS = {
    ("prop", "Nama Pemohon"): "Andi Wijaya, S.T., M.M.",
    ("prop", "Jabatan Pemohon"): "Direktur Utama",
    ("prop", "Nama Perusahaan/Instansi"): "PT. Bahari Sejahtera Makassar",
    ("prop", "NIB"): "-",
    ("prop", "NPWP"): "01.234.567.8-901.000",
    ("prop", "Nomor Telepon Selular"): "081234567890",
    ("prop", "Surat Elektronik"): "info@baharisejahteramks.co.id",
    ("prop", "Jenis Kegiatan"): "Keramba Jaring Apung",
    ("prop", "Nama Perairan"): "Selat Makassar",
    ("prop", "Luas Kebutuhan Ruang"): "0.59 Ha",
    ("prop", "KBLI"): "50121 - Angkutan Laut Wisata Dalam Negeri",
    ("prop", "Tanggal Penyusunan"): "02 Agustus 2026",
    ("prop_loc", "0"): "Desa Bontolebang",
    ("prop_loc", "1"): "Kecamatan Ujung Tanah",
    ("prop_loc", "2"): "Kabupaten Pangkajene dan Kepulauan",
    ("prop_loc", "3"): "Sulawesi Selatan",
    ("prop", "investasi"): "200.000.0000",
    ("prop", "tenaga_kerja"): "15 Orang",
    ("prop", "tenaga_kerja_asing"): "2 Orang",
    ("prop", "mangrove_spesies"): "Rhizophora mucronata",
    ("prop", "mangrove_persen"): "65",
    ("prop", "mangrove_kondisi"): "Sangat Padat",
    ("prop", "desa_luas_ha"): "250",
    ("prop", "desa_penduduk"): "3400",
}

# Field yang dirender sebagai dropdown (select) alih-alih input teks bebas.
# allow_other=True -> menampilkan input teks tambahan saat "Other" dipilih.
KBLI_OPTIONS = [
    "93295 - Wisata Pantai",
    "93296 - Wisata Agro",
    "93297 - Wisata Tirta",
    "93299 - Aktivitas Hiburan dan Rekreasi Lainnya YTDL",
    "55101 - Aktivitas Hotel Bintang Lima",
    "55102 - Aktivitas Hotel Bintang Empat",
    "55103 - Aktivitas Hotel Bintang Tiga",
    "55104 - Aktivitas Hotel Bintang Dua",
    "55105 - Aktivitas Hotel Bintang Satu",
    "50113 - Angkutan Laut Dalam Negeri untuk Wisata",
    "03110 - Penangkapan Ikan dan Biota Air Lainnya di Laut",
    "03120 - Penangkapan Ikan dan Biota Air Lainnya di Perairan Air Tawar",
    "03211 - Pembudidayaan Ikan Bersirip (Selain Ikan Hias) dan Biota Air Laut Lainnya yang Tidak Dilindungi",
    "03212 - Pembudidayaan Ikan Hias Air Laut yang Tidak Dilindungi",
    "03213 - Pembudidayaan Tumbuhan Air Laut yang Tidak Dilindungi",
    "03214 - Pengembangbiakan Ikan dan Biota Air Laut yang Dilindungi",
    "03231 - Pembudidayaan Ikan Bersirip (Selain Ikan Hias) dan Biota Air Payau Lainnya yang Tidak Dilindungi",
    "03232 - Pembudidayaan Ikan Hias Air Payau yang Tidak Dilindungi",
    "03233 - Pembudidayaan Tumbuhan Air Payau yang Tidak Dilindungi",
    "03234 - Pengembangbiakan Biota Air Payau yang Dilindungi",
]

SELECT_FIELDS = {
    ("prop", "Nama Perairan"): {
        "options": ["Laut Banda", "Laut Sulawesi", "Laut Bali", "Laut Sawu", "Laut Flores",
                     "Selat Makassar", "Teluk Bone", "Teluk Tomini"],
        "allow_other": True,
    },
    ("prop_loc", "3"): {
        "options": ["Sulawesi Selatan", "Sulawesi Barat", "Sulawesi Utara", "Sulawesi Tengah",
                     "Sulawesi Tenggara", "Gorontalo", "Bali", "Nusa Tenggara Barat", "Nusa Tenggara Timur"],
        "allow_other": False,
    },
    ("prop", "Jenis Kegiatan"): {
        "options": ["Pemanfaatan Air Laut untuk Budi Daya", "Keramba Jaring Apung", "Dermaga"],
        "allow_other": True,
    },
    ("prop", "KBLI"): {
        "options": KBLI_OPTIONS,
        "allow_other": True,
    },
    ("prop", "mangrove_kondisi"): {
        "options": ["Sangat Padat", "Sedang", "Jarang"],
        "allow_other": False,
    },
    ("prop", "instalasi_posisi"): {
        "options": ["Permukaan Laut", "Kolom Laut", "Dasar Laut"],
        "allow_other": False,
    },
    ("prop", "mangrove_ada"): {
        "options": ["Terdapat ekosistem mangrove", "Tidak terdapat ekosistem mangrove"],
        "allow_other": False,
    },
    ("prop", "lamun_ada_manual"): {
        "options": ["Terdapat ekosistem lamun", "Tidak terdapat ekosistem lamun"],
        "allow_other": False,
    },
    ("prop", "lamun_kondisi"): {
        "options": ["Baik (Kaya/Sehat)", "Rusak (Kurang Kaya/Kurang Sehat)", "Rusak (Miskin)"],
        "allow_other": False,
    },
    ("prop", "karang_ada"): {
        "options": ["Terdapat ekosistem terumbu karang", "Tidak terdapat ekosistem terumbu karang"],
        "allow_other": False,
    },
    ("prop", "karang_kondisi"): {
        "options": ["Baik Sekali", "Baik", "Sedang", "Buruk"],
        "allow_other": False,
    },
}

# Kriteria baku kondisi ekosistem (sesuai dokumen resmi KRITERIA_KONDISI) --
# dipakai untuk menentukan otomatis label kondisi berdasarkan persentase
# tutupan yang diinput, supaya user tidak perlu pilih manual satu-satu.
# Format: list of (batas_bawah_persen, batas_atas_persen, label_kondisi).
KRITERIA_MANGROVE = [
    (75, 100.001, "Sangat Padat"),
    (50, 75, "Sedang"),
    (0, 50, "Jarang"),
]
KRITERIA_LAMUN = [
    (60, 100.001, "Baik (Kaya/Sehat)"),
    (30, 60, "Rusak (Kurang Kaya/Kurang Sehat)"),
    (0, 30, "Rusak (Miskin)"),
]
KRITERIA_KARANG = [
    (75, 100.001, "Baik Sekali"),
    (50, 75, "Baik"),
    (25, 50, "Sedang"),
    (0, 25, "Buruk"),
]


def klasifikasi_kondisi(persen_str, kriteria):
    """Cari label kondisi yang sesuai dari daftar kriteria berdasarkan nilai
    persentase (string, boleh pakai koma atau titik). Kembalikan '' kalau
    nilainya tidak valid/tidak bisa diparse."""
    try:
        v = float(str(persen_str).replace(",", ".").replace("%", "").strip())
    except (TypeError, ValueError):
        return ""
    for lo, hi, label in kriteria:
        if lo <= v < hi:
            return label
    return ""

# Kode wilayah BPS untuk tiap provinsi (dipakai untuk fetch data Kabupaten/Kecamatan/Desa
# secara berjenjang dari API publik https://emsifa.github.io/api-wilayah-indonesia).
PROVINCE_BPS_IDS = {
    "Sulawesi Selatan": "73",
    "Sulawesi Barat": "76",
    "Sulawesi Utara": "71",
    "Sulawesi Tengah": "72",
    "Sulawesi Tenggara": "74",
    "Gorontalo": "75",
    "Bali": "51",
    "Nusa Tenggara Barat": "52",
    "Nusa Tenggara Timur": "53",
}


def render_select_html(fname, options, allow_other=False):
    opts_html = ['<option value="">-- Pilih --</option>']
    for opt in options:
        opts_html.append(f'<option value="{opt}">{opt}</option>')
    other_html = ""
    if allow_other:
        opts_html.append('<option value="__other__">Lainnya...</option>')
        other_html = f'<input type="text" class="select-other-input" id="{fname}_other" placeholder="Isi lainnya" data-for="{fname}">'
    return (
        f'<select class="select-field" name="{fname}" id="{fname}" data-allow-other="{str(allow_other).lower()}">'
        f'{"".join(opts_html)}</select>{other_html}'
    )


def dukung_item_html(idx, checkbox_name, label):
    """Satu baris item Dokumen Data Dukung: checkbox + (kalau dicentang)
    muncul otomatis kolom Link Google Drive dan Upload File."""
    cb_id = f"dd{idx}"
    return (
        f'<div class="dukung-item">'
        f'<div class="checkbox-row"><input type="checkbox" name="{checkbox_name}" id="{cb_id}" class="dukung-cb" data-target="dukung_detail_{cb_id}">'
        f'<label for="{cb_id}">{label}</label></div>'
        f'<div class="dukung-detail" id="dukung_detail_{cb_id}" style="display:none;">'
        f'<div class="field-row"><label>Link Google Drive Dokumen</label>'
        f'<input type="text" name="{checkbox_name}_drive" placeholder="Tempel link Google Drive di sini"></div>'
        f'<div class="field-row"><label>Atau Upload File</label>'
        f'<input type="file" name="{checkbox_name}_file" accept=".pdf,.doc,.docx,.png,.jpg,.jpeg"></div>'
        f'</div></div>'
    )


def join_dan(items):
    """Gabungkan list string jadi satu kalimat natural pakai kata 'dan' untuk
    item terakhir (mis. ['Permukaan Laut','Dasar Laut'] -> 'Permukaan Laut
    dan Dasar Laut'). Dipakai untuk field checklist yang boleh pilih lebih
    dari satu (mis. Instalasi Bangunan Laut Berada Pada)."""
    items = [i.strip() for i in items if i and i.strip()]
    if not items:
        return ""
    if len(items) == 1:
        return items[0]
    return ", ".join(items[:-1]) + " dan " + items[-1]


# Pemetaan nama field upload (form) -> tag gambar internal, dipakai baik
# untuk membaca file upload (build_prop_data_from_manual_form) maupun untuk
# menampilkan kembali thumbnail gambar yang sudah tersimpan (render_manual_form_page).
IMAGE_FIELD_TAGS = [
    ("img_siteplan", "siteplan"),
    ("img_peta_lokasi", "peta_lokasi"),
    ("img_foto_mangrove", "foto_mangrove"),
    ("img_foto_karang_insitu", "foto_karang_insitu"),
    ("img_dok_kegiatan", "dok_kegiatan_eksisting"),
    ("img_dok_pemanfaatan_sekitar", "dok_pemanfaatan_sekitar"),
    ("img_foto_lamun", "foto_lamun"),
    ("img_aksesibilitas", "gambar_aksesibilitas"),
    ("img_sertifikat_lahan", "sertifikat_lahan"),
    ("img_dok_sosialisasi", "dok_sosialisasi"),
    ("img_dok_pendukung_lainnya", "dok_pendukung_lainnya"),
]


def img_field_html(field_name, label, desc="", multiple=False, note="", saved_previews=None):
    desc_html = f'<div class="ff-hint" style="margin-bottom:6px;">{desc}</div>' if desc else ""
    note_html = f'<span style="font-weight:400;color:var(--muted);"> {note}</span>' if note else ""
    saved_html = ""
    if saved_previews:
        thumbs = "".join(
            f'<div class="img-thumb img-thumb-saved" title="Tersimpan dari isian sebelumnya &mdash; '
            f'tidak perlu diunggah ulang, kecuali ingin menggantinya">'
            f'<img src="{p["url"]}" alt="{label}">'
            f'<span class="img-thumb-badge">\u2713 Tersimpan</span>'
            f'</div>'
            for p in saved_previews
        )
        saved_html = f'<div class="img-preview-list img-preview-saved">{thumbs}</div>'
    return (
        f'<div class="file-field-row">'
        f'<label>{label}{note_html}</label>{desc_html}'
        f'<div class="img-input-row">'
        f'<button type="button" class="img-upload-btn" data-field="{field_name}">{ICONS["doc"]} Upload File</button>'
        f'<div class="img-paste-target" tabindex="0" data-field="{field_name}">{ICONS["upload"]}'
        f'<span class="ipt-text"><b>Ctrl+V</b><br>paste</span></div>'
        f'</div>'
        f'<input type="file" name="{field_name}" id="{field_name}" accept="image/png,image/jpeg,.png,.jpg,.jpeg" multiple style="display:none">'
        f'{saved_html}'
        f'<div class="img-preview-list" id="{field_name}_preview"></div>'
        f'</div>'
    )


def render_login_pegawai_page(error=None):
    # Render homepage SUNGGUHAN (sudah lewat Jinja) sebagai backdrop, lalu
    # sisipkan overlay + kartu login di atasnya tepat sebelum tag </body>.
    # CSS ditulis mandiri (inline) di sini karena halaman backdrop (UPLOAD_HTML)
    # cuma memuat LANDING_CSS, bukan REVIEW_CSS.
    home_rendered = render_template_string(UPLOAD_HTML, error=None)

    error_html = (
        f'<div class="lgo-error">\u26A0 {error}</div>' if error else ""
    )
    overlay = """
<style>
  .login-overlay { position:fixed; inset:0; z-index:9999; background:rgba(10,25,45,.62);
    backdrop-filter:blur(4px); display:flex; align-items:center; justify-content:center; padding:20px; }
  .login-modal-card { background:#fff; border-radius:16px; padding:28px 26px; box-shadow:0 20px 50px rgba(0,0,0,.35);
    max-width:400px; width:100%; }
  .login-modal-card h3 { display:flex; align-items:center; justify-content:center; gap:8px;
    font-size:16px; font-weight:800; color:#123A63; margin:0 0 6px; }
  .login-modal-card h3 svg { width:20px; height:20px; }
  .login-modal-sub { text-align:center; font-size:12.5px; color:#5b6b7c; margin:0 0 18px; }
  .lgo-error { background:#fff3cd; border:1px solid #ffe08a; color:#7a5b00; padding:10px 14px;
    border-radius:8px; font-size:12.5px; margin-bottom:14px; }
  .login-modal-card .field-row { margin-bottom:14px; }
  .login-modal-card label { display:block; font-size:12px; font-weight:700; color:#5b6b7c; margin-bottom:5px; }
  .login-modal-card input[type=text], .login-modal-card input[type=password] {
    width:100%; padding:10px 12px; border:1px solid #d3dde7; border-radius:8px; font-size:14px;
    color:#1c2b3a; background:#fff; box-sizing:border-box; }
  .login-modal-card input:focus { outline:none; border-color:#1E63C7; box-shadow:0 0 0 3px rgba(30,99,199,.15); }
  .login-modal-btn { width:100%; background:linear-gradient(90deg,#2F7FE0,#123A63); color:#fff;
    border:none; padding:13px; border-radius:10px; font-size:14.5px; font-weight:800; cursor:pointer;
    display:flex; align-items:center; justify-content:center; gap:8px; margin-top:4px; box-sizing:border-box; }
  .login-modal-btn:hover { filter:brightness(1.08); }
  .login-modal-btn svg { width:18px; height:18px; flex:none; }
  body { overflow:hidden !important; }
</style>
<div class="login-overlay">
  <div class="login-modal-card">
    """ + error_html + """
    <h3>""" + ICONS["user"] + """ Login Pegawai BPRL Makassar</h3>
    <p class="login-modal-sub">Masuk terlebih dahulu untuk menggunakan aplikasi e-GeRAI KKPRL.</p>
    <form method="POST" action="/login-pegawai">
      <div class="field-row">
        <label>Kode Nama</label>
        <input type="text" name="kode_nama" placeholder="Contoh: 09" required autofocus inputmode="numeric">
      </div>
      <div class="field-row">
        <label>Kode Petugas (Password)</label>
        <input type="password" name="kode_petugas" placeholder="Contoh: 123" required inputmode="numeric">
      </div>
      <button type="submit" class="login-modal-btn">""" + ICONS["user"] + """ Masuk</button>
    </form>
  </div>
</div>
"""
    return home_rendered.replace("</body>", overlay + "</body>")


def render_riwayat_saya_page():
    user = session.get("user")
    if not user:
        return render_login_pegawai_page()

    kode = user.get("kode", "")
    drafts = read_user_drafts(kode) if kode else []

    if drafts:
        rows = "".join(
            f'<tr><td>{d.get("waktu","-")}</td>'
            f'<td>{d.get("nama_pemohon") or "-"}</td>'
            f'<td>{d.get("nama_perusahaan") or "-"}</td>'
            f'<td><a href="/riwayat-saya/lanjutkan/{d.get("job_id")}" class="ex-fill" style="text-decoration:none;display:inline-block;">Lanjutkan Isi</a></td></tr>'
            for d in drafts
        )
        table_html = (
            '<table class="history-table"><thead><tr>'
            "<th>Waktu</th><th>Nama Pemohon</th><th>Nama Perusahaan/Instansi</th><th>Aksi</th>"
            f"</tr></thead><tbody>{rows}</tbody></table>"
        )
    else:
        table_html = '<div class="history-empty">Belum ada riwayat isian form yang tersimpan. Riwayat akan muncul di sini setiap kali Anda mengisi form manual, dan otomatis terhapus sendiri setelah ' + str(DRAFT_MAX_AGE_DAYS) + ' hari.</div>'

    body = """
<div class="review-wrap">
  <div class="review-card">
    <h3>""" + ICONS["chart-bar"] + f""" Riwayat Isian Saya <span style="font-weight:400;color:var(--muted);font-size:12px;">({len(drafts)} tersimpan &middot; otomatis terhapus setelah {DRAFT_MAX_AGE_DAYS} hari)</span></h3>
    {table_html}
  </div>
</div>"""

    return """<!DOCTYPE html>
<html lang="id"><head><meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Riwayat Isian Saya &mdash; e-GeRAI KKPRL</title>
<style>""" + LANDING_CSS + REVIEW_CSS + """</style></head>
<body>
""" + HEADER_HTML + """

<section class="review-hero">
  <h1>""" + ICONS["chart-bar"] + """ Riwayat Isian Saya</h1>
  <p>Daftar isian form manual yang pernah Anda mulai. Klik "Lanjutkan Isi" untuk mengisi ulang form secara otomatis dari data yang tersimpan (lampiran gambar/laporan perlu diunggah ulang).</p>
</section>
""" + body + """
</body></html>"""


def render_history_page():
    logged_in = bool(session.get("user"))

    if not logged_in:
        body = """
<div class="review-wrap">
  <div class="review-card history-login-gate">
    <h3 style="justify-content:center;">""" + ICONS["chart-bar"] + """ Riwayat Penggunaan</h3>
    <p>Riwayat hanya bisa dilihat oleh staf yang sudah login menggunakan akun Google BPRL.</p>
    <a href="/login-pegawai" class="login-btn" style="display:inline-flex;">""" + ICONS["user"] + """ Login untuk Melihat Riwayat</a>
  </div>
</div>"""
    else:
        entries = read_history()
        if entries:
            rows = "".join(
                f"<tr><td>{e.get('waktu','-')}</td><td>{e.get('nama_pemohon','-')}</td>"
                f"<td>{e.get('nama_perusahaan','-')}</td><td>{e.get('diproses_oleh','-')}</td></tr>"
                for e in entries
            )
            table_html = (
                '<table class="history-table"><thead><tr>'
                "<th>Waktu</th><th>Nama Pemohon</th><th>Nama Perusahaan/Instansi</th><th>Diproses Oleh</th>"
                f"</tr></thead><tbody>{rows}</tbody></table>"
            )
        else:
            table_html = '<div class="history-empty">Belum ada riwayat dokumen yang di-generate.</div>'

        body = """
<div class="review-wrap">
  <div class="review-card">
    <h3>""" + ICONS["chart-bar"] + f""" Riwayat Penggunaan <span style="font-weight:400;color:var(--muted);font-size:12px;">({len(entries)} entri terakhir)</span></h3>
    {table_html}
  </div>
</div>"""

    return """<!DOCTYPE html>
<html lang="id"><head><meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Riwayat Penggunaan &mdash; e-GeRAI KKPRL</title>
<style>""" + LANDING_CSS + REVIEW_CSS + """</style></head>
<body>
""" + HEADER_HTML + """

<section class="review-hero">
  <h1>""" + ICONS["chart-bar"] + """ Riwayat Penggunaan</h1>
  <p>Catatan dokumen PKKPRL yang sudah di-generate melalui aplikasi ini (nama pemohon, waktu, dan siapa yang memproses). Dokumen itu sendiri tidak disimpan.</p>
</section>
""" + body + """
</body></html>"""


def render_manual_form_page(error=None, prefill_data=None, job_id=None, saved_images_meta=None, saved_msg=None):
    prop_groups = []
    for group_title, fields in FIELD_GROUPS:
        prop_fields = [f for f in fields if f[0] in ("prop", "prop_loc")]
        if prop_fields:
            prop_groups.append((group_title, prop_fields))

    groups_html = []
    for i, (group_title, fields) in enumerate(prop_groups):
        rows = []
        for source, key, label in fields:
            fname = form_field_name(source, key)
            help_text = FIELD_HELP.get((source, key), "")
            help_html = f'<div class="ff-hint" style="margin:2px 0 6px;">{help_text}</div>' if help_text else ""
            example = EXAMPLE_HINTS.get((source, key), "") if not SELECT_FIELDS.get((source, key)) else ""
            example_html = ""
            if example:
                example_html = (
                    f'<div class="field-example">Contoh: <span class="ex-text">{example}</span>'
                    f'<button type="button" class="ex-fill" data-target="{fname}">Pakai contoh ini</button></div>'
                )
            select_cfg = SELECT_FIELDS.get((source, key))
            if select_cfg:
                opts_html = ['<option value="">-- Pilih --</option>']
                for opt in select_cfg["options"]:
                    if source == "prop_loc" and key == "3":
                        bps_id = PROVINCE_BPS_IDS.get(opt, "")
                        opts_html.append(f'<option value="{opt}" data-bps-id="{bps_id}">{opt}</option>')
                    else:
                        opts_html.append(f'<option value="{opt}">{opt}</option>')
                other_html = ""
                if select_cfg["allow_other"]:
                    opts_html.append('<option value="__other__">Lainnya...</option>')
                    other_html = (
                        f'<input type="text" class="select-other-input" '
                        f'id="{fname}_other" placeholder="Isi nama lainnya" '
                        f'data-for="{fname}">'
                    )
                field_html = (
                    f'<select class="select-field" name="{fname}" id="{fname}" data-allow-other="{str(select_cfg["allow_other"]).lower()}">'
                    f'{"".join(opts_html)}</select>{other_html}'
                )
            elif source == "prop_loc" and key == "2":
                field_html = (
                    f'<input type="text" name="{fname}" id="{fname}" list="dl_kabupaten" '
                    f'placeholder="Isi {label.lower()}" autocomplete="off">'
                    f'<datalist id="dl_kabupaten"></datalist>'
                )
            elif source == "prop_loc" and key == "1":
                field_html = (
                    f'<input type="text" name="{fname}" id="{fname}" list="dl_kecamatan" '
                    f'placeholder="Isi {label.lower()}" autocomplete="off">'
                    f'<datalist id="dl_kecamatan"></datalist>'
                )
            elif source == "prop_loc" and key == "0":
                field_html = (
                    f'<input type="text" name="{fname}" id="{fname}" list="dl_desa" '
                    f'placeholder="Isi {label.lower()}" autocomplete="off">'
                    f'<datalist id="dl_desa"></datalist>'
                )
            elif source == "prop" and key == "investasi":
                field_html = (
                    f'<div class="money-field"><span class="money-prefix">Rp</span>'
                    f'<input type="text" name="{fname}" id="{fname}" class="money-input" '
                    f'inputmode="numeric" placeholder="Isi {label.lower()}"></div>'
                )
            elif source == "prop" and key == "NPWP":
                field_html = (
                    f'<input type="text" name="{fname}" id="{fname}" class="npwp-input" '
                    f'inputmode="numeric" placeholder="Isi {label.lower()}">'
                )
            elif source == "prop" and key == "Nomor Telepon Selular":
                field_html = (
                    f'<input type="text" name="{fname}" id="{fname}" class="digits-only-input" '
                    f'inputmode="numeric" placeholder="Isi {label.lower()}">'
                )
            elif source == "prop" and key in ("NIB", "tenaga_kerja", "tenaga_kerja_asing"):
                field_html = (
                    f'<input type="text" name="{fname}" id="{fname}" class="digits-only-input" '
                    f'inputmode="numeric" placeholder="Isi {label.lower()}">'
                )
            elif source == "prop" and key in ("Luas Kebutuhan Ruang", "desa_luas_ha", "desa_penduduk"):
                field_html = (
                    f'<input type="text" name="{fname}" id="{fname}" class="decimal-only-input" '
                    f'inputmode="decimal" placeholder="Isi {label.lower()}">'
                )
            elif source == "prop" and key == "Tanggal Penyusunan":
                field_html = (
                    f'<input type="date" id="{fname}_picker" class="date-picker-input">'
                    f'<input type="hidden" name="{fname}" id="{fname}">'
                )
            else:
                field_html = f'<input type="text" name="{fname}" id="{fname}" placeholder="Isi {label.lower()}">'
            rows.append(
                f'<div class="field-row"><label>{label}</label>{help_html}'
                f'{field_html}'
                f'{example_html}</div>'
            )
        groups_html.append(
            f'<details class="acc-item"{" open" if i == 0 else ""}>'
            f"<summary>{group_title}</summary>"
            f'<div class="acc-body">{"".join(rows)}</div></details>'
        )

    error_html = f'<div class="error-banner" style="margin-bottom:16px;">\u26A0 {error}</div>' if error else ""

    if saved_msg:
        error_html += f'<div class="error-banner" style="margin-bottom:16px;background:#eaf6ec;border-color:#b7e4c7;">\u2705 {saved_msg}</div>'

    if saved_images_meta:
        tags_txt = ", ".join(m.get("tag", "") for m in saved_images_meta)
        error_html += (
            '<div class="error-banner" style="margin-bottom:16px;background:#eaf6ec;border-color:#b7e4c7;">'
            f"\U0001F4CE {len(saved_images_meta)} gambar tersimpan dari isian sebelumnya ({tags_txt}) &mdash; "
            "tetap akan dipakai walau tidak diunggah ulang. Unggah ulang di kolom terkait kalau ingin menggantinya."
            "</div>"
        )

    # Peta tag gambar -> daftar preview (URL thumbnail) untuk gambar yang
    # sudah tersimpan lewat tombol "Simpan" sebelumnya, dikelompokkan per
    # NAMA FIELD form (bukan per tag) supaya gampang disisipkan ke tiap
    # pemanggilan img_field_html() di bawah.
    by_tag = {}
    for m in (saved_images_meta or []):
        by_tag.setdefault(m.get("tag"), []).append(m)
    saved_previews_map = {}
    if job_id:
        for field_name, tag in IMAGE_FIELD_TAGS:
            metas = by_tag.get(tag)
            if metas:
                saved_previews_map[field_name] = [
                    {"url": f"/proposal-manual/draft-image/{job_id}/{m['filename']}"} for m in metas
                ]

    return """<!DOCTYPE html>
<html lang="id"><head><meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Isi Formulir Draft Proposal &mdash; e-GeRAI KKPRL</title>
<style>""" + LANDING_CSS + REVIEW_CSS + """</style></head>
<body>
""" + HEADER_HTML + """

<section class="manual-hero">
  <h1>\U0001F4DD Isi Formulir Draft Proposal PKKPRL</h1>
  <p>Belum punya file Draft Proposal PKKPRL siap pakai? Isi data di bawah ini secara manual, lalu unggah
  Laporan Kondisi Eksisting / Hidro-Oseanografi (PDF). Sistem akan menggabungkan otomatis menjadi
  1 dokumen Word final &mdash; sama seperti alur upload 2 PDF.</p>
</section>

<div class="review-wrap">
  """ + error_html + """
  <form method="POST" action="/proposal-manual" enctype="multipart/form-data" id="manualForm">
    <input type="hidden" name="existing_job_id" value=\"""" + (job_id or "") + """\">

    <div class="manual-upload-card">
      <h3>""" + ICONS["wave"] + """ Laporan Kondisi Eksisting / Hidro-Oseanografi (PDF/Word) <span style="font-weight:400;color:var(--muted);font-size:12px;">&mdash; Opsional</span></h3>
      <div class="ff-hint" style="margin-bottom:10px;">Belum punya dokumennya? Peroleh data Hidro-Oseanografi melalui portal
      <a href="https://huggingface.co/spaces/Fadly2002/Gerai-Pelayanan-BPRL" target="_blank" style="color:var(--blue);font-weight:700;">Gerai Pelayanan Balai Penataan Ruang Laut Makassar</a>,
      unduh hasilnya (PDF atau Word), lalu unggah di bawah ini. Belum sempat siap? Boleh dikosongkan dulu &mdash; pakai tombol <b>"Unduh Draft"</b> di bawah untuk mengunduh draft Proposal saja terlebih dulu, lengkapi Laporannya nanti.</div>
      <div class="dropzone" id="dzManual">
        """ + ICONS["cloud"] + """
        <div class="dz-title">Drag &amp; Drop PDF/Word di sini</div>
        <div class="dz-sub">atau klik untuk memilih file</div>
        <div class="dz-max">Maksimum 10 MB &middot; PDF atau .docx</div>
        <input type="file" name="laporan" id="laporanManual" accept="application/pdf,.docx">
      </div>
      <div class="file-chip" id="chipManual">
        <div class="fc-left">""" + ICONS["check-circle"] + """<span class="fc-name" id="nameManual"></span></div>
        <span class="fc-size" id="sizeManual"></span>
        <button type="button" class="fc-remove" data-target="Manual">""" + ICONS["x"] + """</button>
      </div>
    </div>

    <div class="sticky-bar">
      <div class="sticky-bar-row">
        <button type="submit" formaction="/proposal-manual/simpan" formnovalidate class="btn-draft">""" + ICONS["check-circle"] + """ Simpan</button>
        <button type="submit">""" + ICONS["bolt"] + """ Proses &amp; Lanjut ke Tinjau Data</button>
        <button type="submit" formaction="/proposal-manual/draft" formnovalidate class="btn-draft">""" + ICONS["download"] + """ Unduh Draft</button>
      </div>
    </div>

    <div class="review-card">
      <h3>""" + ICONS["doc"] + """ Data Draft Proposal PKKPRL</h3>
      """ + "".join(groups_html) + """

      <details class="acc-item">
        <summary>Deskripsi Kegiatan</summary>
        <div class="acc-body">
          <div class="field-row">
            <label>Deskripsi Kegiatan</label>
            <div class="ff-hint" style="margin-bottom:6px;">Mohon untuk mengisi dengan deskripsi kegiatan yang akan dilakukan.</div>
            <textarea name="deskripsi_kegiatan" id="deskripsi_kegiatan" rows="4" placeholder="Isi deskripsi kegiatan"></textarea>
            <div class="field-example">Contoh: <span class="ex-text">Kegiatan usaha yang diusulkan adalah pembesaran biota laut budidaya, melalui pengoperasian keramba jaring apung (KJA) sebagai sarana penampungan dan pemeliharaan sementara ikan hidup sebelum dipasarkan. Biota Laut memiliki nilai ekonomi tinggi dengan peluang pasar yang masih terbuka, sehingga kegiatan ini berpotensi memberikan nilai tambah hasil perikanan serta menjadi alternatif diversifikasi usaha bagi nelayan setempat. Pelaksanaan kegiatan diharapkan dapat meningkatkan pendapatan dan kesejahteraan masyarakat pesisir serta mengurangi ketergantungan terhadap jenis ikan lainnya.</span>
            <button type="button" class="ex-fill" data-target="deskripsi_kegiatan">Pakai contoh ini</button></div>
          </div>

          <div class="field-row">
            <label>Manfaat Kegiatan</label>
            <div class="ff-hint" style="margin-bottom:6px;">Mohon untuk mengisi dengan deskripsi manfaat dari kegiatan yang akan dilakukan.</div>
            <textarea name="manfaat_kegiatan" id="manfaat_kegiatan" rows="4" placeholder="Isi manfaat kegiatan"></textarea>
            <div class="field-example">Contoh: <span class="ex-text">Kegiatan pembangunan dan operasional fasilitas budidaya udang vannamei bertujuan untuk mendukung peningkatan produksi perikanan budidaya secara berkelanjutan melalui pemanfaatan ruang laut yang optimal dan sesuai dengan ketentuan yang berlaku. Selain memberikan nilai tambah bagi sektor perikanan, kegiatan ini juga diharapkan dapat mendorong pertumbuhan ekonomi daerah, membuka peluang kerja bagi masyarakat sekitar, serta mendukung penerapan budidaya yang produktif dan berwawasan lingkungan.</span>
            <button type="button" class="ex-fill" data-target="manfaat_kegiatan">Pakai contoh ini</button></div>
          </div>

          <div class="field-row">
            <label>Tujuan Kegiatan</label>
            <div class="ff-hint" style="margin-bottom:6px;">Mohon untuk mengisi dengan deskripsi tujuan dari kegiatan yang akan dilakukan.</div>
            <textarea name="tujuan_kegiatan" id="tujuan_kegiatan" rows="4" placeholder="Isi tujuan kegiatan"></textarea>
            <div class="field-example">Contoh: <span class="ex-text">Tujuan pemanfaatan ruang laut yang diajukan adalah untuk pembangunan fasilitas pemanfaatan air laut bagi kegiatan budidaya, yang berfungsi sebagai sarana penunjang usaha pembudidayaan ikan bersirip (selain ikan hias) serta biota air payau lainnya yang tidak dilindungi. Kegiatan utama perusahaan adalah budidaya udang vannamei. Dalam mendukung pelaksanaan kegiatan utama tersebut, direncanakan pembangunan fasilitas pendukung yang meliputi Instalasi Pengolahan Air Limbah (IPAL), instalasi penyediaan air bersih, dan instalasi kelistrikan.</span>
            <button type="button" class="ex-fill" data-target="tujuan_kegiatan">Pakai contoh ini</button></div>
          </div>

          <div class="field-row">
            <label>Instalasi Bangunan Menetap Di Laut</label>
            <div class="ff-hint" style="margin-bottom:6px;">Contoh: Saluran Inlet atau Outlet</div>
            <input type="text" name="instalasi_bangunan" id="instalasi_bangunan" placeholder="Isi instalasi bangunan menetap di laut">
          </div>

          <div class="field-row">
            <label>Instalasi Bangunan Laut Berada Pada</label>
            <div class="ff-hint" style="margin-bottom:6px;">Bisa pilih lebih dari satu kalau instalasi berada di beberapa posisi sekaligus.</div>
            <div class="checkbox-row"><input type="checkbox" name="instalasi_posisi[]" value="Permukaan Laut" id="instpos_permukaan"><label for="instpos_permukaan">Permukaan Laut</label></div>
            <div class="checkbox-row"><input type="checkbox" name="instalasi_posisi[]" value="Kolom Laut" id="instpos_kolom"><label for="instpos_kolom">Kolom Laut</label></div>
            <div class="checkbox-row"><input type="checkbox" name="instalasi_posisi[]" value="Dasar Laut" id="instpos_dasar"><label for="instpos_dasar">Dasar Laut</label></div>
          </div>

          <div class="field-row">
            <label>Deskripsi Jadwal Kegiatan</label>
            <div class="ff-hint" style="margin-bottom:6px;">Tuliskan rincian kegiatan dan waktu pelaksanaannya, format: [Nama Kegiatan] : [Bulan/Tahun Pelaksanaan]</div>
            <textarea name="jadwal_kegiatan" id="jadwal_kegiatan" rows="4" placeholder="Isi jadwal kegiatan"></textarea>
            <div class="field-example">Contoh: <span class="ex-text">Pengurusan PKKPRL : Bulan 1 - Bulan 3. Pemasangan Keramba Jaring Apung : Bulan 3 - Bulan 5. Operasional Keramba Jaring Apung : Bulan 5 - Bulan 12.</span>
            <button type="button" class="ex-fill" data-target="jadwal_kegiatan">Pakai contoh ini</button></div>
          </div>

          <div class="field-row">
            <label>Dokumen Data Dukung</label>
            <div class="ff-hint" style="margin-bottom:6px;">Dokumen data dukung yang dimiliki oleh pelaku usaha (centang yang sesuai -- akan otomatis masuk ke bagian "IV. Dokumen Persyaratan Lainnya" di draft dokumen)</div>
            """ + dukung_item_html(1, "dukung_nib", "NIB") + """
            """ + dukung_item_html(2, "dukung_sertifikat", "Sertifikat Kepemilikan Lahan Darat") + """
            """ + dukung_item_html(3, "dukung_izin_lingkungan", "Surat Izin Lingkungan") + """
            """ + dukung_item_html(4, "dukung_ba_sosialisasi", "Berita Acara Sosialisasi") + """
            """ + dukung_item_html(5, "dukung_identitas", "Dokumen Identitas dan Legalitas Pemohon/Perusahaan") + """
            """ + dukung_item_html(6, "dukung_survei", "Dokumentasi Survei Lapangan Kondisi Eksisting Lokasi") + """
            """ + dukung_item_html(7, "dukung_peta", "Peta Pendukung (Peta Lokasi, Site Plan, Pola Ruang Wilayah)") + """
            """ + dukung_item_html(8, "dukung_dipa", "DIPA/RKAKL (Sumber Anggaran APBD/APBN) / Lainnya") + """
            """ + dukung_item_html(9, "dukung_sk_kkprl", "SK Penetapan KNMP") + """
            <div id="dukungLainnyaWrap"></div>
            <button type="button" id="btnTambahDukung" class="ex-fill" style="margin-top:8px;">+ Tambah Dokumen Lainnya</button>
          </div>
        </div>
      </details>

      <details class="acc-item">
        <summary>Status Kegiatan</summary>
        <div class="acc-body">
          <div class="field-row">
            <label>Kegiatan Eksisting/Rencana</label>
            <div class="ff-hint" style="margin-bottom:6px;">Apakah kegiatan yang dimohonkan merupakan kegiatan eksisting atau baru akan direncanakan</div>
            <select class="select-field" name="kegiatan_status" id="kegiatan_status">
              <option value="">-- Pilih --</option>
              <option value="Eksisting">Eksisting</option>
              <option value="Rencana">Rencana</option>
              <option value="Eksisting dan Pengembangan">Eksisting dan Pengembangan</option>
            </select>
          </div>
          <div class="ff-hint" style="margin:0 0 8px;">Centang salah satu dari tiap pasangan berikut sesuai kondisi kegiatan yang dimohonkan. Kalau tidak dicentang sama sekali, bagian terkait di draft akan ditandai perlu dilengkapi manual (tidak ditebak otomatis).</div>
          <div class="checkbox-row"><input type="checkbox" name="non_reklamasi" id="cb1"><label for="cb1">Kegiatan Tanpa Reklamasi</label></div>
          <div class="checkbox-row"><input type="checkbox" name="reklamasi" id="cb1b"><label for="cb1b">Kegiatan Reklamasi</label></div>
          <div class="checkbox-row"><input type="checkbox" name="kegiatan_berusaha" id="cb2"><label for="cb2">Termasuk Kegiatan Berusaha</label></div>
          <div class="checkbox-row"><input type="checkbox" name="non_berusaha" id="cb2b"><label for="cb2b">Kegiatan Non Berusaha</label></div>
          <div class="checkbox-row"><input type="checkbox" name="non_strategis" id="cb3"><label for="cb3">Termasuk kegiatan non-strategis nasional</label></div>
        </div>
      </details>

      <details class="acc-item">
        <summary>Data Ekosistem Tambahan</summary>
        <div class="acc-body">
          <div class="field-row">
            <label>Keberadaan Ekosistem Mangrove</label>
            """ + render_select_html("mangrove_ada", SELECT_FIELDS[("prop", "mangrove_ada")]["options"]).replace('id="mangrove_ada"', 'id="mangrove_ada" data-eco-trigger="1"') + """
          </div>
          <div class="eco-children locked" data-for="mangrove_ada">
            <div class="field-row">
              <label>Spesies Mangrove Dominan</label>
              <input type="text" name="prop__mangrove_spesies" id="prop__mangrove_spesies" placeholder="Isi spesies mangrove" disabled>
              <div class="field-example">Contoh: <span class="ex-text">Rhizophora mucronata</span>
              <button type="button" class="ex-fill" data-target="prop__mangrove_spesies">Pakai contoh ini</button></div>
            </div>
            <div class="field-row">
              <label>Persentase Tutupan Mangrove (%)</label>
              <input type="text" name="prop__mangrove_persen" id="prop__mangrove_persen" class="decimal-only-input" inputmode="decimal" placeholder="Isi persentase tutupan mangrove" disabled>
            </div>
            <div class="field-row">
              <label>Kondisi Tutupan Mangrove</label>
              """ + render_select_html("prop__mangrove_kondisi", SELECT_FIELDS[("prop", "mangrove_kondisi")]["options"]).replace('id="prop__mangrove_kondisi"', 'id="prop__mangrove_kondisi" disabled') + """
            </div>
          </div>

          <div class="field-row">
            <label>Keberadaan Ekosistem Lamun</label>
            """ + render_select_html("lamun_ada_manual", SELECT_FIELDS[("prop", "lamun_ada_manual")]["options"]).replace('id="lamun_ada_manual"', 'id="lamun_ada_manual" data-eco-trigger="1"') + """
          </div>
          <div class="eco-children locked" data-for="lamun_ada_manual">
            <div class="field-row">
              <label>Spesies Lamun</label>
              <input type="text" name="lamun_spesies" id="lamun_spesies" placeholder="Isi spesies lamun" disabled>
            </div>
            <div class="field-row">
              <label>Persentase Tutupan Lamun</label>
              <div class="ff-hint" style="margin-bottom:6px;">Dilampirkan dalam bentuk persentase (%)</div>
              <input type="text" name="lamun_persen" id="lamun_persen" class="decimal-only-input" inputmode="decimal" placeholder="Isi persentase tutupan lamun" disabled>
            </div>
            <div class="field-row">
              <label>Kondisi Lamun</label>
              """ + render_select_html("lamun_kondisi", SELECT_FIELDS[("prop", "lamun_kondisi")]["options"]).replace('id="lamun_kondisi"', 'id="lamun_kondisi" disabled') + """
            </div>
          </div>

          <div class="field-row">
            <label>Keberadaan Ekosistem Terumbu Karang</label>
            """ + render_select_html("karang_ada", SELECT_FIELDS[("prop", "karang_ada")]["options"]).replace('id="karang_ada"', 'id="karang_ada" data-eco-trigger="1"') + """
          </div>
          <div class="eco-children locked" data-for="karang_ada">
            <div class="field-row">
              <label>Spesies Terumbu Karang</label>
              <input type="text" name="karang_spesies" id="karang_spesies" placeholder="Isi spesies terumbu karang" disabled>
            </div>
            <div class="field-row">
              <label>Persentase Tutupan Terumbu Karang</label>
              <div class="ff-hint" style="margin-bottom:6px;">Dicantumkan dalam bentuk %</div>
              <input type="text" name="karang_persen_manual" id="karang_persen_manual" class="decimal-only-input" inputmode="decimal" placeholder="Isi persentase tutupan terumbu karang" disabled>
            </div>
            <div class="field-row">
              <label>Kondisi Terumbu Karang</label>
              """ + render_select_html("karang_kondisi", SELECT_FIELDS[("prop", "karang_kondisi")]["options"]).replace('id="karang_kondisi"', 'id="karang_kondisi" disabled') + """
            </div>
          </div>
        </div>
      </details>

      <details class="acc-item">
        <summary>Pemanfaatan Ruang Laut Sekitar (Opsional)</summary>
        <div class="acc-body">
          <div class="ff-hint" style="margin-bottom:12px;">Isi kondisi di 4 penjuru arah mata angin di sekitar lokasi kegiatan. Akan otomatis tersusun jadi kalimat deskripsi di dokumen final.</div>
          <div class="field-row">
            <label>Sebelah Utara</label>
            <input type="text" name="batas_utara" id="batas_utara" placeholder="Contoh: area penangkapan ikan skala kecil dan kawasan pemukiman nelayan berjarak sekitar 1 km">
          </div>
          <div class="field-row">
            <label>Sebelah Timur</label>
            <input type="text" name="batas_timur" id="batas_timur" placeholder="Contoh: area penangkapan ikan skala kecil dan kawasan pelabuhan lokal berjarak sekitar 1.2 km">
          </div>
          <div class="field-row">
            <label>Sebelah Selatan</label>
            <input type="text" name="batas_selatan" id="batas_selatan" placeholder="Contoh: area penangkapan ikan skala kecil serta koridor kabel/pipa bawah laut berjarak sekitar 1.4 km">
          </div>
          <div class="field-row">
            <label>Sebelah Barat</label>
            <input type="text" name="batas_barat" id="batas_barat" placeholder="Contoh: area penangkapan ikan skala kecil serta kegiatan Keramba Jaring Apung (KJA) berjarak sekitar 1.1 km">
          </div>
          <div class="field-row">
            <label>Deskripsi Tambahan (Opsional)</label>
            <div class="ff-hint" style="margin-bottom:6px;">Kalau ada info tambahan di luar 4 arah mata angin di atas, isi di sini -- akan ditambahkan setelah kalimat otomatis.</div>
            <textarea name="deskripsi_pemanfaatan_sekitar" id="deskripsi_pemanfaatan_sekitar" rows="3" placeholder="Isi deskripsi tambahan (opsional)"></textarea>
          </div>
        </div>
      </details>

      <details class="acc-item">
        <summary>Sosial Ekonomi &amp; Aksesibilitas Lanjutan (Opsional)</summary>
        <div class="acc-body">
          <div class="field-row">
            <label>Mata Pencaharian Masyarakat Desa</label>
            <textarea name="mata_pencaharian" id="mata_pencaharian" rows="4" placeholder="Isi mata pencaharian masyarakat desa"></textarea>
            <div class="field-example">Contoh: <span class="ex-text">Mata pencaharian masyarakat desa didominasi oleh aktivitas yang berkaitan dengan karakter pesisir. Nelayan menjadi salah satu pekerjaan utama, didukung oleh potensi perairan yang memiliki sumber daya ikan, biota laut, padang lamun, mangrove, dan terumbu karang.</span>
            <button type="button" class="ex-fill" data-target="mata_pencaharian">Pakai contoh ini</button></div>
          </div>
          <div class="field-row">
            <label>Sumber Data Sosek</label>
            <input type="text" name="sumber_data_sosek" id="sumber_data_sosek" placeholder="Isi sumber data sosial ekonomi">
            <div class="field-example">Contoh: <span class="ex-text">Badan Pusat Statistik</span>
            <button type="button" class="ex-fill" data-target="sumber_data_sosek">Pakai contoh ini</button></div>
          </div>
          <div class="field-row">
            <label>Tahun Data Sosek</label>
            <input type="text" name="tahun_data_sosek" id="tahun_data_sosek" class="digits-only-input" inputmode="numeric" placeholder="Isi tahun data">
          </div>
          <div class="field-row">
            <label>Aksesibilitas Lokasi</label>
            <div class="ff-hint" style="margin-bottom:6px;">Deskripsi aksesibilitas dari titik poin lokasi yang mudah dikenali ke lokasi area yang dimohonkan, termasuk jarak dan waktu tempuh.</div>
            <textarea name="aksesibilitas_lokasi" id="aksesibilitas_lokasi" rows="4" placeholder="Isi aksesibilitas lokasi"></textarea>
            <div class="field-example">Contoh: <span class="ex-text">Aksesibilitas menuju Desa Tapulaga, Kecamatan Soropia, Kabupaten Konawe, dari Bandara Haluoleo Kendari dapat ditempuh melalui jalur darat. Titik awal perjalanan adalah Bandara Haluoleo yang berada di Desa Ambaipua, Kecamatan Ranomeeto, Kabupaten Konawe Selatan, dengan akses utama melalui Jalan Wolter Monginsidi/poros bandara menuju Kota Kendari. Dari kawasan bandara, perjalanan dilanjutkan menuju pusat Kota Kendari melalui koridor Ranomeeto&ndash;Baruga&ndash;Kendari. Setelah memasuki Kota Kendari, perjalanan diarahkan ke kawasan Kota Lama/Kecamatan Kendari, kemudian berlanjut menuju jalur Kendari&ndash;Toronipa yang menjadi akses utama ke wilayah pesisir Kecamatan Soropia.</span>
            <button type="button" class="ex-fill" data-target="aksesibilitas_lokasi">Pakai contoh ini</button></div>
          </div>
        </div>
      </details>

      <details class="acc-item">
        <summary>Titik Koordinat Batas Area (Opsional)</summary>
        <div class="acc-body">
          <div class="field-row">
            <label>Upload File Koordinat (Excel/CSV/Word/Gambar)</label>
            <div class="ff-hint" style="margin-bottom:6px;">Kalau sudah punya file/screenshot berisi daftar titik koordinat, unggah di sini -- akan otomatis dibaca dan dianalisis, tidak perlu ketik ulang manual. Untuk format gambar (PNG/JPG), pembacaan otomatis pakai AI dan butuh beberapa detik.</div>
            <input type="file" name="upload_koordinat" id="upload_koordinat" accept=".xlsx,.xlsm,.csv,.docx,.png,.jpg,.jpeg">
          </div>
          <div class="field-row">
            <label>Atau Ketik/Tempel Manual &mdash; format: Longitude [spasi] Latitude (nomor titik otomatis)</label>
            <div class="ff-hint" style="margin-bottom:6px;">Sesuai format resmi. Kalau file di atas diisi, ini akan digabung otomatis dengan hasil dari file.</div>
            <textarea name="koordinat_manual" id="koordinat_manual" rows="4" placeholder="122.650194        -3.934945&#10;122.649197        -3.935361&#10;122.649261        -3.935530&#10;122.650258        -3.935114"></textarea>
            <div class="field-example">Contoh: <span class="ex-text">122.650194        -3.934945&#10;122.649197        -3.935361&#10;122.649261        -3.935530&#10;122.650258        -3.935114</span>
            <button type="button" class="ex-fill" data-target="koordinat_manual">Pakai contoh ini</button></div>
          </div>
        </div>
      </details>

      <details class="acc-item">
        <summary>Lampiran Gambar (Opsional)</summary>
        <div class="acc-body">
          <div class="ff-hint" style="margin-bottom:14px;">Ada 2 cara mengisi tiap gambar &mdash; <b>Upload File</b> untuk pilih dari komputer, atau kotak <b>Ctrl+V</b> untuk paste dari clipboard (misal screenshot). Keduanya bisa dipakai berkali-kali secara bergantian; file akan terus bertambah, tidak saling mengganti.</div>

          """ + img_field_html("img_siteplan", "Gambaran Rencana Tapak Site",
                                "Unggah gambaran rencana tapak site dari kegiatan yang dimohonkan. Maks. 10 MB.",
                                note="(bisa lebih dari 1)", saved_previews=saved_previews_map.get("img_siteplan")) + """

          """ + img_field_html("img_peta_lokasi", "Peta Lokasi",
                                "Unggah visualisasi peta lokasi yang dimohonkan dalam bentuk citra satelit yang telah dilengkapi dengan poligon batas area permohonan. Maks. 10 MB.",
                                saved_previews=saved_previews_map.get("img_peta_lokasi")) + """

          <div class="field-row">
            <label>Sumber Peta</label>
            <div class="ff-hint" style="margin-bottom:6px;">Mohon untuk mencantumkan sumber peta yang diambil</div>
            <input type="text" name="sumber_peta" id="sumber_peta" placeholder="Isi sumber peta">
            <div class="field-example">Contoh: <span class="ex-text">Layar Pinisi, Arcgis, Google Earth dll</span>
            <button type="button" class="ex-fill" data-target="sumber_peta">Pakai contoh ini</button></div>
          </div>

          """ + img_field_html("img_foto_mangrove", "Foto Kondisi Mangrove",
                                saved_previews=saved_previews_map.get("img_foto_mangrove")) + """

          """ + img_field_html("img_foto_karang_insitu", "Foto Survei Terumbu Karang",
                                saved_previews=saved_previews_map.get("img_foto_karang_insitu")) + """

          """ + img_field_html("img_dok_kegiatan", "Dokumentasi Kegiatan Eksisting/Rencana",
                                "Unggah gambar eksisting atau rencana dari kegiatan yang dimohonkan.",
                                saved_previews=saved_previews_map.get("img_dok_kegiatan")) + """

          """ + img_field_html("img_dok_pemanfaatan_sekitar", "Dokumentasi Pemanfaatan Ruang Laut Sekitar",
                                "Maksimal 3 dokumentasi.", note="(bisa lebih dari 1)",
                                saved_previews=saved_previews_map.get("img_dok_pemanfaatan_sekitar")) + """

          """ + img_field_html("img_foto_lamun", "Dokumentasi Ekosistem Lamun",
                                saved_previews=saved_previews_map.get("img_foto_lamun")) + """

          """ + img_field_html("img_aksesibilitas", "Gambar Peta Aksesibilitas Menuju Lokasi",
                                saved_previews=saved_previews_map.get("img_aksesibilitas")) + """

          """ + img_field_html("img_sertifikat_lahan", "Sertifikat Kepemilikan Lahan Darat",
                                saved_previews=saved_previews_map.get("img_sertifikat_lahan")) + """

          """ + img_field_html("img_dok_sosialisasi", "Dokumen Hasil Sosialisasi",
                                "Berita acara atau surat pernyataan tidak keberatan dari masyarakat.",
                                saved_previews=saved_previews_map.get("img_dok_sosialisasi")) + """

          """ + img_field_html("img_dok_pendukung_lainnya", "Dokumen Pendukung Lainnya", note="(bisa lebih dari 1)",
                                saved_previews=saved_previews_map.get("img_dok_pendukung_lainnya")) + """
        </div>
      </details>
    </div>
  </form>

  <a href="/" class="back-link">&larr; Kembali ke halaman utama (unggah 2 PDF)</a>
</div>

<script>
function fmtSize(bytes) {
  if (bytes >= 1024*1024) return (bytes/(1024*1024)).toFixed(2) + " MB";
  return Math.ceil(bytes/1024) + " KB";
}
(function() {
  var dz = document.getElementById('dzManual');
  var input = document.getElementById('laporanManual');
  var chip = document.getElementById('chipManual');
  var nameEl = document.getElementById('nameManual');
  var sizeEl = document.getElementById('sizeManual');
  function showFile(file) {
    if (!file) { chip.classList.remove('show'); dz.style.display = 'block'; return; }
    nameEl.textContent = file.name;
    nameEl.title = file.name;
    sizeEl.textContent = fmtSize(file.size);
    chip.classList.add('show');
    dz.style.display = 'none';
  }
  dz.addEventListener('click', function(e) { if (!e.target.closest('.file-chip')) input.click(); });
  input.addEventListener('change', function() { if (input.files[0]) showFile(input.files[0]); });
  dz.addEventListener('dragover', function(e) { e.preventDefault(); dz.classList.add('dragover'); });
  dz.addEventListener('dragleave', function() { dz.classList.remove('dragover'); });
  dz.addEventListener('drop', function(e) {
    e.preventDefault(); dz.classList.remove('dragover');
    if (e.dataTransfer.files[0]) { input.files = e.dataTransfer.files; showFile(input.files[0]); }
  });
  document.querySelector('.fc-remove[data-target="Manual"]').addEventListener('click', function() {
    input.value = ''; dz.style.display = 'block'; chip.classList.remove('show');
  });
})();

// Tombol "Pakai contoh ini" -> isi input teks dengan contoh, siap diedit
document.querySelectorAll('.ex-fill').forEach(function(btn) {
  btn.addEventListener('click', function() {
    var target = document.getElementById(btn.dataset.target);
    var exampleText = btn.previousElementSibling.textContent;
    if (target) { target.value = exampleText; target.focus(); }
  });
});

// Nilai Investasi: format ribuan otomatis pakai titik saat mengetik (200000000 -> 200.000.000)
document.querySelectorAll('.money-input').forEach(function(input) {
  input.addEventListener('input', function() {
    var digits = input.value.replace(/\D/g, '');
    input.value = digits ? digits.replace(/\B(?=(\d{3})+(?!\d))/g, '.') : '';
  });
});

// NPWP: kunci hanya boleh angka, titik, dan strip (format resmi: 01.234.567.8-901.000)
document.querySelectorAll('.npwp-input').forEach(function(input) {
  input.addEventListener('input', function() {
    input.value = input.value.replace(/[^0-9.\-]/g, '');
  });
});

// Nomor Telepon Selular: kunci hanya boleh angka
document.querySelectorAll('.digits-only-input').forEach(function(input) {
  input.addEventListener('input', function() {
    input.value = input.value.replace(/\D/g, '');
  });
});

// Field desimal (persentase, luas area, dsb): angka + maksimal 1 titik desimal
document.querySelectorAll('.decimal-only-input').forEach(function(input) {
  input.addEventListener('input', function() {
    var v = input.value.replace(/[^0-9.]/g, '');
    var parts = v.split('.');
    if (parts.length > 2) { v = parts[0] + '.' + parts.slice(1).join(''); }
    input.value = v;
  });
});

// Auto-klasifikasi kondisi ekosistem (mangrove/lamun/karang) berdasarkan
// persentase yang diinput, sesuai kriteria baku resmi (KRITERIA_KONDISI).
// User tetap bisa ubah manual kalau perlu -- ini cuma bantu isi otomatis.
(function() {
  var KRITERIA = {
    'prop__mangrove_persen': { select: 'prop__mangrove_kondisi', rules: [[75,999,'Sangat Padat'],[50,75,'Sedang'],[0,50,'Jarang']] },
    'lamun_persen': { select: 'lamun_kondisi', rules: [[60,999,'Baik (Kaya/Sehat)'],[30,60,'Rusak (Kurang Kaya/Kurang Sehat)'],[0,30,'Rusak (Miskin)']] },
    'karang_persen_manual': { select: 'karang_kondisi', rules: [[75,999,'Baik Sekali'],[50,75,'Baik'],[25,50,'Sedang'],[0,25,'Buruk']] },
  };
  Object.keys(KRITERIA).forEach(function(inputId) {
    var input = document.getElementById(inputId);
    var cfg = KRITERIA[inputId];
    var select = document.getElementById(cfg.select);
    if (!input || !select) return;
    input.addEventListener('input', function() {
      var v = parseFloat(input.value.replace(',', '.'));
      if (isNaN(v)) return;
      for (var i = 0; i < cfg.rules.length; i++) {
        var lo = cfg.rules[i][0], hi = cfg.rules[i][1], label = cfg.rules[i][2];
        if (v >= lo && v < hi) {
          select.value = label;
          break;
        }
      }
    });
  });
})();

// Tanggal Penyusunan: pilih lewat kalender, otomatis dikonversi ke format Indonesia (DD Bulan YYYY)
document.querySelectorAll('.date-picker-input').forEach(function(picker) {
  var hidden = document.getElementById(picker.id.replace('_picker', ''));
  var BULAN_ID = ['Januari','Februari','Maret','April','Mei','Juni','Juli','Agustus','September','Oktober','November','Desember'];
  picker.addEventListener('change', function() {
    if (!picker.value) { hidden.value = ''; return; }
    var parts = picker.value.split('-'); // YYYY-MM-DD
    var y = parts[0], m = parseInt(parts[1], 10) - 1, d = parseInt(parts[2], 10);
    hidden.value = d + ' ' + BULAN_ID[m] + ' ' + y;
  });
  // Pastikan kalender selalu muncul walau yang diklik bagian kosong kolom
  // (bukan cuma ikon kalender kecil di ujung), baik klik tunggal maupun ganda.
  function openPicker() {
    if (typeof picker.showPicker === 'function') {
      try { picker.showPicker(); } catch (e) { /* diam-diam gagal, biarkan perilaku bawaan browser */ }
    }
  }
  picker.addEventListener('click', openPicker);
  picker.addEventListener('dblclick', openPicker);
});

// Dokumen Data Dukung: munculkan otomatis kolom Link Drive & Upload saat dicentang
document.querySelectorAll('.dukung-cb').forEach(function(cb) {
  var detail = document.getElementById(cb.dataset.target);
  if (!detail) return;
  cb.addEventListener('change', function() {
    detail.style.display = cb.checked ? 'block' : 'none';
  });
});

// Tombol "+ Tambah Dokumen Lainnya": tambah baris dokumen custom (bisa lebih dari 1)
(function() {
  var wrap = document.getElementById('dukungLainnyaWrap');
  var btn = document.getElementById('btnTambahDukung');
  if (!wrap || !btn) return;
  var counter = 0;
  btn.addEventListener('click', function() {
    counter++;
    var row = document.createElement('div');
    row.className = 'dukung-custom-row';
    row.innerHTML =
      '<input type="text" name="dukung_custom_nama[]" placeholder="Nama dokumen (ketik manual)">' +
      '<input type="text" name="dukung_custom_drive[]" placeholder="Link Google Drive (opsional)">' +
      '<button type="button" class="dukung-remove">Hapus</button>';
    row.querySelector('.dukung-remove').addEventListener('click', function() { row.remove(); });
    wrap.appendChild(row);
  });
})();

// Kunci/buka otomatis field turunan (spesies/persentase/kondisi) berdasarkan
// pilihan "Keberadaan Ekosistem" -- terkunci sampai user pilih "Terdapat ekosistem..."
document.querySelectorAll('[data-eco-trigger]').forEach(function(sel) {
  var childWrap = document.querySelector('.eco-children[data-for="' + sel.id + '"]');
  if (!childWrap) return;
  var note = document.createElement('div');
  note.className = 'eco-lock-note';
  note.textContent = 'Pilih "Terdapat ekosistem..." di atas untuk mengisi bagian ini.';
  childWrap.insertBefore(note, childWrap.firstChild);

  function update() {
    var unlocked = sel.value.indexOf('Terdapat ekosistem') === 0;
    childWrap.classList.toggle('locked', !unlocked);
    note.style.display = unlocked ? 'none' : 'block';
    childWrap.querySelectorAll('input, select').forEach(function(el) { el.disabled = !unlocked; });
  }
  sel.addEventListener('change', update);
  update();
});

// Pasangan checkbox status kegiatan yang saling eksklusif (centang salah
// satu otomatis melepas centang pasangannya, supaya tidak kontradiktif).
[["non_reklamasi", "reklamasi"], ["kegiatan_berusaha", "non_berusaha"]].forEach(function(pair) {
  var a = document.querySelector('input[name="' + pair[0] + '"]');
  var b = document.querySelector('input[name="' + pair[1] + '"]');
  if (!a || !b) return;
  a.addEventListener('change', function() { if (a.checked) b.checked = false; });
  b.addEventListener('change', function() { if (b.checked) a.checked = false; });
});

// Dropdown dengan opsi "Lainnya..." -> munculkan input teks tambahan
document.querySelectorAll('select[data-allow-other="true"]').forEach(function(sel) {
  var otherInput = document.getElementById(sel.id + '_other');
  sel.addEventListener('change', function() {
    if (sel.value === '__other__') { otherInput.classList.add('show'); otherInput.focus(); }
    else { otherInput.classList.remove('show'); otherInput.value = ''; }
  });
});

// Saat submit: kalau dropdown "Lainnya..." dipilih, kirim teks isian bebas-nya, bukan "__other__"
document.getElementById('manualForm').addEventListener('submit', function() {
  document.querySelectorAll('select[data-allow-other="true"]').forEach(function(sel) {
    if (sel.value === '__other__') {
      var otherInput = document.getElementById(sel.id + '_other');
      var hidden = document.createElement('input');
      hidden.type = 'hidden';
      hidden.name = sel.name;
      hidden.value = otherInput.value;
      sel.removeAttribute('name');
      sel.parentNode.appendChild(hidden);
    }
  });
});

// Datalist bertingkat Provinsi -> Kabupaten -> Kecamatan -> Desa/Kelurahan
// menggunakan API publik https://kodewilayah.web.id (punya demo interaktif serupa di situsnya sendiri)
(function() {
  var WILAYAH_API = '/api/wilayah';
  var provSel = document.getElementById('prop_loc__3');
  var kabInput = document.getElementById('prop_loc__2');
  var kecInput = document.getElementById('prop_loc__1');
  var desaInput = document.getElementById('prop_loc__0');
  var dlKabupaten = document.getElementById('dl_kabupaten');
  var dlKecamatan = document.getElementById('dl_kecamatan');
  var dlDesa = document.getElementById('dl_desa');
  if (!provSel || !kabInput || !kecInput || !desaInput) return;

  var cacheRegencies = [];
  var cacheDistricts = [];

  function titleCase(s) {
    return s.toLowerCase().replace(/\b\w/g, function(c) { return c.toUpperCase(); });
  }

  function fillDatalist(dl, items) {
    dl.innerHTML = '';
    items.forEach(function(txt) {
      var opt = document.createElement('option');
      opt.value = txt;
      dl.appendChild(opt);
    });
  }

  function clearFrom(level) {
    if (level <= 2) { cacheRegencies = []; fillDatalist(dlKabupaten, []); }
    if (level <= 1) { cacheDistricts = []; fillDatalist(dlKecamatan, []); }
    if (level <= 0) { fillDatalist(dlDesa, []); }
  }

  provSel.addEventListener('change', function() {
    clearFrom(2);
    kabInput.value = ''; kecInput.value = ''; desaInput.value = '';
    var opt = provSel.selectedOptions[0];
    var bpsId = opt ? opt.dataset.bpsId : '';
    if (!bpsId) return;
    fetch(WILAYAH_API + '/regencies/' + bpsId)
      .then(function(r) { return r.json(); })
      .then(function(res) {
        if (!res.success) return;
        cacheRegencies = res.data.map(function(d) { return { code: d.code, name: titleCase(d.name) }; });
        fillDatalist(dlKabupaten, cacheRegencies.map(function(d) { return d.name; }));
      })
      .catch(function() { /* diam-diam gagal; pengguna tetap bisa isi manual */ });
  });

  kabInput.addEventListener('change', function() {
    clearFrom(1);
    kecInput.value = ''; desaInput.value = '';
    var match = cacheRegencies.find(function(d) { return d.name.toLowerCase() === kabInput.value.trim().toLowerCase(); });
    if (!match) return;
    fetch(WILAYAH_API + '/districts/' + match.code)
      .then(function(r) { return r.json(); })
      .then(function(res) {
        if (!res.success) return;
        cacheDistricts = res.data.map(function(d) { return { code: d.code, name: titleCase(d.name) }; });
        fillDatalist(dlKecamatan, cacheDistricts.map(function(d) { return d.name; }));
      })
      .catch(function() {});
  });

  kecInput.addEventListener('change', function() {
    clearFrom(0);
    desaInput.value = '';
    var match = cacheDistricts.find(function(d) { return d.name.toLowerCase() === kecInput.value.trim().toLowerCase(); });
    if (!match) return;
    fetch(WILAYAH_API + '/villages/' + match.code)
      .then(function(r) { return r.json(); })
      .then(function(res) {
        if (!res.success) return;
        fillDatalist(dlDesa, res.data.map(function(d) { return titleCase(d.name); }));
      })
      .catch(function() {});
  });
})();


// Field gambar: area Paste (khusus Ctrl+V + drag-drop, TIDAK membuka file browser saat diklik)
// terpisah dari tombol Upload File (eksplisit membuka file browser). Kedua cara SELALU menambah
// (accumulate), tidak pernah saling mengganti file yang sudah ada di field yang sama.
document.querySelectorAll('.img-paste-target').forEach(function(target) {
  var fieldName = target.dataset.field;
  var input = document.getElementById(fieldName);
  var previewList = document.getElementById(fieldName + '_preview');
  var uploadBtn = document.querySelector('.img-upload-btn[data-field="' + fieldName + '"]');
  var files = [];

  function refreshInput() {
    var dt = new DataTransfer();
    files.forEach(function(f) { dt.items.add(f); });
    input.files = dt.files;
  }

  function renderPreviews() {
    previewList.innerHTML = '';
    files.forEach(function(file, idx) {
      var thumb = document.createElement('div');
      thumb.className = 'img-thumb';
      var img = document.createElement('img');
      img.src = URL.createObjectURL(file);
      var removeBtn = document.createElement('button');
      removeBtn.type = 'button';
      removeBtn.className = 'img-remove';
      removeBtn.innerHTML = '&times;';
      removeBtn.addEventListener('click', function(e) {
        e.stopPropagation();
        files.splice(idx, 1);
        refreshInput();
        renderPreviews();
      });
      thumb.appendChild(img);
      thumb.appendChild(removeBtn);
      previewList.appendChild(thumb);
    });
  }

  function addFile(file) {
    if (!file) return;
    var allowed = ['image/png', 'image/jpeg'];
    if (allowed.indexOf(file.type) === -1) {
      alert('Format file tidak didukung. Hanya PNG dan JPG/JPEG yang diperbolehkan.');
      return;
    }
    files.push(file);
    refreshInput();
    renderPreviews();
  }

  // Area Paste: HANYA untuk fokus + Ctrl+V + drag-drop. Klik TIDAK membuka file browser.
  target.addEventListener('paste', function(e) {
    var items = (e.clipboardData || window.clipboardData).items;
    for (var i = 0; i < items.length; i++) {
      if (items[i].type.indexOf('image/') === 0) { addFile(items[i].getAsFile()); }
    }
  });
  target.addEventListener('dragover', function(e) { e.preventDefault(); target.classList.add('dragover'); });
  target.addEventListener('dragleave', function() { target.classList.remove('dragover'); });
  target.addEventListener('drop', function(e) {
    e.preventDefault(); target.classList.remove('dragover');
    for (var i = 0; i < e.dataTransfer.files.length; i++) { addFile(e.dataTransfer.files[i]); }
  });

  // Tombol Upload File: eksplisit membuka file browser
  uploadBtn.addEventListener('click', function() { input.click(); });
  input.addEventListener('change', function() {
    for (var i = 0; i < input.files.length; i++) { addFile(input.files[i]); }
  });
});
</script>
""" + (f"""
<script>
// Isi ulang form otomatis dari draft riwayat pengisian sebelumnya ("Lanjutkan")
(function() {{
  var data = {json.dumps(prefill_data)};

  function setField(name, val) {{
    if (typeof val !== 'string') return;
    var els = document.getElementsByName(name);
    if (!els.length) {{
      var byId = document.getElementById(name);
      if (byId) els = [byId];
    }}
    Array.prototype.forEach.call(els, function(el) {{
      if (el.type === 'checkbox') {{ el.checked = (val === true || val === 'true'); }}
      else if (el.tagName === 'SELECT') {{
        el.value = val;
        if (el.value !== val) {{
          // Nilai tersimpan tidak cocok opsi manapun di dropdown -- ini
          // hasil isian "Lainnya..." sebelumnya (custom text). Kalau
          // dropdown ini mendukungnya, pindah ke opsi "Lainnya..." dan isi
          // teks bebasnya, supaya tidak diam-diam kosong/hilang.
          if (el.dataset.allowOther === 'true') {{
            el.value = '__other__';
            var otherInput = document.getElementById(el.id + '_other');
            if (otherInput) {{
              otherInput.value = val;
              otherInput.classList.add('show');
            }}
          }}
        }}
        el.dispatchEvent(new Event('change'));
        el.dispatchEvent(new Event('input'));
        // Field tanggal: input asli yang terlihat adalah <input type="date"
        // id="{{name}}_picker">, sedangkan {{name}} sendiri cuma hidden input
        // penyimpan nilai -- keduanya perlu disamakan supaya terlihat terisi.
        var picker = document.getElementById(el.id + '_picker');
        if (picker) picker.value = val;
      }}
      else if (el.tagName === 'TEXTAREA' || el.tagName === 'INPUT') {{
        el.value = val;
        el.dispatchEvent(new Event('change'));
        el.dispatchEvent(new Event('input'));
        var picker2 = document.getElementById(el.id + '_picker');
        if (picker2) picker2.value = val;
      }}
    }});
  }}

  // "_lokasi_parts" (Desa/Kecamatan/Kabupaten/Provinsi) diproses LEBIH DULU,
  // sebelum aturan umum "key.startsWith('_') -> lewati" di bawah -- kalau
  // tidak, field ini akan ikut terlewati juga karena namanya diawali "_".
  //
  // PENTING soal urutan: skrip datalist bertingkat (Provinsi -> Kabupaten ->
  // Kecamatan -> Desa) otomatis MENGOSONGKAN semua level di BAWAHNYA setiap
  // kali sebuah level berubah (supaya user tidak salah pilih kombinasi usang).
  // Kalau kita isi dari Desa ke Provinsi (searah array), mengisi Provinsi
  // PALING AKHIR akan langsung menghapus Kabupaten/Kecamatan/Desa yang baru
  // saja diisi. Maka isi SEARAH CASCADE: Provinsi -> Kabupaten -> Kecamatan
  // -> Desa, supaya tiap pengisian terjadi SETELAH pengosongan level di
  // bawahnya oleh level sebelumnya, bukan sebelum.
  if (Array.isArray(data._lokasi_parts)) {{
    var lp = data._lokasi_parts; // [desa, kecamatan, kabupaten, provinsi]
    setField('prop_loc__3', lp[3]);
    setField('prop_loc__2', lp[2]);
    setField('prop_loc__1', lp[1]);
    setField('prop_loc__0', lp[0]);
  }}

  // "instalasi_posisi" disimpan sebagai satu string gabungan (mis. "Permukaan
  // Laut dan Dasar Laut"), tapi elemennya sekarang berupa beberapa checkbox
  // terpisah (name="instalasi_posisi[]") -- pecah lagi jadi per-checkbox.
  if (typeof data.instalasi_posisi === 'string' && data.instalasi_posisi) {{
    var posisiTerpilih = data.instalasi_posisi.split(/,\\s*| dan /);
    document.querySelectorAll('input[name="instalasi_posisi[]"]').forEach(function(cb) {{
      cb.checked = posisiTerpilih.indexOf(cb.value) !== -1;
    }});
  }}

  Object.keys(data).forEach(function(key) {{
    if (key.startsWith('_')) return;
    var val = data[key];
    if (key === 'koordinat' && Array.isArray(val)) {{
      var ta = document.getElementById('koordinat_manual');
      if (ta) ta.value = val.map(function(row) {{ return row[1] + '\\t' + row[2]; }}).join('\\n');
      return;
    }}
    if (typeof val === 'boolean') {{
      var cb = document.getElementsByName(key)[0];
      if (cb && cb.type === 'checkbox') cb.checked = val;
      return;
    }}
    if (typeof val !== 'string') return;
    // Field dari FIELD_GROUPS (mis. "Nama Pemohon") pakai nama HTML "prop__Nama_Pemohon"
    // (spasi/garis-miring diganti underscore) -- coba key asli dulu, baru varian ini.
    setField(key, val);
    var transformed = 'prop__' + key.replace(/[ /]/g, '_');
    if (transformed !== key) setField(transformed, val);
  }});
  document.querySelectorAll('details.acc-item').forEach(function(d) {{ d.open = true; }});
}})();
</script>
""" if prefill_data else "") + """
</body></html>"""


def render_review_page(job_id, prop_data, lap_data, preview_html, error=None):
    # Amankan dari karakter yang mirip sintaks Jinja ({{ }}, {% %}) kalau kebetulan
    # muncul di konten dokumen asli pengguna -- supaya tidak salah dievaluasi
    # sebagai template saat halaman ini dibungkus render_template_string().
    preview_html = (
        preview_html.replace("{{", "&#123;&#123;")
        .replace("}}", "&#125;&#125;")
        .replace("{%", "&#123;%")
        .replace("%}", "%&#125;")
    )
    groups_html = []
    for i, (group_title, fields) in enumerate(FIELD_GROUPS):
        rows = []
        for source, key, label in fields:
            value = get_value(source, key, prop_data, lap_data)
            fname = form_field_name(source, key)
            value_escaped = (value or "").replace('"', "&quot;")
            rows.append(
                f'<div class="field-row"><label>{label}</label>'
                f'<input type="text" name="{fname}" value="{value_escaped}"></div>'
            )
        groups_html.append(
            f'<details class="acc-item"{" open" if i == 0 else ""}>'
            f"<summary>{group_title}</summary>"
            f'<div class="acc-body">{"".join(rows)}</div></details>'
        )

    error_html = f'<div class="error-banner" style="margin-bottom:16px;">\u26A0 {error}</div>' if error else ""
    koordinat_pesan = prop_data.get("_koordinat_file_pesan", "")
    if koordinat_pesan:
        icon = "\u2705" if "Berhasil" in koordinat_pesan else "\u2139\ufe0f"
        error_html += f'<div class="error-banner" style="margin-bottom:16px;background:#eaf6ec;border-color:#b7e4c7;">{icon} {koordinat_pesan}</div>'

    return """<!DOCTYPE html>
<html lang="id"><head><meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Tinjau &amp; Koreksi Data &mdash; e-GeRAI KKPRL</title>
<style>""" + LANDING_CSS + REVIEW_CSS + """</style></head>
<body>
""" + HEADER_HTML + """

<section class="review-hero">
  <h1>\U0001F4DD Tinjau &amp; Koreksi Data</h1>
  <p>Periksa hasil ekstraksi di bawah, lalu bandingkan dengan pratinjau dokumen di sebelah kanan.
  Koreksi kolom yang salah, lalu klik "Generate Dokumen Final &amp; Unduh".</p>
</section>

<div class="review-wrap">
  """ + error_html + """
  <form method="POST" action="/finalize">
    <input type="hidden" name="job_id" value=\"""" + job_id + """\">
    <div class="review-grid">
      <div>
        <div class="sticky-bar">
          <button type="submit">""" + ICONS["check-circle"] + """ Generate Dokumen Final &amp; Unduh</button>
        </div>
        <div class="review-card">
          <h3>""" + ICONS["doc"] + """ Data Hasil Ekstraksi (bisa dikoreksi)</h3>
          """ + "".join(groups_html) + """
        </div>
      </div>

      <div class="preview-panel">
        <div class="review-card">
          <h3>""" + ICONS["doc"] + """ Pratinjau Dokumen Lengkap</h3>
          <div class="preview-box">""" + preview_html + """</div>
        </div>
      </div>
    </div>
  </form>

  <a href="/" class="back-link">&larr; Unggah ulang dokumen lain</a>
</div>
</body></html>"""


PUBLIC_PATHS = {"/login-pegawai", "/login", "/auth/callback", "/logout", "/health"}
PUBLIC_PREFIXES = ("/static/",)


@app.before_request
def require_login_pegawai():
    """Seluruh aplikasi WAJIB login dulu (Kode Nama + Kode Petugas) sebelum
    bisa dipakai -- kecuali halaman login itu sendiri, logout, file statis,
    dan health check (dipakai Railway untuk cek server hidup/tidak)."""
    if request.path in PUBLIC_PATHS or request.path.startswith(PUBLIC_PREFIXES):
        return None
    if not session.get("user"):
        return render_login_pegawai_page()
    return None


@app.route("/", methods=["GET"])
def index():
    return render_template_string(UPLOAD_HTML, error=None)


@app.route("/health", methods=["GET"])
def health():
    return {"status": "ok"}


@app.route("/api/wilayah/<level>/<code>")
def api_wilayah_proxy(level, code):
    """Proxy server-side ke api.kodewilayah.web.id supaya request datalist
    Kabupaten/Kecamatan/Desa tidak terhalang CORS di browser (fetch dilakukan
    dari server, browser cuma bicara dengan server sendiri)."""
    import urllib.request
    import json as _json

    if level not in ("regencies", "districts", "villages"):
        return {"success": False, "message": "level tidak valid", "data": []}, 400

    url = f"https://api.kodewilayah.web.id/{level}/{code}"
    try:
        req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
        with urllib.request.urlopen(req, timeout=8) as resp:
            payload = _json.loads(resp.read().decode("utf-8"))
        return payload
    except Exception as e:
        traceback.print_exc()
        return {"success": False, "message": str(e), "data": []}, 502


@app.route("/login-pegawai", methods=["GET"])
def login_pegawai_form():
    return render_login_pegawai_page()


@app.route("/login-pegawai", methods=["POST"])
def login_pegawai_submit():
    kode_nama = (request.form.get("kode_nama") or "").strip()
    kode_petugas = (request.form.get("kode_petugas") or "").strip()
    staff = fetch_staff_list()
    entry = staff.get(kode_nama)
    if entry and entry["password"] == kode_petugas:
        session["user"] = {"name": entry["nama"], "email": "", "picture": "", "kode": kode_nama}
        session.permanent = True
        return redirect(url_for("index"))
    return render_login_pegawai_page(
        error="Kode Nama atau Kode Petugas salah. Mohon periksa kembali."
    )


@app.route("/login")
def login():
    try:
        redirect_uri = url_for("auth_callback", _external=True)
        return google_oauth.authorize_redirect(redirect_uri)
    except Exception:
        traceback.print_exc()
        return render_template_string(UPLOAD_HTML, error=(
            "Login Google belum bisa diproses. Kemungkinan GOOGLE_CLIENT_ID / "
            "GOOGLE_CLIENT_SECRET belum diisi di environment variable server."
        ))


@app.route("/auth/callback")
def auth_callback():
    try:
        token = google_oauth.authorize_access_token()
        user_info = token.get("userinfo") or {}
        session["user"] = {
            "name": user_info.get("name", "Pengguna"),
            "email": user_info.get("email", ""),
            "picture": user_info.get("picture", ""),
        }
    except Exception:
        traceback.print_exc()
    return redirect(url_for("index"))


@app.route("/logout")
def logout():
    session.pop("user", None)
    return redirect(url_for("index"))


@app.route("/history")
def history_page():
    return render_template_string(render_history_page())


@app.route("/riwayat-saya")
def riwayat_saya_page():
    return render_template_string(render_riwayat_saya_page())


@app.route("/riwayat-saya/lanjutkan/<job_id>")
def riwayat_saya_lanjutkan(job_id):
    user = session.get("user")
    if not user or not user.get("kode"):
        return render_template_string(render_login_pegawai_page())
    draft = load_draft(user["kode"], job_id)
    if not draft:
        return render_template_string(render_riwayat_saya_page())
    return render_template_string(render_manual_form_page(
        prefill_data=draft.get("prop_data", {}),
        job_id=job_id,
        saved_images_meta=draft.get("prop_images_meta", []),
    ))


@app.route("/review", methods=["POST"])
def review():
    proposal_file = request.files.get("proposal")
    laporan_file = request.files.get("laporan")

    if not proposal_file or not laporan_file or proposal_file.filename == "" or laporan_file.filename == "":
        return render_template_string(UPLOAD_HTML, error="Mohon unggah kedua file (proposal & laporan)."), 400
    if not proposal_file.filename.lower().endswith((".pdf", ".docx")):
        return render_template_string(UPLOAD_HTML, error="File Draft Proposal harus berformat PDF atau Word (.docx)."), 400
    proposal_ext = ".docx" if proposal_file.filename.lower().endswith(".docx") else ".pdf"
    laporan_ext = ".docx" if laporan_file.filename.lower().endswith(".docx") else ".pdf"
    if not laporan_file.filename.lower().endswith((".pdf", ".docx")):
        return render_template_string(UPLOAD_HTML, error="File Laporan harus berformat PDF atau Word (.docx)."), 400

    job_store.cleanup_old_jobs(JOBS_DIR)

    job_id = uuid.uuid4().hex[:12]
    tmp_dir = os.path.join(UPLOAD_DIR, job_id)
    os.makedirs(tmp_dir, exist_ok=True)
    proposal_path = os.path.join(tmp_dir, "proposal" + proposal_ext)
    laporan_path = os.path.join(tmp_dir, "laporan" + laporan_ext)
    proposal_file.save(proposal_path)
    laporan_file.save(laporan_path)

    try:
        prop_data, prop_images = extract_proposal_with_fallback(proposal_path, log=lambda *_: None)
        lap_data, lap_images = extract_laporan_with_fallback(laporan_path, log=lambda *_: None)

        job_store.save_job(JOBS_DIR, job_id, prop_data, prop_images, lap_data, lap_images)

        preview_docx_path = os.path.join(tmp_dir, "preview.docx")
        build_document(prop_data, prop_images, lap_data, lap_images, preview_docx_path)
        with open(preview_docx_path, "rb") as f:
            preview_html = mammoth.convert_to_html(f).value
    except Exception:
        traceback.print_exc()
        shutil.rmtree(tmp_dir, ignore_errors=True)
        job_store.delete_job(JOBS_DIR, job_id)
        return render_template_string(
            UPLOAD_HTML,
            error="Terjadi kesalahan saat memproses dokumen. Pastikan file Proposal dan Laporan adalah PDF/Word yang valid.",
        ), 500
    finally:
        shutil.rmtree(tmp_dir, ignore_errors=True)

    return render_template_string(render_review_page(job_id, prop_data, lap_data, preview_html))


@app.route("/proposal-manual", methods=["GET"])
def proposal_manual_form():
    return render_template_string(render_manual_form_page())


def build_prop_data_from_manual_form(form, files):
    """Bangun prop_data & prop_images dari form isian manual (dipakai baik
    untuk alur submit-ke-review biasa maupun unduh-draft-langsung)."""
    def file_ext(filename):
        ext = os.path.splitext(filename)[1].lstrip(".").lower()
        return "jpg" if ext == "jpeg" else ext

    ALLOWED_IMAGE_EXT = ("jpg", "jpeg", "png")
    prop_images = []

    prop_data = {}
    prop_data, _ = apply_form_values(form, prop_data, {})
    prop_data["non_reklamasi"] = "non_reklamasi" in form
    prop_data["reklamasi"] = "reklamasi" in form
    prop_data["kegiatan_berusaha"] = "kegiatan_berusaha" in form
    prop_data["non_berusaha"] = "non_berusaha" in form
    prop_data["non_strategis"] = "non_strategis" in form
    prop_data["kegiatan_status"] = form.get("kegiatan_status", "")
    prop_data["sumber_peta"] = form.get("sumber_peta", "")
    prop_data["deskripsi_kegiatan"] = form.get("deskripsi_kegiatan", "")
    prop_data["manfaat_kegiatan"] = form.get("manfaat_kegiatan", "")
    prop_data["tujuan_kegiatan"] = form.get("tujuan_kegiatan", "")
    prop_data["instalasi_bangunan"] = form.get("instalasi_bangunan", "")
    prop_data["instalasi_posisi"] = join_dan(form.getlist("instalasi_posisi[]"))
    prop_data["jadwal_kegiatan"] = form.get("jadwal_kegiatan", "")
    DUKUNG_ITEMS = [
        ("dukung_nib", "NIB"),
        ("dukung_sertifikat", "Sertifikat Kepemilikan Lahan Darat"),
        ("dukung_izin_lingkungan", "Surat Izin Lingkungan"),
        ("dukung_ba_sosialisasi", "Berita Acara Sosialisasi"),
        ("dukung_identitas", "Dokumen Identitas dan Legalitas Pemohon/Perusahaan"),
        ("dukung_survei", "Dokumentasi Survei Lapangan Kondisi Eksisting Lokasi"),
        ("dukung_peta", "Peta Pendukung (Peta Lokasi, Site Plan, Pola Ruang Wilayah)"),
        ("dukung_dipa", "DIPA/RKAKL (Sumber Anggaran APBD/APBN) / Lainnya"),
        ("dukung_sk_kkprl", "SK Penetapan KNMP"),
    ]
    dukung_list = []
    dukung_detail = []
    for field_name, label in DUKUNG_ITEMS:
        if field_name in form:
            entry = {"label": label, "drive": form.get(f"{field_name}_drive", "").strip(), "file": ""}
            f = files.get(f"{field_name}_file")
            if f and f.filename:
                entry["file"] = f.filename
                ext = os.path.splitext(f.filename)[1].lstrip(".").lower()
                if ext in ("png", "jpg", "jpeg"):
                    prop_images.append({"tag": "dukung_dokumen", "bytes": f.read(), "ext": file_ext(f.filename)})
            dukung_list.append(label)
            dukung_detail.append(entry)
    for nama_custom, drive_custom in zip(form.getlist("dukung_custom_nama[]"), form.getlist("dukung_custom_drive[]")):
        nama_custom = nama_custom.strip()
        if nama_custom:
            dukung_list.append(nama_custom)
            dukung_detail.append({"label": nama_custom, "drive": drive_custom.strip(), "file": ""})
    prop_data["dokumen_data_dukung"] = ", ".join(dukung_list)
    prop_data["dokumen_data_dukung_detail"] = dukung_detail

    prop_data["mangrove_ada"] = form.get("mangrove_ada", "")
    prop_data["mangrove_spesies"] = form.get("prop__mangrove_spesies", "")
    prop_data["mangrove_persen"] = form.get("prop__mangrove_persen", "")
    prop_data["mangrove_kondisi"] = form.get("prop__mangrove_kondisi", "")
    prop_data["lamun_ada_manual"] = form.get("lamun_ada_manual", "")
    prop_data["lamun_spesies"] = form.get("lamun_spesies", "")
    prop_data["lamun_persen"] = form.get("lamun_persen", "")
    prop_data["lamun_kondisi"] = form.get("lamun_kondisi", "")
    prop_data["karang_ada"] = form.get("karang_ada", "")
    prop_data["karang_spesies"] = form.get("karang_spesies", "")
    prop_data["karang_persen_manual"] = form.get("karang_persen_manual", "")
    prop_data["karang_kondisi"] = form.get("karang_kondisi", "")

    prop_data["batas_utara"] = form.get("batas_utara", "")
    prop_data["batas_timur"] = form.get("batas_timur", "")
    prop_data["batas_selatan"] = form.get("batas_selatan", "")
    prop_data["batas_barat"] = form.get("batas_barat", "")
    prop_data["deskripsi_pemanfaatan_sekitar"] = form.get("deskripsi_pemanfaatan_sekitar", "")
    prop_data["mata_pencaharian"] = form.get("mata_pencaharian", "")
    prop_data["sumber_data_sosek"] = form.get("sumber_data_sosek", "")
    prop_data["tahun_data_sosek"] = form.get("tahun_data_sosek", "")
    prop_data["aksesibilitas_lokasi"] = form.get("aksesibilitas_lokasi", "")

    koordinat = []
    koordinat_file_pesan = None
    koordinat_file = files.get("upload_koordinat")
    if koordinat_file and koordinat_file.filename:
        hasil_file, koordinat_file_pesan = parse_koordinat_file(koordinat_file)
        koordinat.extend(hasil_file)
    for line in form.get("koordinat_manual", "").splitlines():
        line = line.strip()
        if not line:
            continue
        if "|" in line:
            parts = [p.strip() for p in line.split("|")]
            if len(parts) == 3 and any(parts):
                koordinat.append([str(len(koordinat) + 1), parts[1], parts[2]])
        else:
            parts = line.split()
            if len(parts) == 2:
                koordinat.append([str(len(koordinat) + 1), parts[0], parts[1]])
    prop_data["koordinat"] = koordinat
    if koordinat_file_pesan:
        prop_data["_koordinat_file_pesan"] = koordinat_file_pesan

    for field_name, tag in IMAGE_FIELD_TAGS:
        for f in files.getlist(field_name):
            if f and f.filename and os.path.splitext(f.filename)[1].lstrip(".").lower() in ALLOWED_IMAGE_EXT:
                prop_images.append({"tag": tag, "bytes": f.read(), "ext": file_ext(f.filename)})

    return prop_data, prop_images


EMPTY_LAP_DATA = {}  # dipakai saat generate draft tanpa Laporan Hidro-Oseanografi


@app.route("/proposal-manual/draft-image/<job_id>/<path:filename>")
def proposal_manual_draft_image(job_id, filename):
    """Sajikan kembali satu file gambar yang sudah tersimpan lewat tombol
    \"Simpan\", supaya bisa ditampilkan sebagai thumbnail preview di form.
    Hanya bisa diakses oleh pemilik draft yang sedang login (dicocokkan
    lewat kode pegawai di session)."""
    user = session.get("user")
    if not user or not user.get("kode"):
        return "", 403
    safe_kode = _safe_kode(user["kode"])
    # Cegah path traversal -- filename harus persis nama file, tanpa "/" atau "..".
    filename = os.path.basename(filename)
    img_dir = _draft_images_dir(safe_kode, job_id)
    fpath = os.path.join(img_dir, filename)
    if not os.path.isfile(fpath) or not os.path.abspath(fpath).startswith(os.path.abspath(img_dir)):
        return "", 404
    return send_file(fpath)


@app.route("/proposal-manual/simpan", methods=["POST"])
def proposal_manual_simpan():
    """Simpan isian form (SEMUA field + gambar yang sudah diunggah) sebagai
    draft, TANPA memproses/menggabungkan dokumen apa pun -- supaya pengguna
    bisa berhenti sejenak dan melanjutkan isian nanti tanpa kehilangan apa
    yang sudah dikerjakan. Beda dari \"Unduh Draft\" yang langsung membuat
    file Word; tombol ini murni menyimpan progres isian."""
    user = session.get("user")
    if not user or not user.get("kode"):
        return render_template_string(render_login_pegawai_page())

    try:
        prop_data, prop_images = build_prop_data_from_manual_form(request.form, request.files)
        existing_job_id = request.form.get("existing_job_id", "").strip()
        job_id = existing_job_id or uuid.uuid4().hex[:12]
        prop_images = merge_draft_images(user["kode"], job_id, prop_images)
        save_draft(user["kode"], job_id, prop_data, prop_images)
    except Exception:
        traceback.print_exc()
        return render_template_string(render_manual_form_page(
            error="Terjadi kesalahan saat menyimpan draft. Silakan coba lagi.",
        )), 500

    n_img = len(prop_images)
    pesan = f"Draft tersimpan ({n_img} gambar ikut tersimpan)." if n_img else "Draft tersimpan."
    # Baca ulang metadata gambar dari draft yang baru saja disimpan supaya
    # thumbnail-nya langsung tampil di halaman ini juga (bukan cuma
    # setelah reload lewat "Lanjutkan").
    saved_draft = load_draft(user["kode"], job_id)
    saved_images_meta = (saved_draft or {}).get("prop_images_meta", [])
    return render_template_string(render_manual_form_page(
        prefill_data=prop_data,
        job_id=job_id,
        saved_msg=pesan,
        saved_images_meta=saved_images_meta,
    ))


@app.route("/proposal-manual", methods=["POST"])
def proposal_manual_submit():
    laporan_file = request.files.get("laporan")
    laporan_ext = None
    if laporan_file and laporan_file.filename:
        if not laporan_file.filename.lower().endswith((".pdf", ".docx")):
            return render_template_string(render_manual_form_page(error="File Laporan harus berformat PDF atau Word (.docx).")), 400
        laporan_ext = ".docx" if laporan_file.filename.lower().endswith(".docx") else ".pdf"

    job_store.cleanup_old_jobs(JOBS_DIR)

    job_id = uuid.uuid4().hex[:12]
    tmp_dir = os.path.join(UPLOAD_DIR, job_id)
    os.makedirs(tmp_dir, exist_ok=True)
    laporan_path = None
    if laporan_ext:
        laporan_path = os.path.join(tmp_dir, "laporan" + laporan_ext)
        laporan_file.save(laporan_path)

    try:
        prop_data, prop_images = build_prop_data_from_manual_form(request.form, request.files)
        current_user = session.get("user")
        draft_job_id = request.form.get("existing_job_id", "").strip() or job_id
        if current_user and current_user.get("kode"):
            prop_images = merge_draft_images(current_user["kode"], draft_job_id, prop_images)

        if laporan_path:
            lap_data, lap_images = extract_laporan_with_fallback(laporan_path, log=lambda *_: None)
        else:
            lap_data, lap_images = dict(EMPTY_LAP_DATA), []

        job_store.save_job(JOBS_DIR, job_id, prop_data, prop_images, lap_data, lap_images)
        if current_user and current_user.get("kode"):
            save_draft(current_user["kode"], draft_job_id, prop_data, prop_images)

        preview_docx_path = os.path.join(tmp_dir, "preview.docx")
        build_document(prop_data, prop_images, lap_data, lap_images, preview_docx_path)
        with open(preview_docx_path, "rb") as f:
            preview_html = mammoth.convert_to_html(f).value
    except Exception:
        traceback.print_exc()
        shutil.rmtree(tmp_dir, ignore_errors=True)
        job_store.delete_job(JOBS_DIR, job_id)
        return render_template_string(render_manual_form_page(error="Terjadi kesalahan saat memproses. Pastikan file Laporan adalah PDF/Word yang valid.")), 500
    finally:
        shutil.rmtree(tmp_dir, ignore_errors=True)

    return render_template_string(render_review_page(job_id, prop_data, lap_data, preview_html))


@app.route("/proposal-manual/draft", methods=["POST"])
def proposal_manual_draft():
    """Unduh draft Proposal langsung (TANPA menunggu/menggabung Laporan
    Hidro-Oseanografi) -- untuk staf yang ingin cepat mengunduh draft
    berisi data yang sudah diisi sejauh ini, sebelum melengkapi Laporan
    di kesempatan berikutnya."""
    try:
        prop_data, prop_images = build_prop_data_from_manual_form(request.form, request.files)
        lap_data, lap_images = dict(EMPTY_LAP_DATA), []

        current_user = session.get("user")
        draft_job_id = request.form.get("existing_job_id", "").strip() or uuid.uuid4().hex[:12]
        if current_user and current_user.get("kode"):
            prop_images = merge_draft_images(current_user["kode"], draft_job_id, prop_images)
            save_draft(current_user["kode"], draft_job_id, prop_data, prop_images)

        tmp_path = os.path.join(OUTPUT_DIR, f"Draft_{uuid.uuid4().hex[:12]}.docx")
        build_document(prop_data, prop_images, lap_data, lap_images, tmp_path)
        log_history_entry(
            nama_pemohon=prop_data.get("Nama Pemohon", ""),
            nama_perusahaan=prop_data.get("Nama Perusahaan/Instansi", ""),
            sumber="Unduh Draft (Tanpa Laporan)",
        )
    except Exception:
        traceback.print_exc()
        return render_template_string(render_manual_form_page(
            error="Terjadi kesalahan saat membuat draft. Silakan coba lagi."
        )), 500

    @after_this_request
    def cleanup(response):
        try:
            os.remove(tmp_path)
        except OSError:
            pass
        return response

    perusahaan = prop_data.get("Nama Perusahaan/Instansi", "PKKPRL").replace(" ", "_").replace(".", "")
    download_name = f"Draft_Proposal_PKKPRL_{perusahaan}.docx"
    return send_file(
        tmp_path,
        as_attachment=True,
        download_name=download_name,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@app.route("/finalize", methods=["POST"])
def finalize():
    job_id = request.form.get("job_id", "")
    loaded = job_store.load_job(JOBS_DIR, job_id)
    if not loaded:
        return render_template_string(
            UPLOAD_HTML,
            error="Sesi review sudah kedaluwarsa atau tidak ditemukan. Mohon unggah ulang dokumennya.",
        ), 400

    prop_data, prop_images, lap_data, lap_images = loaded
    prop_data, lap_data = apply_form_values(request.form, prop_data, lap_data)

    output_path = os.path.join(OUTPUT_DIR, f"Proposal_Final_{job_id}.docx")
    try:
        build_document(prop_data, prop_images, lap_data, lap_images, output_path)
        log_history_entry(
            nama_pemohon=prop_data.get("Nama Pemohon", ""),
            nama_perusahaan=prop_data.get("Nama Perusahaan/Instansi", ""),
            sumber="Generate Dokumen Final",
        )
    except Exception:
        traceback.print_exc()
        return render_template_string(
            UPLOAD_HTML, error="Terjadi kesalahan saat membuat dokumen final. Silakan coba lagi."
        ), 500
    finally:
        job_store.delete_job(JOBS_DIR, job_id)

    @after_this_request
    def cleanup(response):
        try:
            os.remove(output_path)
        except OSError:
            pass
        return response

    perusahaan = prop_data.get("Nama Perusahaan/Instansi", "PKKPRL").replace(" ", "_").replace(".", "")
    download_name = f"Proposal_Teknis_PKKPRL_{perusahaan}.docx"

    return send_file(
        output_path,
        as_attachment=True,
        download_name=download_name,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port, threaded=True, debug=False)
