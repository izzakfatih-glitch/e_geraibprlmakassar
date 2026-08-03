"""
Modul pembangun dokumen Word (proposal PKKPRL final) dari data & gambar
hasil ekstraksi (lihat extract.py).
"""
import io
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1F, 0x4E, 0x79)
LIGHTBLUE = "DCE6F1"
FONT = "Calibri"

NA = "[data tidak terdeteksi otomatis \u2013 mohon lengkapi manual]"


def g(d, key, default=NA):
    v = d.get(key)
    if v is None or v == "":
        return default
    return v


def shade_cell(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def set_font(run, size=11, bold=False, italic=False, color=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), FONT)


BULAN_ID_KE_NOMOR = {
    "januari": 1, "februari": 2, "maret": 3, "april": 4, "mei": 5, "juni": 6,
    "juli": 7, "agustus": 8, "september": 9, "oktober": 10, "november": 11, "desember": 12,
}
BULAN_ID_SINGKAT = ["Jan", "Feb", "Mar", "Apr", "Mei", "Jun", "Jul", "Ags", "Sep", "Okt", "Nov", "Des"]


def parse_jadwal_kegiatan(text):
    """Uraikan teks 'Deskripsi Jadwal Kegiatan' menjadi daftar aktivitas.
    Format yang diharapkan (satu atau lebih, dipisah titik):
    "Nama Kegiatan : Bulan X - Bulan Y." Mengembalikan list of
    (nama_kegiatan, bulan_mulai, bulan_selesai). Kalau tidak ada yang cocok
    sama sekali, mengembalikan list kosong (pemanggil lalu memakai fallback
    teks polos biasa)."""
    if not text:
        return []
    pattern = re.compile(
        r"([^.:]+?)\s*:\s*Bulan\s*(\d+)\s*(?:-|s/d|sampai)\s*Bulan\s*(\d+)\s*\.?",
        re.IGNORECASE,
    )
    hasil = []
    for m in pattern.finditer(text):
        nama = m.group(1).strip(" .")
        try:
            mulai = int(m.group(2))
            selesai = int(m.group(3))
        except ValueError:
            continue
        if nama and selesai >= mulai:
            hasil.append((nama, mulai, selesai))
    return hasil


def parse_tanggal_indonesia(text):
    """Uraikan tanggal format 'D Bulan YYYY' (mis. '3 Agustus 2026') jadi
    (bulan, tahun). Mengembalikan None kalau gagal."""
    if not text:
        return None
    m = re.search(r"(\d{1,2})\s+([A-Za-z]+)\s+(\d{4})", text)
    if not m:
        return None
    nama_bulan = m.group(2).strip().lower()
    bulan = BULAN_ID_KE_NOMOR.get(nama_bulan)
    if not bulan:
        return None
    return bulan, int(m.group(3))


def build_gantt_table(builder, activities, start_bulan_tahun):
    """Buat tabel Gantt chart (rencana jadwal pelaksanaan) mirip format resmi:
    baris tahun, baris bulan (Jan-Des disingkat), lalu 1 baris per kegiatan
    dengan sel yang diarsir gelap pada rentang bulan aktivitas tsb berjalan."""
    max_bulan = max(a[2] for a in activities)
    start_month, start_year = start_bulan_tahun

    # Bangun label kalender aktual untuk tiap "Bulan N" (N=1..max_bulan)
    kalender = []
    m, y = start_month, start_year
    for _ in range(max_bulan):
        kalender.append((BULAN_ID_SINGKAT[m - 1], y))
        m += 1
        if m > 12:
            m = 1
            y += 1

    doc = builder.doc
    ncols = 1 + max_bulan
    table = doc.add_table(rows=0, cols=ncols)
    table.style = "Table Grid"

    # Baris 1: label tahun (digabung/merge per kelompok tahun yang sama)
    row_year = table.add_row()
    shade_cell(row_year.cells[0], "1F4E79")
    row_year.cells[0].text = ""
    i = 0
    while i < max_bulan:
        y = kalender[i][1]
        j = i
        while j < max_bulan and kalender[j][1] == y:
            j += 1
        start_cell = row_year.cells[1 + i]
        if j - i > 1:
            end_cell = row_year.cells[1 + j - 1]
            start_cell = start_cell.merge(end_cell)
        shade_cell(start_cell, "1F4E79")
        p = start_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(y))
        set_font(r, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        i = j

    # Baris 2: header "Kegiatan" + nama bulan singkat
    row_bulan = table.add_row()
    shade_cell(row_bulan.cells[0], "1F4E79")
    p0 = row_bulan.cells[0].paragraphs[0]
    r0 = p0.add_run("Kegiatan")
    set_font(r0, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    for i, (nama_bulan, _) in enumerate(kalender):
        cell = row_bulan.cells[1 + i]
        shade_cell(cell, "1F4E79")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(nama_bulan)
        set_font(r, size=9.5, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))

    # Baris kegiatan
    for nama, mulai, selesai in activities:
        row = table.add_row()
        cell0 = row.cells[0]
        p0 = cell0.paragraphs[0]
        r0 = p0.add_run(nama)
        set_font(r0, size=9.5, bold=True)
        for i in range(max_bulan):
            bulan_ke = i + 1
            cell = row.cells[1 + i]
            cell.text = ""
            if mulai <= bulan_ke <= selesai:
                shade_cell(cell, "1F4E79")

    # Lebar kolom: kolom "Kegiatan" lebih lebar, kolom bulan sempit & seragam
    table.columns[0].width = Cm(3.6)
    for i in range(max_bulan):
        table.columns[1 + i].width = Cm(1.0)
    for row in table.rows:
        row.cells[0].width = Cm(3.6)
        for i in range(max_bulan):
            row.cells[1 + i].width = Cm(1.0)

    return table


class Builder:
    def __init__(self):
        self.doc = Document()
        style = self.doc.styles["Normal"]
        style.font.name = FONT
        style.font.size = Pt(11)
        section = self.doc.sections[0]
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.0)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)

    def h1(self, text):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(text)
        set_font(r, size=15, bold=True, color=NAVY)
        return p

    def h2(self, text):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(text)
        set_font(r, size=13, bold=True, color=NAVY)
        return p

    def h3(self, text):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(text)
        set_font(r, size=12, bold=True)
        return p

    def p(self, text, italic=False, justify=True):
        para = self.doc.add_paragraph()
        para.paragraph_format.space_after = Pt(8)
        para.paragraph_format.line_spacing = 1.25
        if justify:
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = para.add_run(text)
        set_font(r, size=11, italic=italic)
        return para

    def labeled(self, label, text):
        para = self.doc.add_paragraph()
        para.paragraph_format.space_after = Pt(8)
        para.paragraph_format.line_spacing = 1.25
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r1 = para.add_run(label + ": ")
        set_font(r1, size=11, bold=True)
        r2 = para.add_run(text)
        set_font(r2, size=11)
        return para

    def bullet(self, text):
        para = self.doc.add_paragraph(style="List Bullet")
        r = para.add_run(text)
        set_font(r, size=11)
        return para

    def caption(self, text):
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_after = Pt(12)
        r = para.add_run(text)
        set_font(r, size=10, italic=True)
        return para

    def image(self, img_bytes, width_cm=14):
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.add_picture(io.BytesIO(img_bytes), width=Cm(width_cm))
        return para

    def image_missing(self, label):
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = para.add_run(f"[GAMBAR '{label}' TIDAK DITEMUKAN DI DOKUMEN SUMBER]")
        set_font(r, size=10, italic=True, color=RGBColor(0xAA, 0x00, 0x00))
        return para

    def page_break(self):
        self.doc.add_page_break()

    def kv_table(self, rows, col_widths_cm=(5.0, 10.0)):
        table = self.doc.add_table(rows=0, cols=2)
        table.style = "Table Grid"
        for k, v in rows:
            row = table.add_row()
            row.cells[0].width = Cm(col_widths_cm[0])
            row.cells[1].width = Cm(col_widths_cm[1])
            p0 = row.cells[0].paragraphs[0]
            r0 = p0.add_run(k)
            set_font(r0, size=10.5, bold=True)
            shade_cell(row.cells[0], LIGHTBLUE)
            row.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p1 = row.cells[1].paragraphs[0]
            r1 = p1.add_run(str(v))
            set_font(r1, size=10.5)
            row.cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        return table

    def data_table(self, header, rows, header_color=None):
        header_color = header_color or "1F4E79"
        ncols = len(header)
        table = self.doc.add_table(rows=0, cols=ncols)
        table.style = "Table Grid"
        hrow = table.add_row()
        for i, htext in enumerate(header):
            cell = hrow.cells[i]
            shade_cell(cell, header_color)
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = para.add_run(htext)
            set_font(r, size=10.5, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for row_vals in rows:
            row = table.add_row()
            for i, val in enumerate(row_vals):
                cell = row.cells[i]
                para = cell.paragraphs[0]
                if i > 0:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = para.add_run(str(val))
                set_font(r, size=10.5)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        return table

    def save(self, path):
        self.doc.save(path)


def get_image_bytes(images, tag, nth=0):
    matches = [im for im in images if im["tag"] == tag]
    if len(matches) > nth:
        return matches[nth]["bytes"]
    return None


def build_document(prop, prop_imgs, lap, lap_imgs, output_path):
    b = Builder()

    lokasi_parts = prop.get("_lokasi_parts") or []
    desa = lokasi_parts[0] if len(lokasi_parts) > 0 else g(prop, "_desa")
    kecamatan = lokasi_parts[1] if len(lokasi_parts) > 1 else desa
    kabupaten = lokasi_parts[2] if len(lokasi_parts) > 2 else NA
    provinsi = lokasi_parts[3] if len(lokasi_parts) > 3 else NA
    lokasi_lengkap = f"Desa {desa}, Kecamatan {kecamatan}, {kabupaten}, Provinsi {provinsi}"

    perusahaan = g(prop, "Nama Perusahaan/Instansi")
    perairan = g(prop, "Nama Perairan")
    luas = g(prop, "Luas Kebutuhan Ruang")
    jenis_kegiatan = g(prop, "Jenis Kegiatan")

    # ================= COVER =================
    title = b.doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("PROPOSAL TEKNIS")
    set_font(r, size=20, bold=True, color=NAVY)

    sub = b.doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("PERMOHONAN PERSETUJUAN KESESUAIAN KEGIATAN\nPEMANFAATAN RUANG LAUT (PKKPRL)")
    set_font(r, size=14, bold=True)

    note = b.doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = note.add_run("Disusun mengacu pada Peraturan Menteri Kelautan dan Perikanan Nomor 28 Tahun 2021 "
                     "tentang Penyelenggaraan Penataan Ruang Laut")
    set_font(r, size=10, italic=True)
    note.paragraph_format.space_after = Pt(20)

    b.kv_table([
        ("Nama Pemohon", g(prop, "Nama Pemohon")),
        ("Jabatan Pemohon", g(prop, "Jabatan Pemohon")),
        ("Nama Perusahaan/Instansi", perusahaan),
        ("NIB", g(prop, "NIB")),
        ("NPWP", g(prop, "NPWP")),
        ("Nomor Telepon Selular", g(prop, "Nomor Telepon Selular")),
        ("Surat Elektronik", g(prop, "Surat Elektronik")),
        ("Jenis Kegiatan", jenis_kegiatan),
        ("Lokasi Kegiatan", lokasi_lengkap),
        ("Nama Perairan", perairan),
        ("Luas Kebutuhan Ruang", luas),
        ("KBLI", g(prop, "KBLI")),
        ("Tanggal Penyusunan", g(prop, "Tanggal Penyusunan")),
    ])
    b.page_break()

    # ================= BAB I =================
    b.h1("I. RENCANA BANGUNAN DAN INSTALASI LAUT")
    b.h2("Pendahuluan")
    b.p(f"Proposal teknis ini disusun sebagai bagian dari persyaratan permohonan Persetujuan Kesesuaian "
        f"Kegiatan Pemanfaatan Ruang Laut (PKKPRL), sebagaimana diatur dalam Peraturan Pemerintah Nomor 21 "
        f"Tahun 2021 tentang Penyelenggaraan Penataan Ruang, Peraturan Menteri Kelautan dan Perikanan Nomor "
        f"28 Tahun 2021 tentang Penyelenggaraan Penataan Ruang Laut, serta ketentuan pelaksanaan pada sistem "
        f"OSS Berbasis Risiko.")
    b.p(f"{perusahaan} yang diwakili oleh {g(prop, 'Nama Pemohon')} berencana menyelenggarakan kegiatan "
        f"berusaha berupa {jenis_kegiatan}. Rencana kegiatan ini berlokasi di {lokasi_lengkap}, menggunakan "
        f"perairan {perairan} dengan total kebutuhan luas ruang laut yang dimohonkan sebesar {luas}.")

    b.h2("A. Rencana Kegiatan Utama dan Penunjang")
    b.h3("1. Uraian Kegiatan")
    b.p(f"Kegiatan yang dimohonkan adalah {jenis_kegiatan}, dengan kebutuhan ruang laut seluas {luas}.")
    deskripsi_keg = prop.get("deskripsi_kegiatan", "")
    if deskripsi_keg:
        b.labeled("Deskripsi Kegiatan", deskripsi_keg)
    manfaat_keg = prop.get("manfaat_kegiatan", "")
    if manfaat_keg:
        b.labeled("Manfaat Kegiatan", manfaat_keg)
    tujuan_keg = prop.get("tujuan_kegiatan", "")
    if tujuan_keg:
        b.labeled("Tujuan Kegiatan", tujuan_keg)
    invest = g(prop, "investasi")
    invest_str = f"Rp{invest}" if invest != NA else NA
    tenaga = g(prop, "tenaga_kerja")
    tenaga_asing = g(prop, "tenaga_kerja_asing", "0")
    b.p(f"Rencana tenaga kerja yang digunakan berjumlah {tenaga} orang per siklus, dengan tenaga kerja asing "
        f"berjumlah {tenaga_asing}. Total komitmen pendanaan investasi kegiatan ini sebesar {invest_str}, "
        f"mencakup perencanaan teknis, pengadaan sarana-prasarana, operasional, serta pengelolaan lingkungan hidup.")

    instalasi_bangunan = prop.get("instalasi_bangunan", "")
    instalasi_posisi = prop.get("instalasi_posisi", "")
    if instalasi_bangunan or instalasi_posisi:
        posisi_txt = f", berada pada {instalasi_posisi.lower()}" if instalasi_posisi else ""
        bangunan_txt = instalasi_bangunan if instalasi_bangunan else "instalasi penunjang kegiatan"
        b.p(f"Instalasi bangunan menetap di laut yang direncanakan berupa {bangunan_txt}{posisi_txt}.")

    dukung = prop.get("dokumen_data_dukung", "")
    if dukung:
        b.p(f"Dokumen data dukung yang telah dimiliki oleh pelaku usaha meliputi: {dukung}.")

    h3_2 = b.h3("2. Kegiatan Eksisting atau Rencana yang Akan Dimohonkan")
    status_map = {
        "Eksisting": "merupakan kegiatan yang sudah berjalan (eksisting)",
        "Rencana": "merupakan kegiatan yang baru akan direncanakan",
        "Eksisting dan Pengembangan": "merupakan kegiatan eksisting yang akan dikembangkan lebih lanjut",
    }
    status_txt = status_map.get(prop.get("kegiatan_status", ""), "")
    status_sentence = f" Kegiatan ini {status_txt}." if status_txt else ""
    b.p(f"Kegiatan rencana yang dimohonkan adalah {jenis_kegiatan} yang berada di {lokasi_lengkap}, menggunakan "
        f"perairan {perairan} dengan total kebutuhan luas ruang laut sebesar {luas}.{status_sentence} Pengajuan PKKPRL dilakukan "
        f"dalam rangka pemenuhan perizinan dasar di lokasi yang dimohonkan sebelum mengajukan perizinan lanjutan.")
    img_no = 1
    site_img = get_image_bytes(prop_imgs, "siteplan")
    if site_img:
        b.image(site_img, width_cm=13)
        b.caption(f"Gambar {img_no}. Peta Rencana Tapak (Site Plan) Kegiatan {perusahaan}.")
    else:
        b.image_missing("siteplan")
        b.caption(f"Gambar {img_no}. Peta Rencana Tapak (Site Plan) Kegiatan {perusahaan}.")
    img_no += 1

    dok_kegiatan_img = get_image_bytes(prop_imgs, "dok_kegiatan_eksisting")
    if dok_kegiatan_img:
        b.image(dok_kegiatan_img, width_cm=13)
        b.caption(f"Gambar {img_no}. Dokumentasi Kegiatan Eksisting/Rencana yang Dimohonkan.")
        img_no += 1

    b.h3("3. Rencana Jadwal Pelaksanaan Kegiatan Utama dan Pendukungnya")
    activities = parse_jadwal_kegiatan(prop.get("jadwal_kegiatan", ""))
    bangunan_txt2 = instalasi_bangunan if instalasi_bangunan else "instalasi penunjang kegiatan"
    posisi_txt2 = f" yang berada di {instalasi_posisi}" if instalasi_posisi else ""
    if activities:
        b.p(f"Adapun kegiatan utama yang akan dilakukan ialah pembangunan dan pengembangan lokasi {jenis_kegiatan} "
            f"akan dilakukan sebagaimana ditampilkan pada Tabel 1. Seluruh bangunan merupakan {bangunan_txt2}{posisi_txt2}.")
        start_bt = parse_tanggal_indonesia(prop.get("Tanggal Penyusunan", "")) or (
            __import__("datetime").datetime.now().month, __import__("datetime").datetime.now().year
        )
        build_gantt_table(b, activities, start_bt)
        b.caption("Tabel 1. Rencana Jadwal Pelaksanaan Kegiatan Utama dan Pendukungnya.")
    else:
        jadwal = prop.get("jadwal_kegiatan", "")
        if jadwal:
            b.labeled("Jadwal Kegiatan", jadwal)
        else:
            b.p(NA)

    b.h3("4. Reklamasi / Non-Reklamasi")
    reklamasi_txt = "tanpa reklamasi" if prop.get("non_reklamasi") else "dengan reklamasi"
    b.p(f"Kegiatan {jenis_kegiatan} yang dilakukan oleh {perusahaan} merupakan kegiatan yang dilaksanakan {reklamasi_txt}.")

    b.h2("B. Kegiatan Berusaha atau Non-Berusaha")
    berusaha_txt = "kegiatan berusaha" if prop.get("kegiatan_berusaha") else "kegiatan non-berusaha"
    b.p(f"Kegiatan {jenis_kegiatan} yang dilakukan {perusahaan} di {lokasi_lengkap}, yang menggunakan perairan "
        f"{perairan}, merupakan {berusaha_txt}.")

    b.h2("C. Kegiatan Strategis Nasional atau Nonstrategis Nasional")
    strategis_txt = "kegiatan non-strategis nasional/dasar" if prop.get("non_strategis") else "kegiatan strategis nasional"
    b.p(f"Rencana kegiatan pemanfaatan ruang laut ini tergolong sebagai {strategis_txt}. Penetapan status ini "
        f"digunakan sebagai acuan untuk memenuhi persyaratan teknis permohonan PKKPRL.")

    b.h2("D. Peta Lokasi")
    koord = prop.get("koordinat") or []
    if koord:
        b.p("Peta lokasi/plotting batas-batas area yang dimohonkan PKKPRL ditunjukkan oleh titik koordinat berikut:")
        b.data_table(["Nomor Titik", "Longitude", "Latitude"], koord)
        b.caption("Tabel 1. Titik Koordinat Batas Area Permohonan PKKPRL.")
    else:
        b.p(NA)

    peta_lokasi_img = get_image_bytes(prop_imgs, "peta_lokasi")
    if peta_lokasi_img:
        b.image(peta_lokasi_img, width_cm=11)
    else:
        b.image_missing("peta_lokasi")
    b.caption(f"Gambar {img_no}. Peta Lokasi dan Sebaran Titik Koordinat Rencana Kegiatan.")
    img_no += 1
    sumber_peta = prop.get("sumber_peta", "")
    if sumber_peta:
        b.caption(f"Sumber Peta: {sumber_peta}")

    b.h2("E. Deskripsi Luas/Panjang yang Dibutuhkan")
    b.p(f"Luas perairan yang dimohonkan KKPRL adalah seluas {luas} yang terletak di perairan {perairan}, "
        f"{lokasi_lengkap}.")

    b.page_break()

    # ================= BAB II =================
    b.h1("II. INFORMASI PEMANFAATAN RUANG LAUT")
    b.p(f"Berdasarkan hasil identifikasi, pemanfaatan ruang laut eksisting di sekitar lokasi kegiatan "
        f"{jenis_kegiatan} dari {perusahaan} termasuk dalam wilayah perairan {perairan}, pada administrasi "
        f"{lokasi_lengkap}.")
    b.p("Berdasarkan hasil survei/pengamatan langsung, tidak terdapat pemanfaatan ruang laut oleh pihak lain "
        "di sekitar lokasi permohonan. Rencana kegiatan disusun dengan memperhatikan kepentingan nelayan "
        "tradisional dan masyarakat, serta tidak menghalangi akses pelayaran yang sudah ada.")
    batas_utara = prop.get("batas_utara", "")
    batas_timur = prop.get("batas_timur", "")
    batas_selatan = prop.get("batas_selatan", "")
    batas_barat = prop.get("batas_barat", "")
    if batas_utara or batas_timur or batas_selatan or batas_barat:
        kalimat_arah = ["Pemanfaatan ruang laut di sekitar lokasi permohonan didominasi oleh aktivitas "
                        "penangkapan ikan skala kecil di seluruh penjuru arah mata angin."]
        if batas_utara:
            kalimat_arah.append(f"Di sebelah utara, lokasi berbatasan dengan {batas_utara}.")
        if batas_timur:
            kalimat_arah.append(f"Pada sisi timur, kawasan berbatasan dengan {batas_timur}.")
        if batas_selatan:
            kalimat_arah.append(f"Sementara itu, di sebelah selatan terdapat {batas_selatan}.")
        if batas_barat:
            kalimat_arah.append(f"Di sebelah barat berbatasan dengan {batas_barat}.")
        b.p(" ".join(kalimat_arah))

    deskripsi_sekitar = prop.get("deskripsi_pemanfaatan_sekitar", "")
    if deskripsi_sekitar:
        b.p(deskripsi_sekitar)

    foto_pantai_list = [im for im in prop_imgs if im["tag"] == "foto_pantai"]
    if foto_pantai_list:
        for idx, im in enumerate(foto_pantai_list):
            b.image(im["bytes"], width_cm=11)
            label = "Kondisi Eksisting Perairan dan Garis Pantai" if idx == 0 else "Dokumentasi Lapangan Tambahan Kondisi Garis Pantai"
            b.caption(f"Gambar {img_no}. {label} di Sekitar Lokasi Permohonan.")
            img_no += 1
    else:
        b.image_missing("foto_pantai")
        b.caption(f"Gambar {img_no}. Kondisi Eksisting Perairan dan Garis Pantai di Sekitar Lokasi Permohonan.")
        img_no += 1

    dok_sekitar_list = [im for im in prop_imgs if im["tag"] == "dok_pemanfaatan_sekitar"]
    for im in dok_sekitar_list:
        b.image(im["bytes"], width_cm=11)
        b.caption(f"Gambar {img_no}. Dokumentasi Kegiatan Pemanfaatan Ruang Laut Sekitar.")
        img_no += 1

    b.page_break()

    # ================= BAB III =================
    b.h1("III. DATA KONDISI TERKINI LOKASI DAN SEKITARNYA")

    b.h2("A. Ekosistem Sekitar")
    b.h3("1. Mangrove")
    mangrove_ada = prop.get("mangrove_ada", "")
    if mangrove_ada == "Tidak terdapat ekosistem mangrove":
        b.p("Berdasarkan hasil pengamatan langsung kondisi pesisir di sekitar lokasi kegiatan, tidak "
            "teridentifikasi keberadaan ekosistem mangrove pada area yang dimohonkan.")
    else:
        spesies = g(prop, "mangrove_spesies")
        persen_mgv = g(prop, "mangrove_persen")
        kondisi_mgv = g(prop, "mangrove_kondisi")
        b.p(f"Berdasarkan hasil pengamatan langsung kondisi pesisir di sekitar lokasi kegiatan, terdapat ekosistem "
            f"mangrove yang didominasi oleh jenis {spesies}, dengan persentase tutupan mencapai {persen_mgv}% "
            f"pada kondisi {kondisi_mgv}.")
    mgv_img = get_image_bytes(prop_imgs, "foto_mangrove")
    if mgv_img:
        b.image(mgv_img, width_cm=11)
    else:
        b.image_missing("foto_mangrove")
    b.caption(f"Gambar {img_no}. Kondisi Tutupan Vegetasi Mangrove di Sekitar Lokasi Kegiatan.")
    img_no += 1

    b.h3("2. Lamun")
    ada_lamun = lap.get("ada_lamun")
    jarak_eko = g(lap, "eko_jarak_terdekat_km")
    if ada_lamun:
        b.p(f"Berdasarkan data sekunder hasil pengamatan lapangan awal, tidak teridentifikasi keberadaan "
            f"ekosistem lamun (seagrass) secara langsung pada titik pengambilan sampel di area yang dimohonkan. "
            f"Namun demikian, berdasarkan analisis spasial basis data ekosistem nasional, ekosistem padang "
            f"lamun teridentifikasi berada di sekitar lokasi, dengan jarak ekosistem terdekat dari titik pusat "
            f"rencana kegiatan sekitar {jarak_eko} km. Disarankan verifikasi lapangan lanjutan untuk memastikan "
            f"keberadaan dan luasan ekosistem lamun secara lebih akurat.")
    else:
        b.p("Berdasarkan data sekunder perairan di sekitar lokasi kegiatan, tidak teridentifikasi keberadaan "
            "ekosistem lamun (seagrass) pada area yang dimohonkan.")

    lamun_ada_manual = prop.get("lamun_ada_manual", "")
    if lamun_ada_manual == "Terdapat ekosistem lamun":
        lamun_spesies = g(prop, "lamun_spesies")
        lamun_persen = g(prop, "lamun_persen")
        lamun_kondisi = g(prop, "lamun_kondisi")
        b.p(f"Berdasarkan hasil pengamatan pemohon di lapangan, teridentifikasi ekosistem lamun yang didominasi "
            f"oleh jenis {lamun_spesies}, dengan persentase tutupan mencapai {lamun_persen}% pada kondisi {lamun_kondisi}.")
        lamun_img = get_image_bytes(prop_imgs, "foto_lamun")
        if lamun_img:
            b.image(lamun_img, width_cm=11)
            b.caption(f"Gambar {img_no}. Dokumentasi Ekosistem Lamun di Sekitar Lokasi Kegiatan.")
            img_no += 1

    b.h3("3. Terumbu Karang")
    karang_ha = g(lap, "eko_karang_ha")
    karang_pct = g(lap, "eko_karang_pct")
    lainnya_ha = g(lap, "eko_lainnya_ha")
    lainnya_pct = g(lap, "eko_lainnya_pct")
    terbuka_ha = g(lap, "eko_terbuka_ha")
    terbuka_pct = g(lap, "eko_terbuka_pct")
    total_ha = g(lap, "eko_total_ha")
    b.p(f"Hasil survei in-situ pada perairan di sekitar lokasi menunjukkan dijumpainya koloni terumbu karang "
        f"pada beberapa titik substrat berbatu. Berdasarkan analisis spasial basis data ekosistem, dari total "
        f"area kajian seluas {total_ha} Ha, tutupan terumbu karang tercatat seluas {karang_ha} Ha ({karang_pct}%), "
        f"diikuti substrat dasar non-terumbu seluas {lainnya_ha} Ha ({lainnya_pct}%), dan area laut terbuka "
        f"tanpa ekosistem seluas {terbuka_ha} Ha ({terbuka_pct}%).")
    b.data_table(
        ["Jenis Tutupan", "Luas (Ha)", "Persentase (%)"],
        [
            ["Terumbu Karang", karang_ha, karang_pct],
            ["Lainnya (substrat dasar non-terumbu)", lainnya_ha, lainnya_pct],
            ["Area Laut Terbuka (tanpa ekosistem)", terbuka_ha, terbuka_pct],
            ["Total Area Kajian", total_ha, "100,0"],
        ],
    )
    b.caption("Tabel 2. Rincian Tutupan Ekosistem pada Area Kajian Spasial di Sekitar Titik Pusat Rencana Kegiatan.")

    karang_ada_manual = prop.get("karang_ada", "")
    if karang_ada_manual == "Terdapat ekosistem terumbu karang":
        karang_spesies_m = g(prop, "karang_spesies")
        karang_persen_m = g(prop, "karang_persen_manual")
        karang_kondisi_m = g(prop, "karang_kondisi")
        b.p(f"Berdasarkan hasil pengamatan pemohon di lapangan, teridentifikasi ekosistem terumbu karang yang "
            f"didominasi oleh jenis {karang_spesies_m}, dengan persentase tutupan mencapai {karang_persen_m}% "
            f"pada kondisi {karang_kondisi_m}.")
    elif karang_ada_manual == "Tidak terdapat ekosistem terumbu karang":
        b.p("Berdasarkan hasil pengamatan pemohon di lapangan, tidak teridentifikasi keberadaan ekosistem "
            "terumbu karang secara langsung pada area yang dimohonkan.")

    karang_img = get_image_bytes(prop_imgs, "foto_karang_insitu")
    if karang_img:
        b.image(karang_img, width_cm=11)
    else:
        b.image_missing("foto_karang_insitu")
    b.caption(f"Gambar {img_no}. Dokumentasi Survei In-Situ Koloni Terumbu Karang di Perairan Sekitar Lokasi Kegiatan.")
    img_no += 1

    eko_img = get_image_bytes(lap_imgs, "peta_ekosistem")
    if eko_img:
        b.image(eko_img, width_cm=11)
    else:
        b.image_missing("peta_ekosistem")
    b.caption(f"Gambar {img_no}. Peta Sebaran Spasial Ekosistem Pesisir di Sekitar Titik Pusat Rencana Kegiatan.")
    img_no += 1
    b.p(f"Jarak ekosistem terdekat dari titik pusat rencana kegiatan adalah {jarak_eko} km, sehingga mitigasi "
        f"dampak perlu difokuskan pada upaya penghindaran (avoidance) terhadap area terumbu karang, pengendalian "
        f"sedimen, serta pengelolaan kualitas air.")

    b.h2("B. Hidro-Oseanografi")
    b.h3("1. Gelombang")
    b.p(f"Tinggi gelombang signifikan (Hs) rata-rata tercatat sebesar {g(lap,'hs_rata')} meter, sedangkan Hs "
        f"maksimum ekstrem tercatat sebesar {g(lap,'hs_maks')} meter dengan arah dominan dari "
        f"{g(lap,'hs_arah')}\u00b0. Parameter ini menjadi acuan utama dalam desain ketahanan struktur bangunan laut "
        f"terhadap beban gelombang ekstrem.")
    gel_img = get_image_bytes(lap_imgs, "mawar_gelombang")
    if gel_img:
        b.image(gel_img, width_cm=9)
    else:
        b.image_missing("mawar_gelombang")
    b.caption(f"Gambar {img_no}. Mawar Gelombang Ekstrem pada Titik Pusat Rencana Kegiatan (Arah Dominan {g(lap,'hs_arah')}\u00b0).")
    img_no += 1

    b.h3("2. Arus")
    b.p(f"Kecepatan arus rata-rata tercatat sebesar {g(lap,'arus_rata')} m/detik, dengan kecepatan maksimum "
        f"ekstrem sebesar {g(lap,'arus_maks')} m/detik dan arah dominan dari {g(lap,'arus_arah')}\u00b0. Parameter "
        f"ini menjadi indikator potensi gerusan (scouring) di sekitar struktur bangunan laut.")
    arus_img = get_image_bytes(lap_imgs, "mawar_arus")
    if arus_img:
        b.image(arus_img, width_cm=9)
    else:
        b.image_missing("mawar_arus")
    b.caption(f"Gambar {img_no}. Mawar Arus pada Titik Pusat Rencana Kegiatan (Arah Dominan {g(lap,'arus_arah')}\u00b0).")
    img_no += 1

    b.data_table(
        ["Parameter", "Nilai Rata-rata", "Nilai Ekstrem", "Arah Dominan"],
        [
            ["Tinggi Gelombang Signifikan (Hs)", f"{g(lap,'hs_rata')} m", f"{g(lap,'hs_maks')} m", f"{g(lap,'hs_arah')}\u00b0"],
            ["Kecepatan Arus", f"{g(lap,'arus_rata')} m/detik", f"{g(lap,'arus_maks')} m/detik", f"{g(lap,'arus_arah')}\u00b0"],
        ],
    )
    b.caption("Tabel 3. Ringkasan Parameter Gelombang dan Arus pada Titik Pusat Rencana Kegiatan.")

    b.h3("3. Pasang Surut")
    b.p(f"Perairan ini memiliki tipe pasang surut {g(lap,'tipe_pasut')} (Bilangan Formzahl {g(lap,'formzahl')}), "
        f"dengan tunggang air (tidal range) sebesar {g(lap,'tidal_range')} meter, elevasi tertinggi (HAT) "
        f"sebesar +{g(lap,'hat')} meter, dan elevasi terendah (LAT) sebesar {g(lap,'lat')} meter.")
    b.data_table(
        ["Parameter Pasang Surut", "Elevasi"],
        [
            ["Highest Astronomical Tide (HAT)", f"+{g(lap,'hat')} m"],
            ["Mean Sea Level (MSL)", f"{g(lap,'msl')} m"],
            ["Lowest Astronomical Tide (LAT)", f"{g(lap,'lat')} m"],
            ["Tidal Range", f"{g(lap,'tidal_range')} m"],
            ["Bilangan Formzahl", f"{g(lap,'formzahl')} ({g(lap,'tipe_pasut')})"],
        ],
    )
    b.caption("Tabel 4. Parameter Pasang Surut pada Lokasi Kegiatan.")
    pasut_img = get_image_bytes(lap_imgs, "siklus_pasut")
    if pasut_img:
        b.image(pasut_img, width_cm=13)
    else:
        b.image_missing("siklus_pasut")
    b.caption(f"Gambar {img_no}. Grafik Fluktuasi Pasang Surut Selama 14 Hari (Tipe {g(lap,'tipe_pasut')}).")
    img_no += 1

    b.h2("C. Profil Dasar Laut")
    b.p(f"Kedalaman pada titik pusat lokasi kegiatan tercatat sebesar {g(lap,'batimetri_titik_pusat')} meter "
        f"terhadap Lowest Water Spring (LWS). Hasil pemeruman pada profil garis batimetri sepanjang lintasan "
        f"{g(lap,'batimetri_panjang_lintasan')} km menunjukkan kedalaman terdalam mencapai "
        f"{g(lap,'batimetri_terdalam')} meter.")
    bati_img = get_image_bytes(lap_imgs, "profil_batimetri")
    if bati_img:
        b.image(bati_img, width_cm=13)
    else:
        b.image_missing("profil_batimetri")
    b.caption(f"Gambar {img_no}. Profil Garis Batimetri pada Lintasan Pemeruman Titik Pusat Rencana Kegiatan.")
    img_no += 1

    b.h2("D. Kondisi Sosial Ekonomi Masyarakat")
    sumber_sosek = prop.get("sumber_data_sosek", "") or "Badan Pusat Statistik"
    tahun_sosek = prop.get("tahun_data_sosek", "")
    tahun_txt = f" tahun {tahun_sosek}" if tahun_sosek else ""
    b.p(f"Berdasarkan data sekunder {sumber_sosek}{tahun_txt}, {desa} memiliki luas wilayah {g(prop,'desa_luas_ha')} "
        f"Ha dengan jumlah penduduk sebanyak {g(prop,'desa_penduduk')} jiwa. Kehadiran rencana kegiatan ini "
        f"diharapkan dapat mendukung struktur sosial-ekonomi kawasan secara harmonis dan melibatkan konsultasi "
        f"publik dengan kelompok nelayan setempat sebelum pelaksanaan konstruksi.")
    mata_pencaharian = prop.get("mata_pencaharian", "")
    if mata_pencaharian:
        b.labeled("Mata Pencaharian Masyarakat Desa", mata_pencaharian)

    b.h2("E. Aksesibilitas Lokasi dan Sekitarnya")
    aksesibilitas_manual = prop.get("aksesibilitas_lokasi", "")
    if aksesibilitas_manual:
        b.p(aksesibilitas_manual)
    else:
        b.p(f"Aksesibilitas menuju lokasi kegiatan di {lokasi_lengkap} dapat ditempuh melalui jalur darat maupun laut.")

    akses_img = get_image_bytes(prop_imgs, "gambar_aksesibilitas")
    if akses_img:
        b.image(akses_img, width_cm=13)
        b.caption(f"Gambar {img_no}. Peta Aksesibilitas Menuju Lokasi Kegiatan.")
        img_no += 1

    pola_img = get_image_bytes(prop_imgs, "peta_pola_ruang")
    if pola_img:
        b.image(pola_img, width_cm=13)
    else:
        b.image_missing("peta_pola_ruang")
    b.caption(f"Gambar {img_no}. Peta Rencana Pola Ruang Wilayah dan Posisi Lokasi Permohonan.")
    img_no += 1

    b.page_break()
    b.h1("IV. DOKUMEN PERSYARATAN LAINNYA")
    b.p(f"Dokumen pendukung untuk permohonan PKKPRL yang diajukan oleh {perusahaan} meliputi:")
    b.bullet("Sertifikat Kepemilikan Lahan Darat.")
    b.bullet("Dokumen identitas dan legalitas pemohon/perusahaan.")
    b.bullet("Dokumentasi survei lapangan kondisi eksisting lokasi.")
    b.bullet("Peta pendukung (peta lokasi, peta site plan, dan peta pola ruang wilayah).")

    sertifikat_img = get_image_bytes(prop_imgs, "sertifikat_lahan")
    if sertifikat_img:
        b.image(sertifikat_img, width_cm=13)
        b.caption(f"Gambar {img_no}. Sertifikat Kepemilikan Lahan Darat.")
        img_no += 1

    sosialisasi_img = get_image_bytes(prop_imgs, "dok_sosialisasi")
    if sosialisasi_img:
        b.image(sosialisasi_img, width_cm=13)
        b.caption(f"Gambar {img_no}. Dokumen Hasil Sosialisasi dengan Masyarakat Sekitar.")
        img_no += 1

    pendukung_lain_list = [im for im in prop_imgs if im["tag"] == "dok_pendukung_lainnya"]
    for im in pendukung_lain_list:
        b.image(im["bytes"], width_cm=13)
        b.caption(f"Gambar {img_no}. Dokumen Pendukung Lainnya.")
        img_no += 1

    b.p(f"Demikian proposal teknis ini disusun sebagai bagian dari kelengkapan administrasi dan teknis "
        f"permohonan PKKPRL atas nama {perusahaan}.", italic=True)
    b.p("Catatan: Dokumen ini dibangkitkan otomatis oleh aplikasi penggabung proposal PKKPRL dari dua sumber "
        "dokumen. Mohon verifikasi kembali seluruh data dan gambar sebelum digunakan untuk pengajuan resmi.",
        italic=True)

    b.save(output_path)
    return img_no - 1
