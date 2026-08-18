"""
Modul ekstraksi data & gambar dari:
1. Draft Proposal PKKPRL (PDF)
2. Laporan Teknis Hidro-Oseanografi / Kondisi Eksisting Ekosistem (PDF)

Catatan teknis: PDF sumber mengekstrak teks dengan satu kata per baris pada
beberapa bagian, sehingga semua regex dijalankan terhadap teks yang sudah
dinormalisasi (semua whitespace/newline diubah menjadi satu spasi).
"""
import os
import re
import hashlib
import fitz  # PyMuPDF


def norm(s):
    s = (s or "").replace("\u200b", " ").replace("\ufeff", " ").replace("\xa0", " ")
    return re.sub(r"\s+", " ", s).strip()


def klasifikasi_tipe_pasut(formzahl_str):
    """Tentukan tipe pasang surut dari NILAI ANGKA Bilangan Formzahl (F),
    berdasarkan kategori baku berikut:
      0    < F <= 0.25 : Semidiurnal (pasang surut harian ganda)
      0.25 < F <= 1.50 : Mixed Semidiurnal (campuran condong harian ganda)
      1.50 < F <= 3.00 : Mixed Diurnal (campuran condong harian tunggal)
      F > 3.00         : Diurnal (pasang surut harian tunggal)
    PENTING: tipe ini dihitung dari ANGKA Formzahl-nya sendiri, BUKAN dicari
    lewat teks naratif di sekitarnya (mis. "...diklasifikasikan sebagai
    X,...") -- pendekatan berbasis teks itu rapuh dan pernah terbukti salah
    ambil teks yang sama sekali tidak relevan (mis. dari baris tabel sumber
    yang typo/salah isi, seperti '(habitat biota laut)' alih-alih nama tipe
    pasutnya), padahal angka Formzahl-nya sendiri sudah cukup untuk
    menentukan tipe secara pasti dan konsisten di seluruh dokumen.
    Return "" kalau formzahl_str tidak valid/kosong (biar pemanggil bisa
    fallback ke default)."""
    try:
        f = float(str(formzahl_str).replace(",", ".").strip())
    except (TypeError, ValueError):
        return ""
    if f <= 0:
        return ""
    if f <= 0.25:
        return "Semidiurnal"
    if f <= 1.50:
        return "Mixed Semidiurnal"
    if f <= 3.00:
        return "Mixed Diurnal"
    return "Diurnal"


# ----------------------------------------------------------------------
# PROPOSAL PDF
# ----------------------------------------------------------------------

PROPOSAL_LABELS = [
    "Nama Pemohon", "Jabatan Pemohon", "Nama Perusahaan/Instansi", "NIB",
    "NPWP", "Nomor Telepon Selular", "Surat Elektronik", "Jenis Kegiatan",
    "Lokasi Kegiatan", "Nama Perairan", "Luas Kebutuhan Ruang", "KBLI",
    "Tanggal Penyusunan",
]

# Heading pattern -> section tag, dicek berurutan sesuai kemunculan teks
PROPOSAL_HEADINGS = [
    (r"2\.\s*Kegiatan Eksisting", "siteplan"),
    (r"E\.\s*RENCANA TAPAK", "siteplan"),
    (r"D\.\s*PETA LOKASI", "peta_lokasi"),
    (r"II\.\s*INFORMASI PEMANFAATAN", "foto_pantai"),
    (r"1\.\s*Mangrove", "foto_mangrove"),
    (r"2\.\s*Lamun", "foto_mangrove"),
    (r"3\.\s*Terumbu Karang", "terumbu_karang_section"),
    (r"E\.\s*AKSESIBILITAS", "peta_pola_ruang"),
]


def _parse_proposal_text(full_text, full_text_raw):
    """Semua regex parsing Draft Proposal PKKPRL (field teks, lokasi, koordinat,
    investasi, dst), dipakai bersama baik sumbernya PDF maupun DOCX."""
    data = {}

    for idx, label in enumerate(PROPOSAL_LABELS):
        next_label = PROPOSAL_LABELS[idx + 1] if idx + 1 < len(PROPOSAL_LABELS) else r"I\.\s*RENCANA"
        pattern = re.escape(label) + r"\s*(.*?)\s*" + next_label
        m = re.search(pattern, full_text, re.DOTALL)
        data[label] = norm(m.group(1)) if m else ""

    lokasi_raw = data.get("Lokasi Kegiatan", "")
    data["_lokasi_parts"] = [p for p in re.split(r"\s{1,}", lokasi_raw) if p.strip()]

    # Lokasi Kegiatan pada dokumen sumber ada 2 format yang pernah ditemui:
    # (a) 4 baris terpisah (desa/kecamatan/kabupaten/provinsi masing-masing
    #     baris sendiri) -- ambil dari teks MENTAH supaya batas baris tidak
    #     hilang; atau
    # (b) satu baris dipisah koma, mis. "Desa X, Kecamatan Y, Kabupaten Z,
    #     Provinsi W" -- format ini yang paling umum dijumpai di dokumen
    #     instansi pemerintah.
    m_loc = re.search(r"Lokasi Kegiatan\s*\n(.*?)\nNama Perairan", full_text_raw, re.DOTALL)
    if m_loc:
        lines = [norm(x) for x in m_loc.group(1).split("\n") if norm(x)]
        if len(lines) >= 4:
            data["_lokasi_parts"] = lines[:4]
        elif len(lines) == 1 and lines[0].count(",") >= 2:
            # format (b): satu baris, pisahkan per koma lalu buang label
            # "Desa"/"Kecamatan"/"Provinsi" di depan tiap bagian (kabupaten/
            # kota TIDAK dibuang labelnya karena template pemanggil tidak
            # menambahkan prefiks sendiri untuk bagian ini).
            comma_parts = [norm(x) for x in lines[0].split(",") if norm(x)]
            if len(comma_parts) >= 4:
                cleaned = []
                for i, part in enumerate(comma_parts[:4]):
                    if i in (0, 1, 3):  # desa, kecamatan, provinsi
                        part = re.sub(r"^(desa|kecamatan|provinsi)\s+", "", part, flags=re.IGNORECASE)
                    cleaned.append(part)
                data["_lokasi_parts"] = cleaned
            elif lines:
                data["_lokasi_parts"] = lines
        elif lines:
            data["_lokasi_parts"] = lines

    m = re.search(r"PT\.\s*[A-Z .]+?(?=\s+yang diwakili|\s+berencana)", full_text)
    if m:
        data["Nama Perusahaan/Instansi"] = norm(m.group(0))

    # Titik koordinat: dokumen sumber kadang memakai notasi "BT"/"LS"
    # (Bujur Timur/Lintang Selatan), kadang notasi internasional "E"/"S"
    # (East/South). Regex ini menangkap keduanya, dan berapapun jumlah titiknya.
    koord = re.findall(
        r"(\d+)\s+(\d+°\s*\d+'\s*[\d,]+\"\s*(?:BT|E))\s+(\d+°\s*\d+'\s*[\d,]+\"\s*(?:LS|S))",
        full_text,
    )
    data["koordinat"] = koord

    m = re.search(r"komitmen\s*pendanaan\s*investasi\s*secara\s*keseluruhan\s*sebesar\s*([\d.,]+)", full_text)
    data["investasi"] = m.group(1).rstrip(".") if m else ""

    m = re.search(r"berjumlah\s*(\d+)\s*orang\s*per\s*siklus", full_text)
    data["tenaga_kerja"] = m.group(1) if m else ""
    m = re.search(r"tenaga\s*kerja\s*asing\s*berjumlah\s*(\d+)", full_text)
    data["tenaga_kerja_asing"] = m.group(1) if m else "0"

    m = re.search(
        r"didominasi\s*oleh\s*([A-Za-z][A-Za-z .]+?)\s*dengan\s*persentase\s*tutupan\s*mangrove\s*(\d+)%\s*kondisi\s*([A-Za-z ]+?)\s*serta",
        full_text,
    )
    if m:
        data["mangrove_spesies"] = norm(m.group(1))
        data["mangrove_persen"] = m.group(2)
        data["mangrove_kondisi"] = norm(m.group(3))

    m = re.search(r"seluas\s*(\d+)\s*Ha\s*dengan\s*jumlah\s*penduduk\s*sebanyak\s*([\d.]+)", full_text)
    if m:
        data["desa_luas_ha"] = m.group(1)
        data["desa_penduduk"] = m.group(2).rstrip(".")

    data["non_reklamasi"] = "tanpa reklamasi" in full_text.lower()
    data["kegiatan_berusaha"] = "kegiatan berusaha" in full_text.lower()
    data["non_strategis"] = "non-strategis nasional" in full_text.lower()
    return data


def _nearby_caption_tag(para_texts_norm, pi, patterns, before=3, after=3):
    """Cari caption gambar di paragraf-paragraf SEKITAR posisi gambar (pi),
    baik sebelum maupun sesudah, diprioritaskan dari yang paling dekat --
    karena konvensi penulisan caption berbeda-beda antar dokumen (ada yang
    menaruh caption di atas gambar, ada yang di bawah)."""
    offsets = [0]
    for d in range(1, max(before, after) + 1):
        if d <= after:
            offsets.append(d)
        if d <= before:
            offsets.append(-d)
    for off in offsets:
        idx = pi + off
        if 0 <= idx < len(para_texts_norm):
            cand = para_texts_norm[idx]
            if not cand:
                continue
            tag = _tag_from_caption(cand, patterns)
            if tag:
                return tag
    return None


def _tag_proposal_image(current_section, section_img_count):
    if current_section == "terumbu_karang_section":
        return "foto_pantai" if section_img_count == 0 else "foto_karang_insitu"
    return current_section or "lainnya"


# ----------------------------------------------------------------------
# PENCOCOKAN GAMBAR -> TAG BERDASARKAN CAPTION ASLI DI DEKAT GAMBAR
# ----------------------------------------------------------------------
# Cara lama menandai gambar berdasar section/heading TERAKHIR yang terlihat
# di halaman/paragraf sebelumnya -- ini gampang keliru kalau gambar (mis.
# dokumentasi/lampiran foto) posisinya di dokumen sumber tidak persis
# mengikuti urutan section di template (contoh nyata: screenshot rapat yang
# kebetulan muncul di section "Rencana Tapak" ikut ditandai sebagai foto
# site plan). Caption asli ("Gambar N. <deskripsi>") yang menempel tepat
# di gambarnya sendiri jauh lebih dipercaya karena tidak bergantung pada
# section/urutan fisik gambar di dokumen sumber -- maka ini jadi metode
# UTAMA, dengan cara lama sebagai fallback kalau caption tidak ditemukan.

CAPTION_RE = re.compile(r"^(gambar|lampiran|foto)\s*\d*[.:]?\s*", re.IGNORECASE)

CAPTION_PATTERNS_PROPOSAL = [
    (r"site\s*plan|rencana\s*tapak", "siteplan"),
    (r"peta\s*lokasi|sebaran\s*titik\s*koordinat|\bkoordinat\b", "peta_lokasi"),
    (r"garis\s*pantai|eksisting\s*perairan", "foto_pantai"),
    (r"mangrove", "foto_mangrove"),
    (r"in-?situ|terumbu\s*karang", "foto_karang_insitu"),
    (r"pola\s*ruang", "peta_pola_ruang"),
]

CAPTION_PATTERNS_LAPORAN = [
    # Pola spesifik (2+ kata / frasa khas) dicek DULU supaya caption yang
    # kebetulan menyebut lebih dari satu istilah domain (mis. "...Ekosistem
    # Pesisir yang berpotensi terdampak Gelombang...") tetap tertandai sesuai
    # topik UTAMA captionnya, bukan kata pertama yang match.
    (r"mawar\s*gelombang", "mawar_gelombang"),
    (r"mawar\s*arus", "mawar_arus"),
    (r"(fluktuasi\s*)?pasang\s*surut|siklus\s*pasut", "siklus_pasut"),
    (r"tutupan\s*ekosistem|analisis\s*spasial(\s*\w+){0,4}\s*ekosistem|(peta\s*)?sebaran\s*(spasial\s*)?ekosistem|peta\s*ekosistem", "peta_ekosistem"),
    (r"profil\s*(garis\s*)?batimetri|kontur\s*(&|dan)?\s*layout|(kontur\s*)?batimetri", "profil_batimetri"),
    # Fallback frasa/kata lebih umum -- dicek belakangan supaya tidak
    # "mencuri" caption yang sebenarnya lebih cocok ke pola spesifik di atas.
    (r"ekosistem\s*pesisir|\bekosistem\b", "peta_ekosistem"),
    (r"layout\s*area", "profil_batimetri"),
    (r"\bgelombang\b", "mawar_gelombang"),
    (r"\barus\b", "mawar_arus"),
    (r"\bpasut\b", "siklus_pasut"),
]

# Urutan prioritas TAG (bukan urutan pola) dipakai oleh resolve_laporan_tags
# di bawah -- tiap tag "mengklaim" kandidat gambar dengan caption paling
# spesifik lebih dulu, sebelum tag lain mengambil sisanya. Ini mencegah efek
# domino: kalau satu caption gagal cocok, gambar itu TIDAK lagi "mencuri"
# tag milik gambar lain di section berikutnya (bug lama) -- ia cukup jatuh
# ke "lainnya".
LAPORAN_TAG_PRIORITY = ["mawar_gelombang", "mawar_arus", "siklus_pasut", "peta_ekosistem", "profil_batimetri"]


def _tag_from_caption(caption_text, patterns):
    """Cocokkan teks caption (biasanya berbunyi 'Gambar N. <deskripsi>') ke
    tag yang tepat berdasar kata kunci pada DESKRIPSI caption itu sendiri --
    bukan berdasar section/heading yang mungkin jauh sebelumnya di dokumen.
    Return None kalau tidak ada pola yang cocok, supaya pemanggil bisa
    memakai fallback lama."""
    tag, _rank = _tag_from_caption_ranked(caption_text, patterns)
    return tag


def _tag_from_caption_ranked(caption_text, patterns):
    """Sama seperti _tag_from_caption, tapi juga mengembalikan indeks pola
    yang cocok (makin kecil = makin spesifik, karena 'patterns' disusun dari
    yang paling spesifik ke paling umum). Dipakai resolve_laporan_tags untuk
    memprioritaskan caption yang PALING SPESIFIK ketika lebih dari satu
    gambar sama-sama menyinggung topik yang sama (mis. dua gambar sama-sama
    soal ekosistem, tapi yang satu captionnya cuma \"Ekosistem Pesisir\"
    (umum) sementara yang lain \"Analisis Spasial - Tutupan Ekosistem Area\"
    (lebih spesifik/rinci) -- yang lebih spesifik itu yang harus dipakai)."""
    if not caption_text:
        return None, None
    text = norm(caption_text)
    for rank, (pattern, tag) in enumerate(patterns):
        if re.search(pattern, text, re.IGNORECASE):
            return tag, rank
    return None, None


def _page_caption_for_image(page, img_rect):
    """Cari blok teks caption yang paling dekat secara vertikal dengan
    sebuah gambar pada halaman PDF yang sama (konvensi umum: caption
    'Gambar N. ...' diletakkan tepat di bawah gambarnya; kalau tidak ada,
    coba blok tepat di atasnya sebagai upaya kedua)."""
    if img_rect is None:
        return None
    try:
        blocks = page.get_text("blocks")
    except Exception:
        return None
    below, above = [], []
    for b in blocks:
        x0, y0, x1, y1, text = b[0], b[1], b[2], b[3], b[4]
        text = (text or "").strip()
        if not text:
            continue
        if y0 >= img_rect.y1 - 2:
            below.append((y0 - img_rect.y1, text))
        elif y1 <= img_rect.y0 + 2:
            above.append((img_rect.y0 - y1, text))
    below.sort(key=lambda t: t[0])
    above.sort(key=lambda t: t[0])
    for _, text in below[:2]:
        if CAPTION_RE.match(norm(text)):
            return norm(text)
    for _, text in above[:1]:
        if CAPTION_RE.match(norm(text)):
            return norm(text)
    return None


def extract_proposal(pdf_path):
    doc = fitz.open(pdf_path)
    full_text_raw = "\n".join(page.get_text() for page in doc)
    full_text = norm(full_text_raw)

    data = _parse_proposal_text(full_text, full_text_raw)

    # ---------------- IMAGES (sequential heading-tracking per halaman) ----------------
    seen_hash = set()
    images = []
    current_section = None
    section_img_count = 0

    for pnum in range(len(doc)):
        page = doc[pnum]
        page_text_norm = norm(page.get_text())

        matches = []
        for pattern, tag in PROPOSAL_HEADINGS:
            for mm in re.finditer(pattern, page_text_norm, re.IGNORECASE):
                matches.append((mm.start(), tag))
        if matches:
            matches.sort(key=lambda x: x[0])
            new_section = matches[-1][1]
            if new_section != current_section:
                current_section = new_section
                section_img_count = 0

        imglist = page.get_images(full=True)
        for i, img in enumerate(imglist):
            xref = img[0]
            base = doc.extract_image(xref)
            data_bytes = base["image"]
            h = hashlib.md5(data_bytes).hexdigest()
            w, ht = base.get("width", 0), base.get("height", 0)
            if w < 150 or ht < 150:
                continue

            if h in seen_hash:
                continue
            seen_hash.add(h)

            # 1) coba tandai dari CAPTION ASLI di dekat gambar ini di
            #    halaman yang sama (metode utama, paling akurat).
            rects = page.get_image_rects(xref)
            img_rect = rects[0] if rects else None
            caption = _page_caption_for_image(page, img_rect)
            caption_tag = _tag_from_caption(caption, CAPTION_PATTERNS_PROPOSAL)

            if caption_tag:
                tag = caption_tag
            else:
                # 2) fallback: tebak dari section/heading terakhir yang
                #    terlihat (perilaku lama).
                tag = _tag_proposal_image(current_section, section_img_count)
                section_img_count += 1

            images.append({
                "tag": tag, "bytes": data_bytes, "ext": base["ext"],
                "width": w, "height": ht, "page": pnum + 1,
            })

    doc.close()
    return data, images


def extract_proposal_docx(docx_path):
    """Versi ekstraksi Draft Proposal PKKPRL dari file Word (.docx). Menelusuri
    paragraf secara berurutan (mirip pelacakan per-halaman pada versi PDF,
    tapi di sini per-paragraf -- resolusinya malah lebih presisi karena docx
    tidak punya batas halaman yang tetap)."""
    from docx import Document
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    def _iter_block_items(document):
        """Iterasi Paragraph & Table dalam URUTAN ASLI kemunculannya di body
        dokumen. python-docx secara default memisahkan `doc.paragraphs` dan
        `doc.tables` ke dua koleksi terpisah, sehingga urutan campuran
        aslinya (mis. tabel field di awal, lalu paragraf narasi sesudahnya)
        hilang -- padahal parsing regex label->nilai di bawah bergantung
        pada urutan tampilan asli itu."""
        for child in document.element.body.iterchildren():
            if child.tag == qn("w:p"):
                yield Paragraph(child, document)
            elif child.tag == qn("w:tbl"):
                yield Table(child, document)

    doc = Document(docx_path)

    # PENTING: baca paragraf & tabel sesuai URUTAN ASLI kemunculannya di body
    # dokumen (bukan semua paragraf dulu baru semua tabel di akhir) --
    # parsing regex field di bawah ini mengasumsikan label & nilainya
    # berurutan seperti tampilan dokumen. Kalau tabel field (yang biasanya
    # ada di AWAL dokumen) malah ditumpuk ke akhir teks, regex label bisa
    # salah nyangkut ke kemunculan kata yang sama di kalimat/caption lain
    # sebelum tabel yang sesungguhnya ditemukan.
    parts = []
    for block in _iter_block_items(doc):
        if isinstance(block, Table):
            for row in block.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        else:
            parts.append(block.text)
    full_text_raw = "\n".join(parts)
    full_text = norm(full_text_raw)
    data = _parse_proposal_text(full_text, full_text_raw)

    # ---------------- IMAGES ----------------
    # Metode utama: cocokkan tiap gambar ke tag berdasar CAPTION ASLI yang
    # menempel di paragraf yang sama / 1-3 paragraf sesudahnya (konvensi
    # umum: "Gambar N. <deskripsi>" ditulis tepat di bawah gambar). Kalau
    # tidak ada caption yang cocok pola manapun, baru pakai fallback lama
    # (tebak dari section/heading terakhir yang terlihat).
    paragraphs = doc.paragraphs
    para_texts_norm = [norm(p.text) for p in paragraphs]

    seen_hash = set()
    images = []
    current_section = None
    section_img_count = 0

    for pi, para in enumerate(paragraphs):
        para_text_norm = para_texts_norm[pi]
        if para_text_norm:
            matches = []
            for pattern, tag in PROPOSAL_HEADINGS:
                for mm in re.finditer(pattern, para_text_norm, re.IGNORECASE):
                    matches.append((mm.start(), tag))
            if matches:
                matches.sort(key=lambda x: x[0])
                new_section = matches[-1][1]
                if new_section != current_section:
                    current_section = new_section
                    section_img_count = 0

        for run in para.runs:
            blips = run._element.findall(".//" + qn("a:blip"))
            for blip in blips:
                rId = blip.get(qn("r:embed"))
                if not rId:
                    continue
                try:
                    image_part = doc.part.related_parts[rId]
                except KeyError:
                    continue
                data_bytes = image_part.blob
                h = hashlib.md5(data_bytes).hexdigest()
                if h in seen_hash:
                    continue
                seen_hash.add(h)

                caption_tag = _nearby_caption_tag(para_texts_norm, pi, CAPTION_PATTERNS_PROPOSAL)

                if caption_tag:
                    tag = caption_tag
                else:
                    tag = _tag_proposal_image(current_section, section_img_count)
                    section_img_count += 1

                ext = image_part.content_type.split("/")[-1] if "/" in image_part.content_type else "png"
                images.append({"tag": tag, "bytes": data_bytes, "ext": ext, "width": 0, "height": 0, "page": 0})

    return data, images


# ----------------------------------------------------------------------
# LAPORAN HIDRO-OSEANOGRAFI PDF
# ----------------------------------------------------------------------

# Urutan gambar sesuai template (dikonfirmasi manual): gelombang, arus,
# pasut, ekosistem, batimetri. Dipakai karena resource gambar pada PDF ini
# dibagi/shared di semua halaman sehingga pencocokan caption per-halaman
# tidak selalu reliable; urutan xref lebih konsisten untuk template ini.
LAPORAN_IMAGE_ORDER = ["mawar_gelombang", "mawar_arus", "siklus_pasut", "peta_ekosistem", "profil_batimetri"]

# Heading/judul gambar -> tag, dicek di teks SEBELUM gambar muncul (per halaman
# untuk PDF, per paragraf untuk DOCX) supaya gambar ditandai sesuai konteks
# aslinya -- bukan cuma tebakan urutan tetap yang bisa keliru kalau urutan
# gambar di dokumen sumber ternyata berbeda dari yang diasumsikan.
LAPORAN_HEADINGS = [
    (r"mawar\s*gelombang", "mawar_gelombang"),
    (r"mawar\s*arus", "mawar_arus"),
    (r"(fluktuasi\s*)?pasang\s*surut", "siklus_pasut"),
    (r"siklus\s*pasut", "siklus_pasut"),
    (r"(peta\s*)?sebaran\s*(spasial\s*)?ekosistem", "peta_ekosistem"),
    (r"peta\s*ekosistem", "peta_ekosistem"),
    (r"batimetri", "profil_batimetri"),
    (r"kontur\s*batimetri", "profil_batimetri"),
]

# Heading section SEDERHANA bergaya "II. BATIMETRI", "III. GELOMBANG", dst
# (tanpa embel-embel "mawar"/"peta") -- banyak dipakai di laporan versi Word.
# HANYA dicocokkan ke baris PENDEK yang PERSIS berisi nama section itu saja
# (pakai ^...$), supaya tidak salah tangkap kalimat pembuka/pendahuluan yang
# menyebut banyak istilah sekaligus dalam satu kalimat panjang (mis. "...
# kondisi batimetri, dinamika gelombang, pola arus, ... serta sebaran
# ekosistem pesisir ..." -- kalimat seperti ini BUKAN heading, harus diabaikan).
LAPORAN_SECTION_HEADING_PATTERNS = [
    (r"^([ivx]+\.?\s*)?gelombang$", "mawar_gelombang"),
    (r"^([ivx]+\.?\s*)?arus$", "mawar_arus"),
    (r"^([ivx]+\.?\s*)?pasang\s*surut$", "siklus_pasut"),
    (r"^([ivx]+\.?\s*)?ekosistem(\s*pesisir)?$", "peta_ekosistem"),
    (r"^([ivx]+\.?\s*)?batimetri$", "profil_batimetri"),
]
LAPORAN_HEADING_MAX_LEN = 40  # baris lebih panjang dari ini dianggap kalimat/prosa, bukan heading


def resolve_laporan_tags(candidates):
    """Tentukan tag akhir tiap gambar Laporan Hidro-Oseanografi TANPA efek
    domino/tabrakan. 'candidates' adalah list dict per-gambar (urut sesuai
    posisi di dokumen), masing2 minimal berisi 'caption_norm' (teks caption
    ternormalisasi di dekat gambar itu, boleh kosong) dan 'heading_tag'
    (tag dari section/heading terdekat yang mendahului gambar itu secara
    posisi, boleh None).

    Cara lama menandai gambar SATU PER SATU secara berurutan: begitu satu
    caption gagal cocok ke pola manapun, gambar itu langsung "mencomot" tag
    berikutnya yang belum kepakai dari daftar urutan tetap -- ini memicu efek
    domino, SEMUA gambar setelahnya ikut bergeser salah tag (persis kasus
    nyata yang dilaporkan: grafik pasang surut tertandai sebagai peta
    ekosistem, dst). Fungsi ini menghindarinya lewat dua tahap terpisah:
    1) setiap TAG (sesuai prioritas) mengklaim kandidat gambar PERTAMA
       (urutan dokumen) yang caption-nya memang cocok ke tag itu;
    2) tag yang sampai tahap ini masih belum ada pengklaimnya, baru dicoba
       diisi dari heading_tag gambar yang tersisa.
    Gambar yang tetap tidak dapat tag setelah kedua tahap itu diberi
    'lainnya' -- BUKAN tag section lain yang salah.
    """
    n = len(candidates)
    result = [None] * n
    claimed = set()

    ranked = [
        _tag_from_caption_ranked(c.get("caption_norm", ""), CAPTION_PATTERNS_LAPORAN) for c in candidates
    ]
    best_caption_tag = [r[0] for r in ranked]
    specificity = [r[1] for r in ranked]

    for tag in LAPORAN_TAG_PRIORITY:
        # Di antara semua kandidat yang caption-nya cocok ke tag ini, pilih
        # yang PALING SPESIFIK dulu (rank pola terkecil); kalau seri, yang
        # muncul lebih dulu di dokumen menang -- BUKAN asal kandidat pertama
        # yang ditemukan, supaya caption yang lebih rinci/akurat diutamakan
        # dibanding caption umum yang kebetulan nongol lebih dulu.
        best_i = None
        for i in range(n):
            if result[i] is not None or best_caption_tag[i] != tag:
                continue
            if best_i is None or specificity[i] < specificity[best_i]:
                best_i = i
        if best_i is not None:
            result[best_i] = tag
            claimed.add(tag)

    for tag in LAPORAN_TAG_PRIORITY:
        if tag in claimed:
            continue
        for i in range(n):
            if result[i] is None and candidates[i].get("heading_tag") == tag:
                result[i] = tag
                claimed.add(tag)
                break

    for i in range(n):
        if result[i] is None:
            result[i] = "lainnya"
    return result


def _detect_laporan_tag(text_before, current_tag):
    """Cari heading yang paling terakhir muncul di teks sebelum gambar ini,
    untuk menentukan tag yang tepat. Kalau tidak ada heading baru yang cocok,
    pertahankan tag section terakhir yang diketahui (gambar lanjutan di
    section yang sama)."""
    matches = []
    for pattern, tag in LAPORAN_HEADINGS:
        for m in re.finditer(pattern, text_before, re.IGNORECASE):
            matches.append((m.start(), tag))
    if matches:
        matches.sort(key=lambda x: x[0])
        return matches[-1][1]
    return current_tag


# Heading penanda batas section untuk menangkap NARASI UTUH (bukan cuma
# angka) dari Laporan Hidro-Oseanografi -- supaya deskripsi/analisis asli
# yang ditulis surveyor tetap masuk ke dokumen final, bukan sekadar
# kalimat template generik yang cuma diisi angka.
#
# PENTING: heading dicocokkan terhadap SATU BARIS UTUH (bukan potongan
# substring di tengah kalimat manapun) -- karena beberapa dokumen sumber
# memakai penomoran ANGKA ROMAWI ("I. PENDAHULUAN", "II. BATIMETRI", "III.
# GELOMBANG", dst.), bukan cuma angka biasa ("1. Gelombang") atau huruf
# ("C. Profil Dasar Laut"). Mencocokkan di seluruh teks yang sudah
# digabung jadi satu paragraf besar (tanpa peduli baris) gampang salah
# nyangkut ke kalimat lain yang KEBETULAN menyinggung kata yang sama --
# padahal kalimat itu bukan heading section-nya sendiri (mis. paragraf
# Pendahuluan yang menyebut "...gelombang, arus, pasang surut, serta
# kondisi ekosistem pesisir" secara sekilas, ikut tertangkap sebagai
# section "ekosistem" walau itu cuma kalimat pengantar).
#
# "PENDAHULUAN" dan "PENUTUP" sengaja didaftarkan dengan tag None supaya
# section itu dikenali sebagai BATAS (supaya isinya tidak "bocor" ikut ke
# section lain), tapi kontennya sendiri TIDAK diambil sebagai narasi --
# sesuai permintaan: judul & Pendahuluan sumber tidak perlu ikut masuk ke
# dokumen final.
_NUM_PREFIX = r"(?:[IVXLCDM]{1,4}|\d{1,2}|[A-Z])"

NARASI_HEADING_RE = re.compile(
    r"^" + _NUM_PREFIX + r"\s*[.\-–:]?\s*"
    r"(GELOMBANG"
    r"|ARUS"
    r"|PASANG\s*SURUT|SIKLUS\s*PASUT"
    r"|PROFIL\s*(GARIS\s*)?BATIMETRI|PROFIL\s*DASAR\s*LAUT|BATIMETRI"
    r"|EKOSISTEM(\s*PESISIR)?"
    r"|PENDAHULUAN"
    r"|PENUTUP)"
    r"\s*$",
    re.IGNORECASE,
)

NARASI_TAG_BY_KEYWORD = [
    (r"gelombang", "gelombang"),
    (r"arus", "arus"),
    (r"pasang\s*surut|pasut", "pasut"),
    (r"batimetri|profil\s*dasar\s*laut", "batimetri"),
    (r"ekosistem", "ekosistem"),
    (r"pendahuluan", None),
    (r"penutup", None),
]

MAX_NARASI_LEN = 4000  # batas aman biar tidak "kebablasan" ambil isi dokumen kalau heading berikutnya tidak kedeteksi


def extract_narasi_sections(full_text_raw):
    """Cari baris heading section (lihat catatan di atas), lalu ambil semua
    teks di antara satu heading dengan heading berikutnya sebagai narasi
    utuh section itu. Mengembalikan dict {tag: teks_narasi}. Kalau heading
    tidak ditemukan sama sekali, dict yang dikembalikan kosong (pemanggil
    lalu pakai fallback kalimat template seperti sebelumnya)."""
    lines = (full_text_raw or "").split("\n")

    headings = []  # [(indeks_baris, tag_atau_None)]
    for i, line in enumerate(lines):
        line_norm = norm(line)
        if not line_norm:
            continue
        if NARASI_HEADING_RE.match(line_norm):
            tag = None
            for pattern, t in NARASI_TAG_BY_KEYWORD:
                if re.search(pattern, line_norm, re.IGNORECASE):
                    tag = t
                    break
            headings.append((i, tag))

    if not headings:
        return {}

    narasi = {}
    for idx, (line_i, tag) in enumerate(headings):
        next_line_i = headings[idx + 1][0] if idx + 1 < len(headings) else len(lines)
        chunk_lines = lines[line_i + 1: next_line_i]
        # Buang baris yang sebetulnya CAPTION gambar (mis. "Gambar: PETA
        # VISUAL - Ekosistem Pesisir"), bukan bagian narasi -- kalau tidak
        # disaring, label caption itu ikut "bocor" jadi teks di tengah
        # paragraf narasi hasil generate (pernah kejadian nyata: potongan
        # "Gambar: ANALISIS SPASIAL -- Tutupan Ekosistem Area" nyelip di
        # tengah kalimat narasi ekosistem).
        chunk_lines = [ln for ln in chunk_lines if not CAPTION_RE.match(norm(ln))]
        chunk = norm(" ".join(chunk_lines)).strip(" .:-")
        chunk = chunk[:MAX_NARASI_LEN].strip()
        # tag None (PENDAHULUAN/PENUTUP) sengaja TIDAK disimpan -- lihat
        # catatan di atas. Ambil kemunculan PERTAMA tiap tag yang valid.
        if tag and chunk and tag not in narasi:
            narasi[tag] = chunk
    return narasi


def _parse_laporan_text(full_text, full_text_raw):
    """Semua regex parsing Laporan Hidro-Oseanografi, dipakai bersama baik
    sumbernya PDF maupun DOCX -- karena keduanya sama-sama sudah berupa teks
    biasa pada titik ini, regex tidak peduli asalnya dari format file apa."""
    data = {}
    data["_narasi"] = extract_narasi_sections(full_text_raw)

    m = re.search(r"LOKASI TITIK PUSAT RENCANA KEGIATAN\s*(.+?)\s*I\.\s*PENDAHULUAN", full_text)
    data["lokasi_studi"] = norm(m.group(1)) if m else ""

    m = re.search(r"Kedalaman\s*pada\s*titik\s*pusat\s*tercatat\s*sebesar\s*(-?[\d.]+)\s*meter", full_text)
    data["batimetri_titik_pusat"] = m.group(1) if m else ""

    m = re.search(r"panjang\s*lintasan\s*([\d.]+)\s*kilometer", full_text)
    data["batimetri_panjang_lintasan"] = m.group(1) if m else ""

    m = re.search(r"nilai\s*terdalam\s*mencapai\s*(-?[\d.]+)\s*meter", full_text)
    data["batimetri_terdalam"] = m.group(1) if m else ""

    m = re.search(r"Tinggi\s*Gelombang\s*Signifikan\s*\(Hs\)\s*Rata-rata:\s*([\d.]+)\s*meter", full_text)
    data["hs_rata"] = m.group(1) if m else ""

    m = re.search(r"Maksimum\s*Ekstrem:\s*([\d.]+)\s*meter.*?arah\s*dominan\s*dari\s*([\d.]+)\s*derajat\s*\(sektor\s*barat", full_text)
    if m:
        data["hs_maks"], data["hs_arah"] = m.group(1), m.group(2)

    m = re.search(r"Kecepatan\s*Arus\s*Rata-rata:\s*([\d.]+)\s*meter\s*per\s*detik", full_text)
    data["arus_rata"] = m.group(1) if m else ""

    m = re.search(r"Kecepatan\s*Arus\s*Maksimum\s*Ekstrem:\s*([\d.]+)\s*meter\s*per\s*detik.*?arah\s*dominan\s*dari\s*([\d.]+)\s*derajat", full_text)
    if m:
        data["arus_maks"], data["arus_arah"] = m.group(1), m.group(2)

    m = re.search(r"Highest\s*Astronomical\s*Tide\s*\(HAT\):\s*\+?(-?[\d.]+)\s*meter", full_text)
    data["hat"] = m.group(1) if m else ""
    m = re.search(r"Mean\s*Sea\s*Level\s*\(MSL\):\s*(-?[\d.]+)\s*meter", full_text)
    data["msl"] = m.group(1) if m else ""
    m = re.search(r"Lowest\s*Astronomical\s*Tide\s*\(LAT\):\s*(-?[\d.]+)\s*meter", full_text)
    data["lat"] = m.group(1) if m else ""
    m = re.search(r"Tidal\s*Range:\s*([\d.]+)\s*meter", full_text)
    data["tidal_range"] = m.group(1) if m else ""
    m = re.search(r"Bilangan\s*Formzahl:\s*(\d+\.\d+)", full_text)
    data["formzahl"] = m.group(1) if m else ""
    tipe_dari_angka = klasifikasi_tipe_pasut(data["formzahl"])
    if tipe_dari_angka:
        data["tipe_pasut"] = tipe_dari_angka
    else:
        # fallback kalau angka Formzahl-nya sendiri tidak ketemu di teks --
        # coba cari dari kalimat naratif, baru kalau itu juga gagal pakai
        # default paling umum.
        tipe_m = re.search(r"diklasifikasikan.*?sebagai\s*\**\s*([A-Za-z ]+?)\**\s*,", full_text)
        data["tipe_pasut"] = norm(tipe_m.group(1)) if tipe_m else "Mixed Diurnal"

    m = re.search(r"area\s*rencana\s*kegiatan\s*seluas\s*([\d.]+)\s*Hektar", full_text)
    data["eko_total_ha"] = m.group(1) if m else ""

    m = re.search(r"Terumbu\s*Karang:\s*([\d.]+)\s*Hektar\s*\(([\d.]+)\s*persen", full_text)
    if m:
        data["eko_karang_ha"], data["eko_karang_pct"] = m.group(1), m.group(2)

    m = re.search(r"Lainnya\s*\(termasuk\s*substrat\s*dasar\s*non-terumbu\):\s*([\d.]+)\s*Hektar\s*\(([\d.]+)\s*persen", full_text)
    if m:
        data["eko_lainnya_ha"], data["eko_lainnya_pct"] = m.group(1), m.group(2)

    m = re.search(r"Area\s*Laut\s*Terbuka\s*\(Tanpa\s*Ekosistem\):\s*([\d.]+)\s*Hektar\s*\(([\d.]+)\s*persen", full_text)
    if m:
        data["eko_terbuka_ha"], data["eko_terbuka_pct"] = m.group(1), m.group(2)

    m = re.search(r"jarak\s*ekosistem\s*terdekat\s*dari\s*titik\s*pusat\s*rencana\s*kegiatan\s*adalah\s*([\d.]+)\s*kilometer", full_text)
    data["eko_jarak_terdekat_km"] = m.group(1) if m else ""

    data["ada_lamun"] = "padang lamun teridentifikasi" in full_text.lower()
    return data


def extract_laporan(pdf_path):
    doc = fitz.open(pdf_path)
    full_text_raw = "\n".join(page.get_text() for page in doc)
    full_text = norm(full_text_raw)
    data = _parse_laporan_text(full_text, full_text_raw)

    # ---------------- IMAGES ----------------
    # PENTING: heading dicatat berikut POSISI-nya (halaman + koordinat y),
    # bukan cuma "heading terakhir yang terlihat di teks satu halaman penuh".
    # Kalau tidak begitu, halaman yang memuat lebih dari satu heading (mis.
    # akhir section "Ekosistem" menyambung ke awal section "Gelombang" di
    # halaman yang sama) bisa membuat SEMUA gambar di halaman itu -- termasuk
    # peta ekosistem yang posisinya masih di atas heading "Gelombang" --
    # ikut tertandai sebagai gambar gelombang. Dengan posisi y, tiap gambar
    # dicocokkan ke heading TERDEKAT YANG BENAR-BENAR MENDAHULUINYA secara
    # vertikal di halaman yang sama, bukan heading manapun yang kebetulan
    # ada di halaman itu.
    heading_positions = []  # list of (page_num, y0, tag), terurut sesuai urutan baca
    for pnum in range(len(doc)):
        page = doc[pnum]
        try:
            blocks = page.get_text("blocks")
        except Exception:
            blocks = []
        for b in blocks:
            y0, text = b[1], (b[4] or "").strip()
            if not text:
                continue
            text_norm = norm(text)
            for pattern, tag in LAPORAN_HEADINGS:
                if re.search(pattern, text_norm, re.IGNORECASE):
                    heading_positions.append((pnum, y0, tag))
                    break
            else:
                if len(text_norm) <= LAPORAN_HEADING_MAX_LEN:
                    for pattern, tag in LAPORAN_SECTION_HEADING_PATTERNS:
                        if re.match(pattern, text_norm, re.IGNORECASE):
                            heading_positions.append((pnum, y0, tag))
                            break

    def tag_before(pnum, y0):
        """Tag dari heading terakhir yang posisinya sebelum (pnum, y0) dalam
        urutan baca dokumen (halaman lebih kecil, atau halaman sama dengan y
        lebih kecil)."""
        best = None
        for hp_pnum, hp_y0, hp_tag in heading_positions:
            if hp_pnum < pnum or (hp_pnum == pnum and hp_y0 <= y0):
                best = hp_tag
            elif hp_pnum > pnum:
                break
        return best

    seen_hash = set()
    candidates = []
    for pnum in range(len(doc)):
        page = doc[pnum]
        imglist = page.get_images(full=True)
        for i, img in enumerate(imglist):
            xref = img[0]
            base = doc.extract_image(xref)
            data_bytes = base["image"]
            h = hashlib.md5(data_bytes).hexdigest()
            w, ht = base.get("width", 0), base.get("height", 0)
            if w < 150 or ht < 150 or h in seen_hash:
                continue
            seen_hash.add(h)

            rects = page.get_image_rects(xref)
            img_rect = rects[0] if rects else None
            caption = _page_caption_for_image(page, img_rect)
            img_y0 = img_rect.y0 if img_rect is not None else 0
            candidates.append({
                "caption_norm": norm(caption) if caption else "",
                "heading_tag": tag_before(pnum, img_y0),
                "bytes": data_bytes, "ext": base["ext"],
                "width": w, "height": ht, "page": pnum + 1,
            })
    doc.close()

    tags = resolve_laporan_tags(candidates)
    images = []
    for c, tag in zip(candidates, tags):
        images.append({
            "tag": tag, "bytes": c["bytes"], "ext": c["ext"],
            "width": c["width"], "height": c["height"], "page": c["page"],
        })
    return data, images


def extract_laporan_docx(docx_path):
    """Versi ekstraksi Laporan Hidro-Oseanografi dari file Word (.docx).
    Dipakai sebagai alternatif PDF -- dokumen Word cenderung lebih stabil
    dibaca dibanding PDF (tidak ada masalah urutan/pemotongan kata saat
    di-convert ke teks), sehingga regex lebih jarang gagal."""
    from docx import Document

    doc = Document(docx_path)

    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    full_text_raw = "\n".join(parts)
    full_text = norm(full_text_raw)
    data = _parse_laporan_text(full_text, full_text_raw)

    # ---------------- IMAGES ----------------
    # Heading dilacak berikut POSISI-nya (indeks paragraf), lalu tiap gambar
    # dicocokkan ke heading TERDEKAT YANG BENAR-BENAR MENDAHULUINYA -- bukan
    # "heading section terakhir yang kebetulan pernah kelihatan" (rawan
    # keliru kalau ada paragraf pembuka yang menyebut banyak istilah section
    # sekaligus dalam satu kalimat, atau heading di paragraf lanjutan gagal
    # cocok sehingga current_tag jadi 'nyangkut' di section sebelumnya).
    from docx.oxml.ns import qn as _qn
    paragraphs = doc.paragraphs
    para_texts_norm = [norm(p.text) for p in paragraphs]

    heading_positions = []  # list of (para_idx, tag), terurut sesuai urutan baca
    for pi, text_norm in enumerate(para_texts_norm):
        if not text_norm:
            continue
        matched = False
        for pattern, tag in LAPORAN_HEADINGS:
            if re.search(pattern, text_norm, re.IGNORECASE):
                heading_positions.append((pi, tag))
                matched = True
                break
        if not matched and len(text_norm) <= LAPORAN_HEADING_MAX_LEN:
            for pattern, tag in LAPORAN_SECTION_HEADING_PATTERNS:
                if re.match(pattern, text_norm, re.IGNORECASE):
                    heading_positions.append((pi, tag))
                    break

    def tag_before(pi):
        best = None
        for hp_pi, hp_tag in heading_positions:
            if hp_pi <= pi:
                best = hp_tag
            else:
                break
        return best

    seen_hash = set()
    candidates = []
    for pi, para in enumerate(paragraphs):
        for run in para.runs:
            blips = run._element.findall(".//" + _qn("a:blip"))
            for blip in blips:
                rId = blip.get(_qn("r:embed"))
                if not rId:
                    continue
                try:
                    image_part = doc.part.related_parts[rId]
                except KeyError:
                    continue
                data_bytes = image_part.blob
                h = hashlib.md5(data_bytes).hexdigest()
                if h in seen_hash:
                    continue
                seen_hash.add(h)

                caption = None
                offsets = [0, 1, -1, 2, -2, 3, -3]
                for off in offsets:
                    idx = pi + off
                    if 0 <= idx < len(para_texts_norm) and para_texts_norm[idx]:
                        if CAPTION_RE.match(para_texts_norm[idx]) or off == 0:
                            caption = para_texts_norm[idx]
                            break

                ext = image_part.content_type.split("/")[-1] if "/" in image_part.content_type else "png"
                candidates.append({
                    "caption_norm": caption or "",
                    "heading_tag": tag_before(pi),
                    "bytes": data_bytes, "ext": ext, "width": 0, "height": 0, "page": 0,
                })

    tags = resolve_laporan_tags(candidates)
    images = []
    for c, tag in zip(candidates, tags):
        images.append({
            "tag": tag, "bytes": c["bytes"], "ext": c["ext"],
            "width": c["width"], "height": c["height"], "page": c["page"],
        })
    return data, images


if __name__ == "__main__":
    import sys
    import json

    prop_data, prop_imgs = extract_proposal(sys.argv[1])
    lap_data, lap_imgs = extract_laporan(sys.argv[2])
    print("=== PROPOSAL DATA ===")
    print(json.dumps({k: v for k, v in prop_data.items() if not k.startswith("_")}, indent=2, ensure_ascii=False))
    print("Images:", [(im["tag"], im["page"], im["width"], im["height"]) for im in prop_imgs])
    print("=== LAPORAN DATA ===")
    print(json.dumps(lap_data, indent=2, ensure_ascii=False))
    print("Images:", [(im["tag"], im["page"], im["width"], im["height"]) for im in lap_imgs])


# ----------------------------------------------------------------------
# WRAPPER DENGAN FALLBACK CLAUDE API
# ----------------------------------------------------------------------

def extract_proposal_with_fallback(pdf_path, use_llm=True, log=print):
    """Sama seperti extract_proposal(), tapi field yang gagal dibaca regex
    akan dicoba diisi ulang lewat Claude API (jika ANTHROPIC_API_KEY diset).
    Mendukung file sumber PDF maupun Word (.docx), dideteksi dari ekstensi."""
    import fitz
    from llm_fallback import llm_fill_missing_fields, PROPOSAL_FIELD_HINTS, api_key_available

    is_docx = pdf_path.lower().endswith(".docx")
    if is_docx:
        data, images = extract_proposal_docx(pdf_path)
    else:
        data, images = extract_proposal(pdf_path)

    missing = [k for k in PROPOSAL_FIELD_HINTS if not data.get(k)]
    if missing and use_llm:
        if api_key_available():
            log(f"      -> {len(missing)} field kosong, mencoba fallback Claude API: {missing}")
            if is_docx:
                from docx import Document
                d = Document(pdf_path)
                full_text = "\n".join(p.text for p in d.paragraphs)
            else:
                doc = fitz.open(pdf_path)
                full_text = "\n".join(p.get_text() for p in doc)
                doc.close()
            filled = llm_fill_missing_fields(full_text, missing, PROPOSAL_FIELD_HINTS)
            for k, v in filled.items():
                if v:
                    data[k] = v
            log(f"      -> berhasil melengkapi {len(filled)} field via Claude API.")
        else:
            log(f"      -> {len(missing)} field kosong, ANTHROPIC_API_KEY tidak diset (lewati fallback).")

    return data, images


def extract_laporan_with_fallback(pdf_path, use_llm=True, log=print):
    import fitz
    from llm_fallback import llm_fill_missing_fields, LAPORAN_FIELD_HINTS, api_key_available

    is_docx = pdf_path.lower().endswith(".docx")
    if is_docx:
        data, images = extract_laporan_docx(pdf_path)
    else:
        data, images = extract_laporan(pdf_path)

    missing = [k for k in LAPORAN_FIELD_HINTS if not data.get(k)]
    if missing and use_llm:
        if api_key_available():
            log(f"      -> {len(missing)} field kosong, mencoba fallback Claude API: {missing}")
            if is_docx:
                from docx import Document
                d = Document(pdf_path)
                full_text = "\n".join(p.text for p in d.paragraphs)
            else:
                doc = fitz.open(pdf_path)
                full_text = "\n".join(p.get_text() for p in doc)
                doc.close()
            filled = llm_fill_missing_fields(full_text, missing, LAPORAN_FIELD_HINTS)
            for k, v in filled.items():
                if v:
                    data[k] = v
            log(f"      -> berhasil melengkapi {len(filled)} field via Claude API.")
        else:
            log(f"      -> {len(missing)} field kosong, ANTHROPIC_API_KEY tidak diset (lewati fallback).")

    return data, images


def extract_full_text_any(file_path):
    """Ekstrak SELURUH teks polos (tanpa parsing field/regex apa pun) dari
    sebuah dokumen PDF, DOCX, atau XLSX -- dipakai untuk fitur Analisis
    Konsistensi Proposal, yang mengirim teks utuh dokumen ke Claude API untuk
    dibaca & dibandingkan langsung (bukan mengandalkan regex kaku, karena
    dokumen yang sudah jadi/ditulis manual bisa berformat sangat bervariasi)."""
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".docx":
        from docx import Document
        doc = Document(file_path)
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        return "\n".join(parts)
    elif ext in (".xlsx", ".xlsm"):
        from openpyxl import load_workbook
        wb = load_workbook(file_path, data_only=True, read_only=True)
        parts = []
        for ws in wb.worksheets:
            parts.append(f"[Sheet: {ws.title}]")
            for row in ws.iter_rows(values_only=True):
                cells = ["" if c is None else str(c) for c in row]
                if any(c.strip() for c in cells):
                    parts.append(" | ".join(cells))
        wb.close()
        return "\n".join(parts)
    else:
        doc = fitz.open(file_path)
        text = "\n".join(page.get_text() for page in doc)
        doc.close()
        return text


def extract_full_text_multi(file_paths):
    """Ekstrak & gabungkan teks dari BEBERAPA file (PDF/DOCX/XLSX) yang
    merupakan bagian dari satu dokumen/proposal yang sama. Tiap file diberi
    penanda nama sumbernya supaya temuan analisis bisa dirujuk ke file asal."""
    parts = []
    for fp in file_paths:
        nama = os.path.basename(fp)
        # buang prefix acak "proposal_0_" / "laporan_0_" dst kalau ada, sisakan nama asli
        teks = extract_full_text_any(fp)
        parts.append(f"\n\n===== BERKAS: {nama} =====\n{teks}")
    return "".join(parts).strip()
